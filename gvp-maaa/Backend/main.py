# pyre-ignore-all-errors
from fastapi.middleware.cors import CORSMiddleware
from fastapi import FastAPI, Depends, HTTPException, UploadFile, File, Form, Body, Query, Request  # type: ignore
from fastapi.middleware.cors import CORSMiddleware  # type: ignore
from io import BytesIO
from fastapi.responses import StreamingResponse  # type: ignore
from sqlalchemy.orm import Session  # type: ignore
from sqlalchemy import text, extract, func, or_, case  # type: ignore
from typing import Optional, List, Dict, Any
from security import hash_password, verify_password  # type: ignore
from mail import send_reset_email  # type: ignore
from schemas import  AlertCreate, AssignSubjectRequest, AttendanceCreate, ResetPasswordRequest, StudentPromotionRequest, TeacherAdminUpdate, TeacherDeleteRequest,TimetableCreate, TimetableResponse,StudentDeleteRequest,SubjectCreate,AttendanceAnalyticsResponse, MarksUpload, AdminOverviewResponse  # type: ignore
import schemas  # type: ignore
from datetime import datetime, timedelta
from database import engine  # type: ignore
try:
    from ml.risk_engine import calculate_risk_movement, calculate_risk_score, get_attendance_trend_label, get_student_attendance_trend, predict_future_risk
    from ml.prediction_engine import forecast_performance
except ImportError:
    def calculate_risk_movement(*args, **kwargs): return "stable"
    def calculate_risk_score(*args, **kwargs): return (0.0, [])
    def get_attendance_trend_label(*args, **kwargs): return "Fluctuating"
    def get_student_attendance_trend(*args, **kwargs): return [None, None, None, None, None]
    def predict_future_risk(*args, **kwargs): return 0
    def forecast_performance(*args, **kwargs): return 0.0
try:
    from ml.recommendation_engine import generate_recommendations
except ImportError:
    def generate_recommendations(*args, **kwargs): return []
from models import (  # type: ignore
    Base,
    Alert,
    AlertRecipient,
    Timetable,
    Subject,
    FacultySubject,
    Attendance,
    Student,
    User,
    AttendanceWarning,
    FacultyMonthlyAttendanceAlert,
    Assignment,
    AssignmentSubmission,
    Resource,
    ResourceAccess,
    Event,
    EventAttendance,
    EventRegistration,
    ExternalEventSubmission,
    SystemSetting,
    SettingsAuditLog,
    Mark,
    StudentProgress,
    TaskLog,
    PlacementCompany,
    PlacementDrive,
    PlacementFeedback,
    PlacementStudentProfile,
    StudentDrive,
)



from apscheduler.schedulers.background import BackgroundScheduler  # type: ignore


from reportlab.lib import colors  # type: ignore
from reportlab.lib.styles import getSampleStyleSheet  # type: ignore
from reportlab.lib.units import inch  # type: ignore
from reportlab.platypus import Table, TableStyle  # type: ignore
from fastapi.responses import FileResponse  # type: ignore
from reportlab.platypus import Table, TableStyle, Paragraph, Spacer, Image  # type: ignore
from reportlab.platypus import SimpleDocTemplate  # type: ignore
from reportlab.lib.styles import ParagraphStyle  # type: ignore
from reportlab.lib.styles import getSampleStyleSheet  # type: ignore
from reportlab.lib.pagesizes import A4  # type: ignore
from reportlab.lib import colors  # type: ignore
from reportlab.lib.units import inch  # type: ignore
from fastapi.responses import FileResponse  # type: ignore
from reportlab.pdfbase.pdfmetrics import stringWidth  # type: ignore
from reportlab.pdfgen import canvas  # type: ignore
from jose import JWTError, jwt  # type: ignore
from datetime import date, timedelta
from services.risk_engine import get_student_risk
from services.alert_rules import generate_student_alerts, NO_DATA_MESSAGE
from services.placement_engine import (
    generate_action_plan,
    get_company_eligibility,
    get_interview_insights,
    get_placement_readiness,
    get_selection_probability,
    get_skill_gap,
)



import pandas as pd  # type: ignore
import os
import json
import shutil
import uuid
import csv

DEPARTMENT_MAP = {
    11: "CSE",
    12: "CSM",
    14: "ECE",
    15: "MECH",
    1: "CIVIL"
}

BRANCH_TO_DEPARTMENT = {name: id for id, name in DEPARTMENT_MAP.items()}


def apply_student_filters(query, branch: Optional[str], year: Optional[int], semester: Optional[int], section: Optional[str]):
    if branch:
        dept_id = BRANCH_TO_DEPARTMENT.get(branch)
        if dept_id is not None:
            query = query.join(User, User.user_id == Student.student_id).filter(User.department_id == dept_id)
        else:
            try:
                query = query.join(User, User.user_id == Student.student_id).filter(User.department_id == int(branch))
            except Exception:
                query = query.filter(Student.section == branch)

    if year is not None:
        query = query.filter(Student.year == year)
    if semester is not None:
        query = query.filter(Student.semester == semester)
    if section:
        query = query.filter(Student.section == section)
    return query


def apply_alert_filters(query, branch: Optional[str], year: Optional[int], semester: Optional[int], section: Optional[str]):
    if branch or year is not None or semester is not None or section:
        query = query.join(Student, Student.student_id == Alert.student_id)

    if branch:
        dept_id = BRANCH_TO_DEPARTMENT.get(branch)
        if dept_id is not None:
            query = query.join(User, User.user_id == Student.student_id).filter(User.department_id == dept_id)
        else:
            query = query.filter(Alert.department == branch)

    if year is not None:
        query = query.filter(Student.year == year)
    if semester is not None:
        query = query.filter(Student.semester == semester)
    if section:
        query = query.filter(Student.section == section)
    return query


from dotenv import load_dotenv  # type: ignore
load_dotenv()
from database import SessionLocal  # type: ignore
from schemas import (  # type: ignore
    LoginRequest,
    StudentSignupRequest,
    TeacherSignupRequest,
    AdminLoginRequest,
    StudentProfileUpdate,
    FacultyProfileUpdate,
    AssignmentCreate,
    AssignmentResponse,
    AssignmentSubmissionCreate,
    AssignmentSubmissionResponse,
    AssignmentDetailResponse,
    StatusUpdateRequest,
    StudentAssignmentSummaryResponse,
    ResourceResponse,
    ResourceAccessRequest,
    EventCreate,
    EventResponse,
    EventAttendanceResponse,
    EventAttendanceUpdate,
    BulkEventAttendanceUpdate,
    EventAlertRequest,
    EventResultUpdate,
    EventRegistrationRequest,
    StudentEventResponse,
    ExternalEventSubmissionCreate,
    ExternalEventSubmissionResponse,
    FacultyExternalSubmissionDetail
)
from models import User, Student, Faculty,Timetable  # type: ignore
from auth import (  # type: ignore
    create_access_token,
    create_reset_token,
    verify_reset_token,
    get_current_user,
    SECRET_KEY,
    ALGORITHM   # ✅ ADD THIS
)






app = FastAPI(title="GVP Academic Analytics Backend")
Base.metadata.create_all(bind=engine)
# -------------------------
# CORS
# -------------------------
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

from fastapi.staticfiles import StaticFiles  # type: ignore

app.mount(
    "/uploads",
    StaticFiles(directory="uploads"),
    name="uploads"
)

settings_cache = {}
settings_cache_loaded = False


def refresh_settings_cache():
    global settings_cache, settings_cache_loaded
    db = SessionLocal()
    try:
        settings_cache = {setting.key: setting.value for setting in db.query(SystemSetting).all()}
        settings_cache_loaded = True
    except Exception as e:
        print("Failed to refresh settings cache:", e)
        settings_cache = {}
        settings_cache_loaded = False
    finally:
        db.close()


def ensure_student_insights_columns():
    db = SessionLocal()
    try:
        db.execute(text("""
            ALTER TABLE students
            ADD COLUMN IF NOT EXISTS intervention_status VARCHAR DEFAULT 'none',
            ADD COLUMN IF NOT EXISTS intervention_type VARCHAR,
            ADD COLUMN IF NOT EXISTS intervention_last_updated TIMESTAMP,
            ADD COLUMN IF NOT EXISTS previous_risk_score NUMERIC(5, 2);
        """))
        db.commit()
    except Exception as e:
        print("Student insights migration failed:", e)
        db.rollback()
    finally:
        db.close()


def ensure_placement_schema():
    db = SessionLocal()
    try:
        db.execute(text("""
            ALTER TABLE users
            ADD COLUMN IF NOT EXISTS role VARCHAR(50);
        """))

        db.execute(text("""
            ALTER TABLE companies
            ADD COLUMN IF NOT EXISTS role_type VARCHAR(50),
            ADD COLUMN IF NOT EXISTS created_at TIMESTAMP DEFAULT NOW();
        """))

        db.execute(text("""
            CREATE TABLE IF NOT EXISTS placement_drives (
                id SERIAL PRIMARY KEY,
                company_id INT REFERENCES companies(id) ON DELETE CASCADE,
                drive_date DATE,
                mode TEXT,
                branches JSONB,
                created_by INT REFERENCES users(user_id),
                created_at TIMESTAMP DEFAULT NOW()
            );
        """))

        db.execute(text("""
            CREATE TABLE IF NOT EXISTS student_drives (
                id SERIAL PRIMARY KEY,
                student_id INT REFERENCES students(student_id) ON DELETE CASCADE,
                drive_id INT REFERENCES placement_drives(id) ON DELETE CASCADE,
                status TEXT DEFAULT 'assigned',
                current_round INT DEFAULT 0,
                final_result TEXT DEFAULT 'pending',
                updated_at TIMESTAMP DEFAULT NOW(),
                UNIQUE(student_id, drive_id)
            );
        """))

        db.execute(text("""
            CREATE TABLE IF NOT EXISTS placement_feedback (
                id SERIAL PRIMARY KEY,
                student_id INT REFERENCES students(student_id) ON DELETE CASCADE,
                drive_id INT REFERENCES placement_drives(id) ON DELETE CASCADE,
                faculty_id INT REFERENCES users(user_id) ON DELETE SET NULL,
                comment TEXT,
                rating INT,
                created_at TIMESTAMP DEFAULT NOW()
            );
        """))

        db.commit()
    except Exception as e:
        print("Placement schema migration failed:", e)
        db.rollback()
    finally:
        db.close()


def get_optional_current_user(request: Request):
    auth_header = request.headers.get("Authorization")
    if not auth_header or not auth_header.lower().startswith("bearer "):
        return None

    token = auth_header.split(" ", 1)[1]
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        return payload
    except JWTError:
        return None


def get_current_setting(key: str):
    if settings_cache_loaded:
        return settings_cache.get(key)
    return get_setting(key)


def process_event_reminders():
    db = SessionLocal()
    try:
        today = date.today()
        target_date = today + timedelta(days=2)
        
        # Find events exactly 2 days out
        events_to_remind = db.query(Event).filter(Event.event_date == target_date).all()
        
        for event in events_to_remind:
            # Find all students in that year
            student_query = db.query(Student).filter(Student.year == event.year)
            
            # If a specific section is targeted (not 'All')
            if event.section and event.section != "All":
                student_query = student_query.filter(Student.section == event.section)
            
            students = student_query.all()
            
            title = f"Reminder: {event.title}"
            message = f"Reminder: The event '{event.title}' will be held on {event.event_date.strftime('%d %b %Y')} at {event.venue}. Don't miss it!"
            
            for s in students:
                # Check if reminder already sent to avoid duplicates (optional but good)
                # For now, simple create
                new_alert = Alert(
                    title=title,
                    message=message,
                    type="reminder",
                    target_role="student",
                    target_type="individual",
                    student_id=s.student_id,
                    faculty_id=event.created_by
                )
                db.add(new_alert)
                db.flush()
                db.add(AlertRecipient(alert_id=new_alert.id, user_id=s.student_id, is_read=False))
                
        db.commit()
    except Exception as e:
        print(f"Error processing event reminders: {e}")
    finally:
        db.close()


@app.on_event("startup")
def startup_event():
    Base.metadata.create_all(bind=engine)
    ensure_student_insights_columns()
    ensure_placement_schema()
    
    # Run automatic migrations for scaling_logs
    try:
        db = SessionLocal()
        db.execute(text("""
            ALTER TABLE scaling_logs
            ADD COLUMN IF NOT EXISTS file_name TEXT,
            ADD COLUMN IF NOT EXISTS snapshot_data JSON;
        """))

        default_settings = {
            "attendance_threshold": 75,
            "cgpa_threshold": 6.5,
            "attendance_alert_enabled": True,
            "cgpa_alert_enabled": True,
            "alert_frequency": "immediate",
            "report_retention_days": 30,
            "analytics_refresh_interval": "daily",
            "session_timeout": 30,
            "report_format": "PDF",
            "marks_format": None,
            "attendance_format": None,
            "assignment_format": None,
            "resources_format": None
        }

        existing_settings = {
            setting.key for setting in db.query(SystemSetting).all()
        }

        for key, value in default_settings.items():
            if key not in existing_settings:
                db.add(SystemSetting(key=key, value=value))

        db.commit()
        refresh_settings_cache()
    except Exception as e:
        print("Schema migration failed:", e)
    finally:
        db.close()

    scheduler = BackgroundScheduler()
    scheduler.add_job(process_event_reminders, "interval", hours=24)
    scheduler.start()




# -------------------------
# Database Dependency
# -------------------------
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


def get_setting(key: str):
    db = None
    try:
        if settings_cache_loaded:
            return settings_cache.get(key)
        db = SessionLocal()
        setting = db.query(SystemSetting).filter(SystemSetting.key == key).first()
        if not setting:
            return None
        return setting.value
    except Exception as e:
        print(f"[ERROR] get_setting failed for {key}: {e}")
        return None
    finally:
        if db:
            db.close()


def get_settings():
    try:
        if settings_cache_loaded:
            return dict(settings_cache)
        db = SessionLocal()
        settings = db.query(SystemSetting).all()
        return {setting.key: setting.value for setting in settings}
    except Exception:
        return {}
    finally:
        db.close()


def get_default_format_for_module(module_name: str):
    defaults = {
        "marks": "excel",
        "attendance": "pdf",
        "assignments": "pdf",
        "resources": "docx"
    }
    return defaults.get(module_name, "pdf")


def get_report_format(module_name: str):
    settings = get_settings()
    module_key = f"{module_name}_format"
    value = settings.get(module_key)
    if isinstance(value, str) and value.strip().lower() in {"pdf", "excel", "docx"}:
        return value.strip().lower()
    return get_default_format_for_module(module_name)


def _normalize_exam_name(exam_name: Optional[str]) -> str:
    if not exam_name:
        return ""
    return str(exam_name).strip().lower().replace("-", "").replace(" ", "")


def _to_float(value) -> float:
    try:
        if value is None:
            return 0.0
        return float(value)
    except Exception:
        return 0.0


def _compute_attendance_percentage(db: Session, student_id: int) -> float:
    rows = db.query(Attendance.status).filter(Attendance.student_id == student_id).all()
    total = len(rows)
    if total == 0:
        return 0.0
    present = sum(1 for row in rows if bool(row.status))
    return round((present / total) * 100, 2)


def _compute_marks_score(db: Session, student_id: int) -> float:
    rows = (
        db.query(Mark.exam, Mark.marks)
        .filter(Mark.student_id == student_id)
        .order_by(Mark.created_at.asc())
        .all()
    )

    mid1_values = []
    mid2_values = []
    assignment_values = []

    for row in rows:
        exam_key = _normalize_exam_name(row.exam)
        marks_value = _to_float(row.marks)
        if marks_value <= 0:
            continue

        if exam_key in ["mid1", "mid01"]:
            mid1_values.append(marks_value)
        elif exam_key in ["mid2", "mid02"]:
            mid2_values.append(marks_value)
        elif "assignment" in exam_key:
            assignment_values.append(marks_value)

    mid1_avg = sum(mid1_values) / len(mid1_values) if mid1_values else 0.0
    mid2_avg = sum(mid2_values) / len(mid2_values) if mid2_values else 0.0
    assignment_avg = sum(assignment_values) / len(assignment_values) if assignment_values else 0.0
    return round(mid1_avg + mid2_avg + assignment_avg, 2)


def _recent_attendance_improved(db: Session, student_id: int) -> bool:
    rows = (
        db.query(Attendance.attendance_date, Attendance.status)
        .filter(Attendance.student_id == student_id)
        .order_by(Attendance.attendance_date.asc())
        .all()
    )

    daily_values = {}
    for row in rows:
        day = row.attendance_date.isoformat() if row.attendance_date else None
        if not day:
            continue
        if day not in daily_values:
            daily_values[day] = {"present": 0, "total": 0}
        daily_values[day]["total"] += 1
        if bool(row.status):
            daily_values[day]["present"] += 1

    series = []
    for day in sorted(daily_values.keys()):
        total = daily_values[day]["total"]
        present = daily_values[day]["present"]
        if total > 0:
            series.append((present / total) * 100)

    if len(series) < 8:
        return False

    recent = series[-7:]
    previous = series[-14:-7] if len(series) >= 14 else series[:-7]
    if not previous:
        return False
    return (sum(recent) / len(recent)) > (sum(previous) / len(previous))


def _recent_marks_improved(db: Session, student_id: int) -> bool:
    rows = (
        db.query(Mark.created_at, Mark.marks)
        .filter(Mark.student_id == student_id)
        .order_by(Mark.created_at.asc())
        .all()
    )

    values = [_to_float(row.marks) for row in rows if _to_float(row.marks) > 0]
    if len(values) < 2:
        return False

    return values[-1] > values[-2]


def _compute_streak_from_logs(log_dates: List[date]) -> int:
    if not log_dates:
        return 0

    unique_dates = sorted(set(log_dates), reverse=True)
    streak = 0
    cursor = date.today()
    date_set = set(unique_dates)

    while cursor in date_set:
        streak += 1
        cursor = cursor - timedelta(days=1)

    return streak


def _sync_student_progress(db: Session, student_id: int) -> StudentProgress:
    today = date.today()

    attendance_bonus_exists = db.query(TaskLog).filter(
        TaskLog.student_id == student_id,
        TaskLog.task_id == "system-attendance-improvement"
    ).first()
    if not attendance_bonus_exists and _recent_attendance_improved(db, student_id):
        db.add(TaskLog(
            student_id=student_id,
            task_id="system-attendance-improvement",
            completed=True,
            verified=True,
            xp_earned=25,
            date=today,
        ))

    marks_bonus_exists = db.query(TaskLog).filter(
        TaskLog.student_id == student_id,
        TaskLog.task_id == "system-marks-improvement"
    ).first()
    if not marks_bonus_exists and _recent_marks_improved(db, student_id):
        db.add(TaskLog(
            student_id=student_id,
            task_id="system-marks-improvement",
            completed=True,
            verified=True,
            xp_earned=40,
            date=today,
        ))

    db.commit()

    task_logs = db.query(TaskLog).filter(TaskLog.student_id == student_id).all()
    total_xp = sum(int(log.xp_earned or 0) for log in task_logs)
    completed_dates = [log.date for log in task_logs if log.completed and log.date]
    streak_days = _compute_streak_from_logs(completed_dates)
    last_active = max(completed_dates) if completed_dates else None

    progress = db.query(StudentProgress).filter(StudentProgress.student_id == student_id).first()
    if not progress:
        progress = StudentProgress(student_id=student_id)
        db.add(progress)

    progress.total_xp = total_xp
    progress.streak_days = streak_days
    progress.last_active_date = last_active
    db.commit()
    db.refresh(progress)
    return progress


def _build_student_tasks(db: Session, student_id: int):
    student = db.query(Student).filter(Student.student_id == student_id).first()
    if not student:
        return {"today": [], "this_week": []}

    attendance_pct = _compute_attendance_percentage(db, student_id)
    marks_score = _compute_marks_score(db, student_id)

    submissions = db.query(AssignmentSubmission).filter(AssignmentSubmission.student_id == student_id).all()
    submitted_assignment_ids = {
        s.assignment_id for s in submissions if bool(s.is_submitted) or (s.status and s.status.lower() == "submitted")
    }

    assignments = (
        db.query(Assignment, Subject)
        .join(Subject, Subject.subject_id == Assignment.subject_id)
        .filter(
            Assignment.year == student.year,
            Assignment.section == student.section,
            Assignment.is_active == True,
        )
        .order_by(Assignment.due_date.asc())
        .all()
    )

    pending_assignment_pairs = [
        (assignment, subject)
        for assignment, subject in assignments
        if assignment.id not in submitted_assignment_ids
    ]

    today_tasks = []
    week_tasks = []

    if attendance_pct < 75:
        today_tasks.append({
            "id": "attendance-low",
            "title": "Attend all classes today",
            "type": "attendance",
            "priority": "HIGH",
            "verificationType": "system",
            "reason": f"Attendance is {attendance_pct:.2f}%. Minimum safe level is 75%.",
        })

    if marks_score < 40:
        week_tasks.append({
            "id": "marks-low",
            "title": "Prepare for next internal exam",
            "type": "marks",
            "priority": "HIGH",
            "verificationType": "system",
            "reason": "Current marks trend is below target range.",
        })

    for assignment, subject in pending_assignment_pairs[:3]:
        days_left = (assignment.due_date.date() - date.today()).days if assignment.due_date else 999
        task = {
            "id": f"assignment-{assignment.id}",
            "title": f"Complete assignment: {assignment.title}",
            "type": "assignment",
            "priority": "HIGH" if days_left <= 2 else "MEDIUM",
            "verificationType": "system",
            "reason": f"{subject.subject_name} assignment pending before deadline.",
        }
        if days_left <= 2:
            today_tasks.append(task)
        else:
            week_tasks.append(task)

    priority_rank = {"HIGH": 0, "MEDIUM": 1, "LOW": 2}
    today_tasks = sorted(today_tasks, key=lambda x: priority_rank.get(x["priority"], 9))[:2]
    week_tasks = sorted(week_tasks, key=lambda x: priority_rank.get(x["priority"], 9))[:3]

    existing_logs = db.query(TaskLog).filter(TaskLog.student_id == student_id).all()
    latest_by_task = {}
    for log in existing_logs:
        prev = latest_by_task.get(log.task_id)
        if not prev or (log.date and prev.date and log.date >= prev.date):
            latest_by_task[log.task_id] = log

    def with_status(task):
        log = latest_by_task.get(task["id"])
        task["completed"] = bool(log.completed) if log else False
        task["verified"] = bool(log.verified) if log else False
        task["xpAwarded"] = int(log.xp_earned or 0) if log else 0
        return task

    return {
        "today": [with_status(task) for task in today_tasks],
        "this_week": [with_status(task) for task in week_tasks],
    }


@app.get("/student/tasks/today")
def get_student_tasks_today(
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "student":
        raise HTTPException(status_code=403, detail="Student only")

    student_id = int(current_user["user_id"])
    return _build_student_tasks(db, student_id)


@app.post("/student/tasks/complete/{task_id}")
def complete_student_task(
    task_id: str,
    payload: schemas.TaskCompleteRequest,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "student":
        raise HTTPException(status_code=403, detail="Student only")

    student_id = int(current_user["user_id"])
    today = date.today()

    existing = db.query(TaskLog).filter(
        TaskLog.student_id == student_id,
        TaskLog.task_id == task_id,
        TaskLog.date == today,
    ).first()
    if existing:
        return {
            "task_id": task_id,
            "completed": bool(existing.completed),
            "verified": bool(existing.verified),
            "xp_earned": int(existing.xp_earned or 0),
        }

    base_xp = 10
    priority_bonus = 20 if str(payload.priority).upper() == "HIGH" else 0
    verified = False
    xp_earned = 0

    if payload.type == "assignment" and task_id.startswith("assignment-"):
        assignment_id = int(task_id.replace("assignment-", ""))
        submission = db.query(AssignmentSubmission).filter(
            AssignmentSubmission.student_id == student_id,
            AssignmentSubmission.assignment_id == assignment_id,
        ).first()
        verified = bool(submission and (submission.is_submitted or (submission.status and submission.status.lower() == "submitted")))
    elif payload.type == "attendance":
        verified = _compute_attendance_percentage(db, student_id) >= 75
    elif payload.type == "marks":
        verified = _compute_marks_score(db, student_id) >= 40
    elif payload.type == "event" and task_id.startswith("event-"):
        event_id = int(task_id.replace("event-", ""))
        reg = db.query(EventRegistration).filter(
            EventRegistration.student_id == student_id,
            EventRegistration.event_id == event_id,
        ).first()
        verified = bool(reg)
    elif payload.type == "study":
        verified = False

    if verified:
        xp_earned = base_xp + priority_bonus

    log = TaskLog(
        student_id=student_id,
        task_id=task_id,
        completed=True,
        verified=verified,
        xp_earned=xp_earned,
        date=today,
    )
    db.add(log)
    db.commit()

    generated_tasks = _build_student_tasks(db, student_id)
    today_task_ids = [task["id"] for task in generated_tasks.get("today", [])]
    all_today_done = True
    if today_task_ids:
        for generated_task_id in today_task_ids:
            item = db.query(TaskLog).filter(
                TaskLog.student_id == student_id,
                TaskLog.task_id == generated_task_id,
                TaskLog.date == today,
                TaskLog.completed == True,
            ).first()
            if not item:
                all_today_done = False
                break

    daily_bonus_exists = db.query(TaskLog).filter(
        TaskLog.student_id == student_id,
        TaskLog.task_id == f"system-daily-complete-{today.isoformat()}"
    ).first()

    if all_today_done and today_task_ids and not daily_bonus_exists:
        db.add(TaskLog(
            student_id=student_id,
            task_id=f"system-daily-complete-{today.isoformat()}",
            completed=True,
            verified=True,
            xp_earned=30,
            date=today,
        ))
        db.commit()

    progress = _sync_student_progress(db, student_id)

    return {
        "task_id": task_id,
        "completed": True,
        "verified": verified,
        "xp_earned": xp_earned,
        "total_xp": progress.total_xp,
    }


@app.get("/student/xp/{student_id}", response_model=schemas.StudentXpResponse)
def get_student_xp(
    student_id: int,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] == "student" and int(current_user["user_id"]) != student_id:
        raise HTTPException(status_code=403, detail="Cannot access another student progress")

    progress = _sync_student_progress(db, student_id)
    return {"total_xp": int(progress.total_xp or 0)}


@app.get("/student/streak/{student_id}", response_model=schemas.StudentStreakResponse)
def get_student_streak(
    student_id: int,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] == "student" and int(current_user["user_id"]) != student_id:
        raise HTTPException(status_code=403, detail="Cannot access another student progress")

    progress = _sync_student_progress(db, student_id)
    return {"streak_days": int(progress.streak_days or 0)}


@app.get("/class/leaderboard/{class_id}", response_model=List[schemas.LeaderboardItem])
def get_class_leaderboard(
    class_id: str,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    parts = class_id.split("-", 1)
    if len(parts) != 2:
        raise HTTPException(status_code=400, detail="Invalid class_id. Use '<year>-<section>'")

    try:
        class_year = int(parts[0])
    except ValueError as ex:
        raise HTTPException(status_code=400, detail="Invalid class year") from ex

    class_section = parts[1]

    class_students = (
        db.query(Student.student_id, User.name)
        .join(User, User.user_id == Student.student_id)
        .filter(Student.year == class_year, Student.section == class_section)
        .all()
    )

    if not class_students:
        return []

    student_ids = [row.student_id for row in class_students]
    xp_rows = (
        db.query(TaskLog.student_id, func.coalesce(func.sum(TaskLog.xp_earned), 0).label("xp"))
        .filter(TaskLog.student_id.in_(student_ids), TaskLog.verified == True)
        .group_by(TaskLog.student_id)
        .all()
    )

    xp_map = {row.student_id: int(row.xp or 0) for row in xp_rows}
    ranking = [
        {
            "student_id": row.student_id,
            "name": row.name,
            "xp": xp_map.get(row.student_id, 0),
        }
        for row in class_students
    ]

    ranking.sort(key=lambda item: item["xp"], reverse=True)

    for index, item in enumerate(ranking):
        item["rank"] = index + 1

    return ranking


@app.get("/admin/overview/attendance-trend", response_model=List[schemas.AdminOverviewTrendPoint])
def admin_overview_attendance_trend(
    branch: Optional[str] = Query(None),
    year: Optional[int] = Query(None),
    semester: Optional[int] = Query(None),
    section: Optional[str] = Query(None),
):
    db = SessionLocal()
    try:
        trend_points = []
        today = date.today()
        for i in range(6, -1, -1):
            day = today - timedelta(days=i)
            daily_query = db.query(Attendance).filter(Attendance.attendance_date == day)
            if branch or year is not None or semester is not None or section:
                daily_query = daily_query.join(Student, Student.student_id == Attendance.student_id)
                if branch:
                    dept_id = BRANCH_TO_DEPARTMENT.get(branch)
                    if dept_id is not None:
                        daily_query = daily_query.join(User, User.user_id == Student.student_id).filter(User.department_id == dept_id)
                if year is not None:
                    daily_query = daily_query.filter(Student.year == year)
                if semester is not None:
                    daily_query = daily_query.filter(Student.semester == semester)
                if section:
                    daily_query = daily_query.filter(Student.section == section)
            total_day = daily_query.count()
            if total_day > 0:
                present_day = daily_query.filter(Attendance.status == True).count()
                trend_points.append({
                    "date": day.strftime("%b %d"),
                    "attendance": round((present_day / total_day) * 100, 2)
                })
        return trend_points
    except Exception as e:
        print(f"Failed to build attendance trend: {e}")
        return []
    finally:
        db.close()


@app.get("/admin/overview", response_model=AdminOverviewResponse)
def admin_overview(
    branch: Optional[str] = Query(None),
    year: Optional[int] = Query(None),
    semester: Optional[int] = Query(None),
    section: Optional[str] = Query(None),
):
    db = SessionLocal()
    try:
        attendance_threshold = get_setting("attendance_threshold") or 75
        cgpa_threshold = get_setting("cgpa_threshold") or 6.5

        student_query = db.query(Student)
        if branch or year is not None or semester is not None or section:
            student_query = apply_student_filters(student_query, branch, year, semester, section)

        total_students = student_query.count()
        teacher_query = db.query(FacultySubject.faculty_id).distinct()
        if branch or year is not None or semester is not None or section:
            if branch:
                dept_id = BRANCH_TO_DEPARTMENT.get(branch)
                if dept_id is not None:
                    teacher_query = teacher_query.join(Subject, Subject.subject_id == FacultySubject.subject_id).filter(Subject.department_id == dept_id)
            if year is not None:
                teacher_query = teacher_query.filter(FacultySubject.year == year)
            if semester is not None:
                teacher_query = teacher_query.join(Subject, Subject.subject_id == FacultySubject.subject_id).filter(Subject.semester == semester)
            if section:
                teacher_query = teacher_query.filter(FacultySubject.section == section)
            total_teachers = teacher_query.count()
        else:
            total_teachers = db.query(Faculty).count()

        student_rows = student_query.all()
        at_risk_students = 0
        attendance_percentages = []
        low_attendance_count = 0

        for student in student_rows:
            total = db.query(Attendance).filter(Attendance.student_id == student.student_id).count()
            present = db.query(Attendance).filter(
                Attendance.student_id == student.student_id,
                Attendance.status == True
            ).count()
            percentage = round((present / total) * 100, 2) if total > 0 else 0.0
            if total > 0:
                attendance_percentages.append(percentage)
            student_cgpa = float(student.cgpa) if student.cgpa is not None else 0.0
            if percentage < attendance_threshold or student_cgpa < cgpa_threshold:
                at_risk_students += 1
            if percentage < attendance_threshold:
                low_attendance_count += 1

        avg_attendance = round(sum(attendance_percentages) / len(attendance_percentages), 2) if attendance_percentages else 0.0
        avg_cgpa_raw = student_query.with_entities(func.avg(Student.cgpa)).scalar() or 0.0
        avg_cgpa = round(float(avg_cgpa_raw), 2)
        attendance_risk_percent = round((at_risk_students / total_students) * 100, 2) if total_students > 0 else 0.0

        faculty_class_counts = db.query(
            FacultySubject.faculty_id,
            func.count(FacultySubject.id).label("class_count")
        ).group_by(FacultySubject.faculty_id)
        if branch or year is not None or semester is not None or section:
            if branch:
                dept_id = BRANCH_TO_DEPARTMENT.get(branch)
                if dept_id is not None:
                    faculty_class_counts = faculty_class_counts.join(Subject, Subject.subject_id == FacultySubject.subject_id).filter(Subject.department_id == dept_id)
            if year is not None:
                faculty_class_counts = faculty_class_counts.filter(FacultySubject.year == year)
            if semester is not None:
                faculty_class_counts = faculty_class_counts.join(Subject, Subject.subject_id == FacultySubject.subject_id).filter(Subject.semester == semester)
            if section:
                faculty_class_counts = faculty_class_counts.filter(FacultySubject.section == section)
        faculty_class_counts = faculty_class_counts.all()
        teacher_counts = {faculty_id: class_count for faculty_id, class_count in faculty_class_counts}
        overloaded = sum(1 for count in teacher_counts.values() if count > 8)
        underutilized = sum(1 for count in teacher_counts.values() if count < 3)
        avg_classes = round(sum(teacher_counts.values()) / len(teacher_counts), 2) if teacher_counts else 0.0

        alert_query = db.query(Alert).filter(
            Alert.created_at >= datetime.utcnow() - timedelta(days=7)
        )
        if branch or year is not None or semester is not None or section:
            alert_query = alert_query.join(Student, Student.student_id == Alert.student_id)
            if branch:
                dept_id = BRANCH_TO_DEPARTMENT.get(branch)
                if dept_id is not None:
                    alert_query = alert_query.join(User, User.user_id == Student.student_id).filter(User.department_id == dept_id)
                else:
                    alert_query = alert_query.filter(Alert.department == branch)
            if year is not None:
                alert_query = alert_query.filter(Student.year == year)
            if semester is not None:
                alert_query = alert_query.filter(Student.semester == semester)
            if section:
                alert_query = alert_query.filter(Student.section == section)

        active_alerts = alert_query.count()
        alert_records = alert_query.order_by(Alert.created_at.desc()).limit(5).all()

        event_query = db.query(Event)
        if year is not None:
            event_query = event_query.filter(Event.year == str(year))
        if section:
            event_query = event_query.filter(Event.section == section)
        today = date.today()
        active_events = event_query.filter(
            or_(Event.status == "ongoing", Event.event_date >= today)
        ).count()
        events_today = event_query.filter(Event.event_date == today).count()
        week_end = today + timedelta(days=7)
        events_this_week = event_query.filter(Event.event_date >= today, Event.event_date <= week_end).count()

        attendance_trend = []
        today = date.today()
        for i in range(6, -1, -1):
            day = today - timedelta(days=i)
            daily_query = db.query(Attendance).filter(Attendance.attendance_date == day)
            if branch or year is not None or semester is not None or section:
                daily_query = daily_query.join(Student, Student.student_id == Attendance.student_id)
                if branch:
                    dept_id = BRANCH_TO_DEPARTMENT.get(branch)
                    if dept_id is not None:
                        daily_query = daily_query.join(User, User.user_id == Student.student_id).filter(User.department_id == dept_id)
                if year is not None:
                    daily_query = daily_query.filter(Student.year == year)
                if semester is not None:
                    daily_query = daily_query.filter(Student.semester == semester)
                if section:
                    daily_query = daily_query.filter(Student.section == section)
            total_day = daily_query.count()
            present_day = daily_query.filter(Attendance.status == True).count()
            attendance_trend.append({
                "date": day.strftime("%Y-%m-%d"),
                "attendance": round((present_day / total_day) * 100, 2) if total_day > 0 else 0.0
            })

        alerts = []
        if attendance_risk_percent >= 25:
            alerts.append({
                "title": "Large attendance risk cluster detected",
                "type": "academic",
                "severity": "high",
                "timestamp": datetime.utcnow().isoformat(),
                "action": "View Students"
            })
        if at_risk_students > 0 and avg_cgpa < cgpa_threshold:
            alerts.append({
                "title": "CGPA decline detected across students",
                "type": "academic",
                "severity": "medium",
                "timestamp": datetime.utcnow().isoformat(),
                "action": "View Students"
            })
        if overloaded > 0:
            alerts.append({
                "title": "Faculty overload identified",
                "type": "faculty",
                "severity": "high",
                "timestamp": datetime.utcnow().isoformat(),
                "action": "Assign Faculty"
            })
        if low_attendance_count > 0 and avg_attendance < attendance_threshold:
            alerts.append({
                "title": "Timetable review recommended",
                "type": "system",
                "severity": "low",
                "timestamp": datetime.utcnow().isoformat(),
                "action": "Fix Timetable"
            })

        filtered_student_ids = [student.student_id for student in student_rows]
        attendance_student_ids = set(id for (id,) in db.query(Attendance.student_id).filter(Attendance.student_id.in_(filtered_student_ids)).distinct().all()) if filtered_student_ids else set()
        marks_student_ids = set(id for (id,) in db.query(Mark.student_id).filter(Mark.student_id.in_(filtered_student_ids)).distinct().all()) if filtered_student_ids else set()
        students_with_both = len(attendance_student_ids & marks_student_ids)

        data_completeness = round((students_with_both / total_students) * 100, 2) if total_students > 0 else 0.0

        return {
            "metrics": {
                "at_risk_students": at_risk_students,
                "attendance_risk_percent": attendance_risk_percent,
                "data_completeness": data_completeness,
                "active_alerts": active_alerts,
                "total_students": total_students,
                "total_teachers": total_teachers,
                "active_events": active_events,
                "events_today": events_today,
                "events_this_week": events_this_week,
            },
            "academic_health": {
                "avg_attendance": avg_attendance,
                "avg_cgpa": avg_cgpa,
                "at_risk_students": at_risk_students,
            },
            "faculty_health": {
                "avg_classes": avg_classes,
                "overloaded": overloaded,
                "underutilized": underutilized,
            },
            "system_health": {
                "active_users": db.query(Attendance).filter(Attendance.attendance_date == today).count(),
                "last_sync": datetime.utcnow().isoformat(),
                "data_completeness": data_completeness,
            },
            "alerts": [
                {
                    "title": alert.title,
                    "type": "system",
                    "severity": "low",
                    "timestamp": alert.created_at.isoformat() if alert.created_at else datetime.utcnow().isoformat(),
                    "action": "View Alerts"
                }
                for alert in alert_records
            ] + alerts,
            "trend": attendance_trend,
        }
    except Exception as e:
        print(f"Failed to build admin overview: {e}")
        return {
            "metrics": {
                "at_risk_students": 0,
                "attendance_risk_percent": 0.0,
                "faculty_overload": 0,
                "active_alerts": 0,
                "total_students": 0,
                "total_teachers": 0,
            },
            "academic_health": {
                "avg_attendance": 0.0,
                "avg_cgpa": 0.0,
                "at_risk_students": 0,
            },
            "faculty_health": {
                "avg_classes": 0.0,
                "overloaded": 0,
                "underutilized": 0,
            },
            "system_health": {
                "active_users": 0,
                "last_sync": datetime.utcnow().isoformat(),
                "data_completeness": 0.0,
            },
            "alerts": [],
            "trend": [],
        }
    finally:
        db.close()

    generic = settings.get("report_format")
    if isinstance(generic, str) and generic.strip().lower() in {"pdf", "excel", "docx"}:
        return generic.strip().lower()

    return get_default_format_for_module(module_name)


def build_attendance_trend(records, view: str = "daily"):
    view_key = (view or "daily").strip().lower()
    if view_key not in {"daily", "weekly", "monthly"}:
        view_key = "daily"

    if view_key == "daily":
        sorted_records = sorted(records, key=lambda r: r.attendance_date)
        return [
            {
                "date": r.attendance_date.strftime("%Y-%m-%d"),
                "percentage": 100 if r.status else 0,
            }
            for r in sorted_records
        ]

    grouped = {}
    for record in records:
        if view_key == "weekly":
            year, week_num, _ = record.attendance_date.isocalendar()
            key = (year, week_num)
            label = f"Week {week_num}"
        else:
            year = record.attendance_date.year
            month_index = record.attendance_date.month
            key = (year, month_index)
            label = record.attendance_date.strftime("%b")

        bucket = grouped.get(key)
        if bucket is None:
            bucket = {"date": label, "present": 0, "total": 0}
            grouped[key] = bucket

        bucket["present"] += 1 if record.status else 0
        bucket["total"] += 1

    sorted_items = sorted(grouped.items(), key=lambda item: item[0])
    trend = []
    for _, bucket in sorted_items:
        percentage = round((bucket["present"] / bucket["total"]) * 100, 2) if bucket["total"] > 0 else 0
        trend.append({"date": bucket["date"], "percentage": percentage})

    return trend

# -------------------------
# Root Check
# -------------------------
@app.get("/")
def root():
    return {"message": "Backend connected to database successfully"}

# -------------------------
# SYSTEM SETTINGS
@app.get("/api/settings")
def get_system_settings(db: Session = Depends(get_db)):
    if settings_cache_loaded:
        payload = dict(settings_cache)
    else:
        settings = db.query(SystemSetting).all()
        payload = {setting.key: setting.value for setting in settings}

    last_updated = None
    for key, value in payload.items():
        if key == "settings_last_updated":
            continue
        setting = db.query(SystemSetting).filter(SystemSetting.key == key).first()
        if setting and setting.updated_at is not None:
            if last_updated is None or setting.updated_at > last_updated:
                last_updated = setting.updated_at

    if last_updated is not None:
        payload["settings_last_updated"] = last_updated.isoformat()

    return payload


@app.post("/api/settings/preview-impact")
def preview_settings_impact(data: schemas.SettingsPreviewRequest, db: Session = Depends(get_db)):
    attendance_threshold = data.attendance_threshold
    cgpa_threshold = data.cgpa_threshold

    students = db.query(Student).filter(Student.is_deleted == False).all()
    at_risk = 0

    for student in students:
        total = db.query(Attendance).filter(Attendance.student_id == student.student_id).count()
        present = db.query(Attendance).filter(
            Attendance.student_id == student.student_id,
            Attendance.status == True
        ).count()
        percentage = round((present / total) * 100, 2) if total > 0 else 0
        student_cgpa = float(student.cgpa) if student.cgpa is not None else 0.0

        if percentage < attendance_threshold or student_cgpa < cgpa_threshold:
            at_risk += 1

    return {"at_risk_students": at_risk}


@app.get("/api/settings/logs")
def get_settings_logs(db: Session = Depends(get_db)):
    logs = db.query(SettingsAuditLog).order_by(SettingsAuditLog.timestamp.desc()).limit(10).all()
    return [
        {
            "id": log.id,
            "key": log.key,
            "old_value": log.old_value,
            "new_value": log.new_value,
            "updated_by": log.updated_by,
            "timestamp": log.timestamp.isoformat() if log.timestamp else None,
        }
        for log in logs
    ]


@app.put("/api/settings")
def update_system_settings(data: schemas.SettingsUpdateRequest, db: Session = Depends(get_db), current_user: dict = Depends(get_optional_current_user)):
    payload = data.dict(exclude_unset=True)
    if not payload:
        raise HTTPException(status_code=400, detail="No settings provided")

    allowed_options = {
        "attendance_threshold": lambda value: isinstance(value, (int, float)) and 0 <= value <= 100,
        "cgpa_threshold": lambda value: isinstance(value, (int, float)) and 0 <= value <= 10,
        "attendance_alert_enabled": lambda value: isinstance(value, bool),
        "cgpa_alert_enabled": lambda value: isinstance(value, bool),
        "alert_frequency": lambda value: isinstance(value, str) and value.strip() in {"immediate", "daily", "weekly"},
        "report_retention_days": lambda value: isinstance(value, int) and value > 0,
        "analytics_refresh_interval": lambda value: isinstance(value, str) and value.strip() in {"daily", "weekly", "monthly"},
        "session_timeout": lambda value: isinstance(value, int) and value > 0,
        "report_format": lambda value: value is None or (isinstance(value, str) and value.strip().lower() in {"pdf", "excel", "docx"}),
        "marks_format": lambda value: value is None or (isinstance(value, str) and value.strip().lower() in {"pdf", "excel", "docx"}),
        "attendance_format": lambda value: value is None or (isinstance(value, str) and value.strip().lower() in {"pdf", "excel", "docx"}),
        "assignment_format": lambda value: value is None or (isinstance(value, str) and value.strip().lower() in {"pdf", "excel", "docx"}),
        "resources_format": lambda value: value is None or (isinstance(value, str) and value.strip().lower() in {"pdf", "excel", "docx"})
    }

    audit_entries = []
    user_label = "system"
    if current_user and isinstance(current_user, dict):
        user_label = str(current_user.get("user_id") or current_user.get("email") or "system")

    for key, value in payload.items():
        validator = allowed_options.get(key)
        if validator is None:
            raise HTTPException(status_code=400, detail=f"Invalid setting key: {key}")
        if not validator(value):
            raise HTTPException(status_code=400, detail=f"Invalid value for {key}")

        setting = db.query(SystemSetting).filter(SystemSetting.key == key).first()
        old_value = setting.value if setting else None

        if setting is None:
            setting = SystemSetting(key=key, value=value)
            db.add(setting)
        else:
            setting.value = value

        audit_entries.append(SettingsAuditLog(
            key=key,
            old_value=old_value,
            new_value=value,
            updated_by=user_label,
            timestamp=datetime.utcnow()
        ))

    db.add_all(audit_entries)
    db.commit()
    refresh_settings_cache()
    return {"message": "Settings updated successfully", "updated_at": datetime.utcnow().isoformat()}


# -------------------------
# LOGIN (Student / Teacher / Admin)
# -------------------------
@app.post("/login")
def login(data: LoginRequest, db: Session = Depends(get_db)):
    print("\n=== LOGIN ATTEMPT ===")
    print("Incoming email:", data.email)
    print("Incoming password length:", len(data.password))
    
    user = db.query(User).filter(User.email == data.email).first()
    print("User found:", user is not None)

    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    password_result = verify_password(data.password, user.password)
    print("Password verification result:", password_result)
    
    # Handle three cases: True (valid), "upgrade" (valid but needs rehashing), False (invalid)
    if password_result == "upgrade":
        # 🔄 AUTO-UPGRADE: Old password format detected
        # Rehash using new SHA256+bcrypt method and save
        print("[!] Upgrading password hash for user:", user.email)
        try:
            user.password = hash_password(data.password)
            db.commit()
            print("[✓] Password upgraded successfully")
        except Exception as e:
            print(f"[!] Password upgrade error: {e}")
            db.rollback()
            # Continue login anyway - upgrade failed but password was valid
    elif password_result != True:
        # Password is invalid
        raise HTTPException(status_code=401, detail="Invalid password")

    # 🔐 CREATE JWT TOKEN (both new and upgraded passwords reach here)
    access_token = create_access_token(
        data={
            "user_id": user.user_id,
            "role": user.role,
            "department_id": user.department_id
        }
    )

    return {
        "access_token": access_token,
        "token_type": "bearer",
        "user_id": user.user_id,
        "name": user.name,
        "role": user.role,
        "department_id": user.department_id
    }


# -------------------------
# STUDENT SIGNUP
# -------------------------
@app.post("/signup/student")
def student_signup(data: StudentSignupRequest, db: Session = Depends(get_db)):

    if not data.email.endswith("@gvpcdpgc.edu.in"):
        raise HTTPException(status_code=400, detail="Only college email allowed")

    if len(data.roll_no) < 6:
        raise HTTPException(status_code=400, detail="Invalid roll number")

    # extract department id (12 from roll number)
    joining_year = int(data.roll_no[2:4])
    department_id = int(data.roll_no[5:7])

    existing_user = db.query(User).filter(User.email == data.email).first()
    if existing_user:
        raise HTTPException(status_code=400, detail="Email already registered")

    # 1️⃣ Create user
    new_user = User(
    name=data.name,
    email=data.email,
    password=hash_password(data.password),
    role="student",
    department_id=department_id
  )


    db.add(new_user)
    db.commit()
    db.refresh(new_user)

    # 2️⃣ Auto-create student profile
    student = Student(
    student_id=new_user.user_id,
    roll_no=data.roll_no,
    joining_year=joining_year,   # ✅ ADD THIS
    year=1,
    semester=1,
    section=None,
    cgpa=0.00
 )


    db.add(student)
    db.commit()

    return {
        "message": "Student signup successful",
        "user_id": new_user.user_id
    }



# -------------------------
# STUDENT PROFILE GET
# -------------------------
@app.get("/student/profile")
def get_student_profile(
    user=Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if user["role"] != "student":
        raise HTTPException(status_code=403, detail="Not authorized")

    student = (
        db.query(Student, User)
        .join(User, Student.student_id == User.user_id)
        .filter(User.user_id == user["user_id"])
        .first()
    )

    if not student:
        raise HTTPException(status_code=404, detail="Student not found")

    student_data, user_data = student

    return {
    "name": user_data.name,
    "email": user_data.email,
    "roll_no": student_data.roll_no,
    "year": student_data.year,
    "semester": student_data.semester,
    "section": student_data.section,
    "class_id": f"{student_data.year}-{student_data.section}" if student_data.year and student_data.section else None,
    "skills": student_data.skills.split(",") if student_data.skills else [],
    "certificates": json.loads(student_data.certificates)
        if student_data.certificates else [],
    "linkedin": student_data.linkedin,
    "github": student_data.github,
    "portfolio": student_data.portfolio,
 }

# -------------------------
# STUDENT PROFILE PUT
# -------------------------
@app.put("/student/profile")
def update_student_profile(
    data: StudentProfileUpdate,
    user=Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if user["role"] != "student":
        raise HTTPException(status_code=403, detail="Not authorized")

    student = db.query(Student).filter(
        Student.student_id == user["user_id"]
    ).first()

    if not student:
        raise HTTPException(status_code=404, detail="Student not found")

    # BASIC INFO
    student.year = data.year
    student.semester = data.semester
    student.linkedin = data.linkedin
    student.github = data.github
    student.portfolio = data.portfolio

    # SKILLS (list → string)
    student.skills = ",".join(data.skills)

    # ✅ CERTIFICATES (list of objects → JSON string)
    student.certificates = json.dumps(
    [c.dict() for c in data.certificates]
 )


    db.commit()

    return {"message": "Profile updated successfully"}


def _require_placement_student(student_id: int, current_user: dict):
    if current_user["role"] != "student" or int(current_user["user_id"]) != int(student_id):
        raise HTTPException(status_code=403, detail="Student only")


def _normalized_role(current_user: dict) -> str:
    role = str(current_user.get("role") or "").strip().lower()
    if role == "faculty":
        return "teacher"
    return role


def authorize(roles: List[str]):
    allowed_roles = {str(role).strip().lower() for role in roles}

    def _checker(current_user: dict = Depends(get_current_user)):
        if _normalized_role(current_user) not in allowed_roles:
            raise HTTPException(status_code=403, detail="Access denied")
        return current_user

    return _checker


def _safe_float(value) -> float:
    try:
        if value is None:
            return 0.0
        return float(value)
    except Exception:
        return 0.0


def _get_student_backlogs(db: Session, student_id: int) -> int:
    profile = db.query(PlacementStudentProfile).filter(PlacementStudentProfile.student_id == student_id).first()
    if profile and profile.backlogs is not None:
        return int(profile.backlogs)
    return 0


def _branch_allowed(student_id: int, branches: List[str], db: Session) -> bool:
    if not branches:
        return True
    user = db.query(User).filter(User.user_id == student_id).first()
    if not user:
        return False
    student_branch = DEPARTMENT_MAP.get(user.department_id, str(user.department_id))
    return student_branch.upper() in {str(branch).strip().upper() for branch in branches}


@app.get("/placement/readiness/{student_id}")
def placement_readiness(
    student_id: int,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    _require_placement_student(student_id, current_user)
    return get_placement_readiness(student_id, db)


@app.get("/placement/eligibility/{student_id}")
def placement_eligibility(
    student_id: int,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    _require_placement_student(student_id, current_user)
    return get_company_eligibility(student_id, db)


@app.get("/placement/skills/{student_id}")
def placement_skills(
    student_id: int,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    _require_placement_student(student_id, current_user)
    return get_skill_gap(student_id, db)


@app.get("/placement/interviews/{student_id}")
def placement_interviews(
    student_id: int,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    _require_placement_student(student_id, current_user)
    return get_interview_insights(student_id, db)


@app.get("/placement/prediction/{student_id}")
def placement_prediction(
    student_id: int,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    _require_placement_student(student_id, current_user)
    return get_selection_probability(student_id, db)


@app.get("/placement/action-plan/{student_id}")
def placement_action_plan(
    student_id: int,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    _require_placement_student(student_id, current_user)
    return generate_action_plan(student_id, db)


# -------------------------
# PLACEMENT MANAGEMENT APIs
# -------------------------
@app.post("/api/companies")
def create_company(
    payload: Dict[str, Any] = Body(...),
    _: dict = Depends(authorize(["admin"])),
    db: Session = Depends(get_db),
):
    company_name = str(payload.get("name") or "").strip()
    if not company_name:
        raise HTTPException(status_code=400, detail="name is required")

    company = PlacementCompany(
        name=company_name,
        min_cgpa=_safe_float(payload.get("min_cgpa")),
        max_backlogs=int(payload.get("max_backlogs") or 0),
        role_type=(payload.get("role_type") or None),
        required_skills=payload.get("required_skills") or [],
    )
    db.add(company)
    db.commit()
    db.refresh(company)
    return {"id": company.id, "message": "Company created"}


@app.get("/api/companies")
def get_companies(
    _: dict = Depends(authorize(["admin", "coordinator", "teacher", "student"])),
    db: Session = Depends(get_db),
):
    rows = db.query(PlacementCompany).order_by(PlacementCompany.created_at.desc(), PlacementCompany.id.desc()).all()
    return [
        {
            "id": row.id,
            "name": row.name,
            "min_cgpa": _safe_float(row.min_cgpa),
            "max_backlogs": int(row.max_backlogs or 0),
            "role_type": row.role_type,
            "required_skills": row.required_skills or [],
            "created_at": row.created_at.isoformat() if row.created_at else None,
        }
        for row in rows
    ]


@app.put("/api/companies/{company_id}")
def update_company(
    company_id: int,
    payload: Dict[str, Any] = Body(...),
    _: dict = Depends(authorize(["admin"])),
    db: Session = Depends(get_db),
):
    row = db.query(PlacementCompany).filter(PlacementCompany.id == company_id).first()
    if not row:
        raise HTTPException(status_code=404, detail="Company not found")

    if "name" in payload:
        row.name = str(payload.get("name") or "").strip() or row.name
    if "min_cgpa" in payload:
        row.min_cgpa = _safe_float(payload.get("min_cgpa"))
    if "max_backlogs" in payload:
        row.max_backlogs = int(payload.get("max_backlogs") or 0)
    if "role_type" in payload:
        row.role_type = payload.get("role_type") or None
    if "required_skills" in payload:
        row.required_skills = payload.get("required_skills") or []

    db.commit()
    return {"message": "Company updated"}


@app.delete("/api/companies/{company_id}")
def delete_company(
    company_id: int,
    _: dict = Depends(authorize(["admin"])),
    db: Session = Depends(get_db),
):
    row = db.query(PlacementCompany).filter(PlacementCompany.id == company_id).first()
    if not row:
        raise HTTPException(status_code=404, detail="Company not found")
    db.delete(row)
    db.commit()
    return {"message": "Company deleted"}


@app.post("/api/drives")
def create_drive(
    payload: Dict[str, Any] = Body(...),
    current_user: dict = Depends(authorize(["admin", "coordinator"])),
    db: Session = Depends(get_db),
):
    company_id = int(payload.get("company_id") or 0)
    company = db.query(PlacementCompany).filter(PlacementCompany.id == company_id).first()
    if not company:
        raise HTTPException(status_code=404, detail="Company not found")

    drive_date = payload.get("drive_date")
    parsed_date = datetime.strptime(str(drive_date), "%Y-%m-%d").date() if drive_date else None
    drive = PlacementDrive(
        company_id=company_id,
        drive_date=parsed_date,
        mode=str(payload.get("mode") or "online"),
        branches=payload.get("branches") or [],
        created_by=int(current_user["user_id"]),
    )
    db.add(drive)
    db.commit()
    db.refresh(drive)
    return {"id": drive.id, "message": "Drive created"}


@app.get("/api/drives")
def get_drives(
    _: dict = Depends(authorize(["admin", "coordinator", "teacher", "student"])),
    db: Session = Depends(get_db),
):
    rows = (
        db.query(PlacementDrive, PlacementCompany)
        .join(PlacementCompany, PlacementCompany.id == PlacementDrive.company_id)
        .order_by(PlacementDrive.drive_date.desc(), PlacementDrive.id.desc())
        .all()
    )
    return [
        {
            "id": drive.id,
            "company_id": company.id,
            "company_name": company.name,
            "drive_date": drive.drive_date.isoformat() if drive.drive_date else None,
            "mode": drive.mode,
            "branches": drive.branches or [],
            "created_by": drive.created_by,
        }
        for drive, company in rows
    ]


@app.post("/api/drives/{drive_id}/assign")
def assign_students_to_drive(
    drive_id: int,
    _: dict = Depends(authorize(["coordinator", "admin", "teacher"])),
    db: Session = Depends(get_db),
):
    drive = db.query(PlacementDrive).filter(PlacementDrive.id == drive_id).first()
    if not drive:
        raise HTTPException(status_code=404, detail="Drive not found")

    company = db.query(PlacementCompany).filter(PlacementCompany.id == drive.company_id).first()
    if not company:
        raise HTTPException(status_code=404, detail="Company not found")

    students = db.query(Student).all()
    assigned = 0
    for student in students:
        cgpa = _safe_float(student.cgpa)
        backlogs = _get_student_backlogs(db, student.student_id)
        if cgpa < _safe_float(company.min_cgpa):
            continue
        if backlogs > int(company.max_backlogs or 0):
            continue
        if not _branch_allowed(student.student_id, drive.branches or [], db):
            continue

        existing = db.query(StudentDrive).filter(StudentDrive.student_id == student.student_id, StudentDrive.drive_id == drive_id).first()
        if existing:
            continue

        db.add(StudentDrive(student_id=student.student_id, drive_id=drive_id))
        assigned += 1

    db.commit()
    return {"message": "Students assigned", "assigned": assigned}


@app.put("/api/student-drives/{student_drive_id}")
def update_student_drive(
    student_drive_id: int,
    payload: Dict[str, Any] = Body(...),
    _: dict = Depends(authorize(["coordinator", "admin", "teacher"])),
    db: Session = Depends(get_db),
):
    row = db.query(StudentDrive).filter(StudentDrive.id == student_drive_id).first()
    if not row:
        raise HTTPException(status_code=404, detail="Student drive record not found")

    if "current_round" in payload:
        row.current_round = int(payload.get("current_round") or 0)
    if "status" in payload:
        row.status = str(payload.get("status") or row.status)
    if "final_result" in payload:
        row.final_result = str(payload.get("final_result") or row.final_result)
    row.updated_at = datetime.utcnow()
    db.commit()
    return {"message": "Student drive updated"}


@app.post("/api/drives/{drive_id}/upload-results")
def upload_drive_results(
    drive_id: int,
    file: UploadFile = File(...),
    _: dict = Depends(authorize(["coordinator", "admin"])),
    db: Session = Depends(get_db),
):
    drive = db.query(PlacementDrive).filter(PlacementDrive.id == drive_id).first()
    if not drive:
        raise HTTPException(status_code=404, detail="Drive not found")

    try:
        raw = file.file.read().decode("utf-8")
        reader = csv.DictReader(raw.splitlines())
        updated = 0
        for row in reader:
            student_id = int(row.get("student_id") or 0)
            if student_id <= 0:
                continue
            record = db.query(StudentDrive).filter(StudentDrive.student_id == student_id, StudentDrive.drive_id == drive_id).first()
            if not record:
                continue
            if row.get("current_round") is not None:
                record.current_round = int(row.get("current_round") or 0)
            if row.get("status"):
                record.status = row.get("status")
            if row.get("final_result"):
                record.final_result = row.get("final_result")
            record.updated_at = datetime.utcnow()
            updated += 1
        db.commit()
        return {"message": "Bulk results processed", "updated": updated}
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=400, detail=f"Invalid CSV: {str(e)}")


@app.get("/api/student-drives")
def get_student_drives_api(
    student_id: Optional[int] = Query(default=None),
    _: dict = Depends(authorize(["teacher", "coordinator", "admin"])),
    db: Session = Depends(get_db),
):
    query = (
        db.query(StudentDrive, PlacementDrive, PlacementCompany)
        .join(PlacementDrive, PlacementDrive.id == StudentDrive.drive_id)
        .join(PlacementCompany, PlacementCompany.id == PlacementDrive.company_id)
    )
    if student_id is not None:
        query = query.filter(StudentDrive.student_id == student_id)

    rows = query.order_by(StudentDrive.updated_at.desc(), StudentDrive.id.desc()).all()
    return [
        {
            "id": sd.id,
            "student_id": sd.student_id,
            "drive_id": drive.id,
            "company_name": company.name,
            "drive_date": drive.drive_date.isoformat() if drive.drive_date else None,
            "status": sd.status,
            "current_round": sd.current_round,
            "final_result": sd.final_result,
        }
        for sd, drive, company in rows
    ]


@app.post("/api/feedback")
def add_feedback(
    payload: Dict[str, Any] = Body(...),
    current_user: dict = Depends(authorize(["teacher", "coordinator", "admin"])),
    db: Session = Depends(get_db),
):
    student_id = int(payload.get("student_id") or 0)
    drive_id = int(payload.get("drive_id") or 0)
    if student_id <= 0 or drive_id <= 0:
        raise HTTPException(status_code=400, detail="student_id and drive_id are required")

    feedback = PlacementFeedback(
        student_id=student_id,
        drive_id=drive_id,
        faculty_id=int(current_user["user_id"]),
        comment=payload.get("comment"),
        rating=int(payload.get("rating")) if payload.get("rating") is not None else None,
    )
    db.add(feedback)
    db.commit()
    return {"message": "Feedback added"}


@app.get("/api/student/placement-summary")
def student_placement_summary(
    current_user: dict = Depends(authorize(["student"])),
    db: Session = Depends(get_db),
):
    student_id = int(current_user["user_id"])
    eligible = len([row for row in get_company_eligibility(student_id, db) if row.get("eligible")])

    rows = (
        db.query(StudentDrive, PlacementDrive)
        .join(PlacementDrive, PlacementDrive.id == StudentDrive.drive_id)
        .filter(StudentDrive.student_id == student_id)
        .all()
    )
    today = datetime.utcnow().date()
    upcoming = 0
    completed = 0
    offers = 0

    for sd, drive in rows:
        result = str(sd.final_result or "pending").lower()
        if result == "selected":
            offers += 1
        if drive.drive_date and drive.drive_date >= today and result == "pending":
            upcoming += 1
        else:
            completed += 1

    return {"eligible": eligible, "upcoming": upcoming, "completed": completed, "offers": offers}


@app.get("/api/student/eligible-companies")
def student_eligible_companies(
    current_user: dict = Depends(authorize(["student"])),
    db: Session = Depends(get_db),
):
    student_id = int(current_user["user_id"])
    return [row for row in get_company_eligibility(student_id, db) if row.get("eligible")]


@app.get("/api/student/upcoming-drives")
def student_upcoming_drives(
    current_user: dict = Depends(authorize(["student"])),
    db: Session = Depends(get_db),
):
    student_id = int(current_user["user_id"])
    today = datetime.utcnow().date()
    rows = (
        db.query(StudentDrive, PlacementDrive, PlacementCompany)
        .join(PlacementDrive, PlacementDrive.id == StudentDrive.drive_id)
        .join(PlacementCompany, PlacementCompany.id == PlacementDrive.company_id)
        .filter(StudentDrive.student_id == student_id)
        .order_by(PlacementDrive.drive_date.asc())
        .all()
    )

    result = []
    for sd, drive, company in rows:
        if drive.drive_date and drive.drive_date >= today and str(sd.final_result or "pending").lower() == "pending":
            result.append(
                {
                    "student_drive_id": sd.id,
                    "company_name": company.name,
                    "drive_date": drive.drive_date.isoformat(),
                    "mode": drive.mode,
                    "status": sd.status,
                    "current_round": sd.current_round,
                }
            )
    return result


@app.get("/api/student/past-drives")
def student_past_drives(
    current_user: dict = Depends(authorize(["student"])),
    db: Session = Depends(get_db),
):
    student_id = int(current_user["user_id"])
    today = datetime.utcnow().date()
    rows = (
        db.query(StudentDrive, PlacementDrive, PlacementCompany)
        .join(PlacementDrive, PlacementDrive.id == StudentDrive.drive_id)
        .join(PlacementCompany, PlacementCompany.id == PlacementDrive.company_id)
        .filter(StudentDrive.student_id == student_id)
        .order_by(PlacementDrive.drive_date.desc())
        .all()
    )

    result = []
    for sd, drive, company in rows:
        is_past = (drive.drive_date and drive.drive_date < today) or str(sd.final_result or "pending").lower() in {"selected", "rejected"}
        if is_past:
            result.append(
                {
                    "student_drive_id": sd.id,
                    "company_name": company.name,
                    "drive_date": drive.drive_date.isoformat() if drive.drive_date else None,
                    "mode": drive.mode,
                    "status": sd.status,
                    "current_round": sd.current_round,
                    "final_result": sd.final_result,
                }
            )
    return result


@app.get("/api/student/placement-intelligence")
def student_placement_intelligence(
    current_user: dict = Depends(authorize(["student"])),
    db: Session = Depends(get_db),
):
    student_id = int(current_user["user_id"])
    readiness = get_placement_readiness(student_id, db)
    skill_gap = get_skill_gap(student_id, db)
    prediction = get_selection_probability(student_id, db)
    interviews = get_interview_insights(student_id, db)
    action_plan = generate_action_plan(student_id, db)

    attendance_rows = db.query(Attendance).filter(Attendance.student_id == student_id).all()
    attendance_pct = round(sum(1 for row in attendance_rows if bool(row.status)) * 100 / len(attendance_rows), 2) if attendance_rows else 0.0

    past_rows = student_past_drives(current_user=current_user, db=db)
    selected_count = len([row for row in past_rows if str(row.get("final_result") or "").lower() == "selected"])
    success_rate = round((selected_count / len(past_rows)) * 100, 2) if past_rows else 0.0

    student = db.query(Student).filter(Student.student_id == student_id).first()
    cgpa_val = _safe_float(student.cgpa if student else 0)
    score = (cgpa_val * 10 * 0.4) + (success_rate * 0.3) + (attendance_pct * 0.3)
    readiness_band = "High" if score >= 75 else "Medium" if score >= 50 else "Low"

    recommendations = []
    if interviews.get("common_weak_area"):
        recommendations.append(f"Improve {interviews['common_weak_area']} skills")
    if skill_gap.get("weak_skills"):
        recommendations.append(f"Focus on {skill_gap['weak_skills'][0]} rounds")
    if not recommendations:
        recommendations.append("Keep your current preparation momentum")

    return {
        "readiness": readiness,
        "skill_gap": skill_gap,
        "prediction": prediction,
        "interviews": interviews,
        "action_plan": action_plan,
        "success_probability": {
            "score": round(score, 2),
            "readiness": readiness_band,
            "success_rate": success_rate,
            "attendance": attendance_pct,
        },
        "recommendations": recommendations,
    }



# -------------------------
# FACULTY SIGNUP
# -------------------------
@app.post("/signup/teacher")
def teacher_signup(data: TeacherSignupRequest, db: Session = Depends(get_db)):

    if not data.email.endswith("@gvpcdpgc.edu.in"):
        raise HTTPException(status_code=400, detail="Only college email allowed")

    existing_user = db.query(User).filter(User.email == data.email).first()
    if existing_user:
        raise HTTPException(status_code=400, detail="Email already registered")

    # Check if employee_id already exists
    existing_employee = db.query(Faculty).filter(Faculty.employee_id == data.employee_id).first()
    if existing_employee:
        raise HTTPException(status_code=400, detail="Employee ID already registered")

    try:
        # 1️⃣ Create user
        new_user = User(
            name=data.name,
            email=data.email,
            password=hash_password(data.password),
            role="faculty",          # ✅ STANDARDIZE ROLE
            department_id=data.department_id
        )
        db.add(new_user)
        db.flush()  # 🔑 get user_id WITHOUT commit

        # 2️⃣ Create faculty
        faculty = Faculty(
            faculty_id=new_user.user_id,
            employee_id=data.employee_id
        )
        db.add(faculty)

        # 3️⃣ Commit ONCE
        db.commit()

        return {
            "message": "Teacher signup successful",
            "user_id": new_user.user_id
        }

    except Exception as e:
        db.rollback()  # 🔥 THIS SAVES YOU
        raise HTTPException(status_code=400, detail=str(e))

# -------------------------
# FACULTY PROFILE GET
# -------------------------
@app.get("/faculty/profile")
def get_faculty_profile(
    user=Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if user["role"] != "faculty":
        raise HTTPException(status_code=403, detail="Not authorized")

    faculty = (
        db.query(Faculty, User)
        .join(User, Faculty.faculty_id == User.user_id)
        .filter(User.user_id == user["user_id"])
        .first()
    )

    if not faculty:
        raise HTTPException(status_code=404, detail="Faculty not found")

    faculty_data, user_data = faculty

    return {
    # ---------- USER ----------
    "name": user_data.name,
    "email": user_data.email,

    # ---------- FACULTY ----------
    "employee_id": faculty_data.employee_id,
    "designation": faculty_data.designation,
    #"department": user_data.department_id,#
    "qualifications": faculty_data.qualifications,
    "experience": faculty_data.experience,

    "phone": faculty_data.phone,
    "bio": faculty_data.bio,

    "linkedin": faculty_data.linkedin,
    "github": faculty_data.github,
    "portfolio": faculty_data.portfolio,

    # ---------- LIST / JSON ----------
    "expertise": faculty_data.expertise.split(",")
        if faculty_data.expertise else [],

    "certifications": json.loads(faculty_data.certifications)
        if faculty_data.certifications else [],

    "publications": json.loads(faculty_data.publications)
        if faculty_data.publications else [],

    "classes": json.loads(faculty_data.classes)
        if faculty_data.classes else []
 }


# -------------------------
# FACULTY PROFILE PUT
# -------------------------
@app.put("/faculty/profile")
def update_faculty_profile(
    data: FacultyProfileUpdate,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "faculty":
        raise HTTPException(status_code=403, detail="Not authorized")

    faculty = db.query(Faculty).filter(
        Faculty.faculty_id == current_user["user_id"]
    ).first()

    user = db.query(User).filter(
        User.user_id == current_user["user_id"]
    ).first()

    if not faculty or not user:
        raise HTTPException(status_code=404, detail="Faculty not found")

    # ----- USERS TABLE -----
    if data.name is not None:
        user.name = data.name

    # ----- FACULTY TABLE -----
    if data.phone is not None:
        faculty.phone = data.phone

    if data.bio is not None:
        faculty.bio = data.bio

    if data.linkedin is not None:
        faculty.linkedin = data.linkedin

    if data.github is not None:
        faculty.github = data.github

    if data.portfolio is not None:
        faculty.portfolio = data.portfolio

    if data.qualifications is not None:
        faculty.qualifications = data.qualifications

    if data.experience is not None:
        faculty.experience = data.experience

    # ----- JSON FIELDS -----
    if data.expertise is not None:
        faculty.expertise = ",".join(data.expertise)

    if data.certifications:
     faculty.certifications = json.dumps(
        [c.dict() for c in data.certifications]
    )

    if data.publications is not None:
        faculty.publications = json.dumps(
            [p.dict() for p in data.publications]
        )

    if data.classes is not None:
        faculty.classes = json.dumps(
            [c.dict() for c in data.classes]
        )

    db.commit()
    return {"message": "Profile updated successfully"}


# -------------------------
# FACULTY – MARK ATTENDANCE
# -------------------------    
@app.post("/faculty/attendance")
def mark_attendance(
    payload: AttendanceCreate,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "faculty":
        raise HTTPException(status_code=403, detail="Faculty only")

    # 🔒 Check faculty assignment
    subject_check = db.query(FacultySubject).filter(
        FacultySubject.faculty_id == current_user["user_id"],
        FacultySubject.subject_id == payload.subject_id,
        FacultySubject.year == payload.year,
        FacultySubject.section == payload.section,
        FacultySubject.is_active == True
    ).first()

    if not subject_check:
        raise HTTPException(status_code=403, detail="Not assigned to this class")

    updated_students = []

    # -----------------------------------
    # UPDATE / CREATE ATTENDANCE RECORDS
    # -----------------------------------
    for record in payload.records:

        existing = db.query(Attendance).filter(
            Attendance.student_id == record.student_id,
            Attendance.subject_id == payload.subject_id,
            Attendance.attendance_date == payload.date
        ).first()

        if existing:
            if existing.status != record.status:
                existing.status = record.status
                updated_students.append(record.student_id)

        else:
            new_attendance = Attendance(
                student_id=record.student_id,
                subject_id=payload.subject_id,
                faculty_id=current_user["user_id"],
                attendance_date=payload.date,
                status=record.status
            )
            db.add(new_attendance)
            updated_students.append(record.student_id)

    # ✅ Commit once after processing all students
    db.commit()

    # -----------------------------------
    # FETCH SUBJECT & FACULTY INFO
    # -----------------------------------
    subject = db.query(Subject).filter(
        Subject.subject_id == payload.subject_id
    ).first()

    

    faculty_user = db.query(User).filter(
        User.user_id == current_user["user_id"]
    ).first()

    # -----------------------------------
    # CREATE ALERTS (ONLY IF UPDATED)
    # -----------------------------------
    for student_id in updated_students:

        # Delete old attendance alerts for this student
        old_alerts = db.query(Alert).filter(
            Alert.type == "attendance",
            Alert.student_id == student_id
        ).all()

        for old in old_alerts:
            db.query(AlertRecipient).filter(
                AlertRecipient.alert_id == old.id
            ).delete()
            db.delete(old)

        db.commit()

        # Create new alert
        new_alert = Alert(
            title="Attendance Updated",
            message=f"{faculty_user.name} updated your attendance for {subject.subject_name} on {payload.date}.",
            type="attendance",
            target_role="student",
            target_type="individual",
            student_id=student_id
        )

        db.add(new_alert)
        db.commit()
        db.refresh(new_alert)

        recipient = AlertRecipient(
            alert_id=new_alert.id,
            user_id=student_id,
            is_read=False
        )

        db.add(recipient)
        db.commit()

    return {"message": "Attendance saved successfully"}




# -------------------------
# FACULTY – GET STUDENTS FOR ATTENDANCE
# -------------------------
@app.get("/faculty/attendance/students")
def get_students_for_attendance(
    year: int,
    section: str,
    subject_id: int,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "faculty":
        raise HTTPException(status_code=403, detail="Faculty only")

    # 🔒 Check faculty assignment
    assignment = db.query(FacultySubject).filter(
        FacultySubject.faculty_id == current_user["user_id"],
        FacultySubject.subject_id == subject_id,
        FacultySubject.year == year,
        FacultySubject.section == section,
        FacultySubject.is_active == True
    ).first()

    print("USER:", current_user["user_id"])
    print("SUBJECT:", subject_id)
    print("YEAR:", year)
    print("SECTION:", section)
    print("ASSIGNMENT:", assignment)

    if not assignment:
        raise HTTPException(status_code=403, detail="Not assigned to this class")

    # Get subject department properly
    subject = db.query(Subject).filter(
        Subject.subject_id == subject_id
    ).first()

    if not subject:
        raise HTTPException(status_code=404, detail="Subject not found")

    department_id = subject.department_id
    print("Subject Department:", department_id)

    test_students = db.query(Student).filter(
        Student.year == year,
        Student.section == section
    ).all()

    print("Students matching year & section:", len(test_students))

    for s in test_students:
        user_obj = db.query(User).filter(
            User.user_id == s.student_id
        ).first()
        print("Student Roll:", s.roll_no, "Dept:", user_obj.department_id)

    students = (
    db.query(Student, User)
    .join(User, Student.student_id == User.user_id)
    .filter(
        User.department_id == department_id,
        Student.year == year,
        Student.section == section,
        User.is_deleted == False
    )
    .order_by(Student.roll_no.asc())   # ✅ SORT BY ROLL
    .all()
 )

    result = []

    for student, user in students:

        # 🔥 Generate last 5 calendar dates (including missing ones)
        today = date.today()

        last_5_status = []

        for i in range(4, -1, -1):  # oldest → newest
            check_date = today - timedelta(days=i)

            record = db.query(Attendance).filter(
                Attendance.student_id == student.student_id,  # type: ignore
                Attendance.subject_id == subject_id,
                Attendance.attendance_date == check_date
            ).first()

            if record:
                last_5_status.append({
                    "status": record.status,
                    "date": check_date
                })
            else:
                last_5_status.append({
                    "status": None,   # 👈 important for grey dot
                    "date": check_date
                })
        
        total = db.query(Attendance).filter(
            Attendance.student_id == student.student_id,  # type: ignore
            Attendance.subject_id == subject_id
        ).count()

        if total == 0:
            continue  # no classes yet

        present = db.query(Attendance).filter(
            Attendance.student_id == student.student_id,  # type: ignore
            Attendance.subject_id == subject_id,
            Attendance.status == True
        ).count()

        percentage = round((present / total) * 100, 2) if total > 0 else 0

        result.append({
            "id": student.student_id,  # type: ignore
            "roll": student.roll_no,  # type: ignore
            "name": user.name,  # type: ignore
            "last_5": last_5_status,
            "percentage": percentage,
            "present": present,
            "total": total
        })

    return result

# -------------------------
# FACULTY – GET LAST 5 ATTENDANCE RECORDS FOR A STUDENT
# -------------------------
@app.get("/faculty/attendance/last5")
def get_last_5_classes(
    subject_id: int,
    student_id: int,
    db: Session = Depends(get_db)
):
    records = db.query(Attendance).filter(
        Attendance.subject_id == subject_id,
        Attendance.student_id == student_id
    ).order_by(Attendance.attendance_date.desc()).limit(5).all()

    return records


# -------------------------
# FACULTY – CHECK IF ATTENDANCE ALREADY EXISTS
# -------------------------
@app.get("/faculty/attendance/check")
def check_attendance_exists(
    subject_id: int,
    date: date,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "faculty":
        raise HTTPException(status_code=403, detail="Faculty only")

    existing = db.query(Attendance).filter(
        Attendance.subject_id == subject_id,
        Attendance.attendance_date == date,
        Attendance.faculty_id == current_user["user_id"]
    ).first()

    return {
        "already_marked": True if existing else False
    }


# -------------------------
# FACULTY – GET ATTENDANCE BY DATE (FOR EDITING)
# -------------------------
@app.get("/faculty/attendance/by-date")
def get_attendance_by_date(
    subject_id: int,
    date: date,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "faculty":
        raise HTTPException(status_code=403, detail="Faculty only")

    records = db.query(Attendance).filter(
        Attendance.subject_id == subject_id,
        Attendance.attendance_date == date,
        Attendance.faculty_id == current_user["user_id"]
    ).all()

    return [
        {
            "student_id": r.student_id,
            "status": r.status
        }
        for r in records
    ]

# -------------------------
# FACULTY – GET SEMESTER ATTENDANCE PERCENTAGE FOR ALL STUDENTS
# -------------------------
@app.get("/faculty/attendance/semester/{subject_id}")
def get_semester_attendance(
    subject_id: int,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "faculty":
        raise HTTPException(status_code=403, detail="Faculty only")

    students = (
        db.query(Student)
        .join(FacultySubject, FacultySubject.year == Student.year)
        .filter(
            FacultySubject.faculty_id == current_user["user_id"],
            FacultySubject.subject_id == subject_id,
            FacultySubject.is_active == True
        )
        .all()
    )

    result = []

    for student in students:
        total = db.query(Attendance).filter(
            Attendance.student_id == student.student_id,
            Attendance.subject_id == subject_id
        ).count()

        present = db.query(Attendance).filter(
            Attendance.student_id == student.student_id,
            Attendance.subject_id == subject_id,
            Attendance.status == True
        ).count()

        percentage = round((present / total) * 100, 2) if total > 0 else 0

        result.append({
            "student_id": student.student_id,
            "percentage": percentage
        })

    return result


# =========================
# FACULTY – UNIVERSAL ATTENDANCE REPORT
# =========================
@app.get("/faculty/attendance/report/{subject_id}")
def attendance_report(
    subject_id: int,
    start_date: Optional[date] = None,
    end_date: Optional[date] = None,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):

    if current_user["role"] != "faculty":
        raise HTTPException(status_code=403, detail="Faculty only")

    # 🔒 Validate assignment
    assignment = db.query(FacultySubject).filter(
        FacultySubject.faculty_id == current_user["user_id"],
        FacultySubject.subject_id == subject_id,
        FacultySubject.is_active == True
    ).first()

    if not assignment:
        raise HTTPException(status_code=403, detail="Not assigned to this subject")

    # ✅ Default to current week if no dates given
    if not start_date or not end_date:
        today = date.today()
        start_date = today - timedelta(days=today.weekday())
        end_date = today

    # Get subject department properly
    subject = db.query(Subject).filter(
        Subject.subject_id == subject_id
    ).first()

    department_id = subject.department_id

    students = (
        db.query(Student, User)
        .join(User, Student.student_id == User.user_id)
        .filter(
            Student.year == assignment.year,
            Student.section == assignment.section,
            User.department_id == department_id,
            User.is_deleted == False
        )
        .all()
    )

    

    student_data = []
    total_records = 0
    total_present = 0
    total_absent = 0

    unique_classes = db.query(Attendance.attendance_date).filter(
        Attendance.subject_id == subject_id,
        Attendance.attendance_date >= start_date,
        Attendance.attendance_date <= end_date
    ).distinct().count()

    for student, user in students:

        records = db.query(Attendance).filter(
            Attendance.student_id == student.student_id,
            Attendance.subject_id == subject_id,
            Attendance.attendance_date >= start_date,
            Attendance.attendance_date <= end_date
        ).all()

        total = unique_classes
        present = len([r for r in records if r.status])
        absent = total - present
        percent = round((present / total) * 100, 2) if total > 0 else 0  # type: ignore

        total_records += total
        total_present += present
        total_absent += absent

        student_data.append({
            "roll": student.roll_no,
            "name": user.name,
            "total_classes": total,
            "present": present,
            "absent": absent,
            "percentage": percent
        })

    student_data.sort(key=lambda x: x["percentage"], reverse=True)

    
    total_entries = total_present + total_absent

    class_average = round(  # type: ignore
        (total_present / total_entries) * 100, 2  # type: ignore
    ) if total_entries > 0 else 0

    present_percentage = round(  # type: ignore
        (total_present / total_entries) * 100, 2  # type: ignore
    ) if total_entries > 0 else 0

    absent_percentage = round(  # type: ignore
        (total_absent / total_entries) * 100, 2  # type: ignore
    ) if total_entries > 0 else 0

    class_average = present_percentage
    

    return {
    "start_date": start_date,
    "end_date": end_date,
    "total_records": unique_classes,
    "present_percentage": present_percentage,
    "absent_percentage": absent_percentage
 }


# =========================
# FACULTY – DOWNLOAD REPORT PDF (PRO VERSION)
# =========================
@app.get("/faculty/attendance/report/{subject_id}/download")
def download_report_pdf(
    subject_id: int,
    start_date: Optional[date] = None,
    end_date: Optional[date] = None,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):

    if current_user["role"] != "faculty":
        raise HTTPException(status_code=403, detail="Faculty only")

    from datetime import timedelta

    if not start_date or not end_date:
        today = date.today()
        start_date = today - timedelta(days=today.weekday())
        end_date = today

    subject = db.query(Subject).filter(
        Subject.subject_id == subject_id
    ).first()

    faculty = db.query(User).filter(
        User.user_id == current_user["user_id"]
    ).first()

    assignment = db.query(FacultySubject).filter(
        FacultySubject.faculty_id == current_user["user_id"],
        FacultySubject.subject_id == subject_id,
        FacultySubject.is_active == True
    ).first()

    if not assignment:
        raise HTTPException(status_code=403, detail="Not assigned")

    unique_classes = db.query(Attendance.attendance_date).filter(
        Attendance.subject_id == subject_id,
        Attendance.attendance_date >= start_date,
        Attendance.attendance_date <= end_date
    ).distinct().count()

    students = (
        db.query(Student, User)
        .join(User, Student.student_id == User.user_id)
        .filter(
            Student.year == assignment.year,
            Student.section == assignment.section,
            User.department_id == subject.department_id,
            User.is_deleted == False
        )
        .all()
    )

    student_rows = []
    total_present = 0
    total_absent = 0

    for student, user in students:

        records = db.query(Attendance).filter(
            Attendance.student_id == student.student_id,
            Attendance.subject_id == subject_id,
            Attendance.attendance_date >= start_date,
            Attendance.attendance_date <= end_date
        ).all()

        total = len(records)
        present = len([r for r in records if r.status])
        absent = total - present
        percent = round((present / total) * 100, 2) if total > 0 else 0  # type: ignore

        total_present += present
        total_absent += absent

        student_rows.append({
            "roll": student.roll_no,
            "name": user.name,
            "total_classes": total,
            "present": present,
            "absent": absent,
            "percentage": percent
        })

    student_rows.sort(key=lambda x: x["percentage"], reverse=True)

    rank = 1
    for i, s in enumerate(student_rows):
        if i > 0 and s["percentage"] < student_rows[i-1]["percentage"]:
            rank = i + 1
        s["rank"] = rank

    highest = student_rows[0] if student_rows else None
    lowest = student_rows[-1] if student_rows else None

    total_entries = total_present + total_absent  # type: ignore
    class_average = round(  # type: ignore
        (total_present / total_entries) * 100, 2  # type: ignore
    ) if total_entries > 0 else 0

    report_format = get_report_format("attendance")
    if report_format == "excel":
        output = BytesIO()
        df = pd.DataFrame(student_rows)
        df.to_excel(output, index=False, engine="openpyxl")
        output.seek(0)
        return StreamingResponse(
            output,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f"attachment; filename=attendance_report_{subject_id}.xlsx"}
        )

    if report_format == "docx":
        try:
            from docx import Document
        except ImportError:
            raise HTTPException(status_code=500, detail="DOCX export support is unavailable")

        document = Document()
        document.add_heading("Attendance Performance Report", level=1)
        document.add_paragraph(f"Subject: {subject.subject_name}")
        document.add_paragraph(f"Date range: {start_date} to {end_date}")
        document.add_paragraph(f"Class Average: {class_average}%")

        table = document.add_table(rows=1, cols=6)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Roll"
        hdr_cells[1].text = "Name"
        hdr_cells[2].text = "Total"
        hdr_cells[3].text = "Present"
        hdr_cells[4].text = "Absent"
        hdr_cells[5].text = "Percentage"

        for row in student_rows:
            cells = table.add_row().cells
            cells[0].text = str(row["roll"])
            cells[1].text = str(row["name"])
            cells[2].text = str(row["total_classes"])
            cells[3].text = str(row["present"])
            cells[4].text = str(row["absent"])
            cells[5].text = f"{row['percentage']}%"

        output = BytesIO()
        document.save(output)
        output.seek(0)
        return StreamingResponse(
            output,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=attendance_report_{subject_id}.docx"}
        )

    # ================= PDF BUILD =================

    file_path = f"attendance_report_{subject_id}.pdf"
    doc = SimpleDocTemplate(
        file_path,
        pagesize=A4,
        rightMargin=40,
        leftMargin=40,
        topMargin=60,
        bottomMargin=40
    )

    elements = []
    styles = getSampleStyleSheet()

    # ===== HEADER =====

    logo = Image("assests/gvp logo.jpg", width=0.9*inch, height=0.9*inch)

    college_style = ParagraphStyle(
        'CollegeStyle',
        parent=styles['Normal'],
        fontSize=13,
        leading=16
    )

    header = Table([[
        logo,
        Paragraph(
            "<b>GAYATRI VIDYA PARISHAD COLLEGE FOR DEGREE AND PG COURSES (A)</b>",
            college_style
        )
    ]], colWidths=[1*inch, 4.8*inch])

    header.setStyle(TableStyle([
        ("VALIGN", (0,0), (-1,-1), "MIDDLE")
    ]))

    elements.append(header)
    elements.append(Spacer(1, 15))

    title_style = ParagraphStyle(
        'CenteredTitle',
        parent=styles['Heading2'],
        alignment=1,  # CENTER
        textColor=colors.HexColor("#1F3A8A"),
        spaceAfter=10
    )

    elements.append(Paragraph(
        "<b>ATTENDANCE PERFORMANCE REPORT</b>",
        title_style
    ))

    elements.append(Spacer(1, 20))

    # ===== SUMMARY BOX (ALL IMPORTANT DATA AT TOP) =====

    summary_data = [
        ["Faculty:", faculty.name],
        ["Subject:", subject.subject_name],
        ["Period:", f"{start_date} to {end_date}"],
        ["Total Classes:", unique_classes],
        ["Top Attendance:", f"{highest['name']} ({highest['percentage']}%)" if highest else "-"],
        ["Low Attendance:", f"{lowest['name']} ({lowest['percentage']}%)" if lowest else "-"],
        ["Class Average:", f"{class_average}%"]
    ]

    summary_table = Table(summary_data, colWidths=[2*inch, 3.8*inch])

    summary_table.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,-1), colors.whitesmoke),
        ("BOX", (0,0), (-1,-1), 1, colors.grey),
        ("INNERGRID", (0,0), (-1,-1), 0.25, colors.grey),
        ("FONTNAME", (0,0), (0,-1), "Helvetica-Bold"),  # Bold left column
        ("VALIGN", (0,0), (-1,-1), "MIDDLE")
    ]))
    elements.append(summary_table)
    elements.append(Spacer(1, 25))

    # ===== STUDENT TABLE =====

    # Sort by percentage DESC, then roll ASC
    student_rows.sort(
        key=lambda x: (-x["percentage"], x["roll"])
    )

    # Assign rank properly
    for index, s in enumerate(student_rows, start=1):
        s["rank"] = index

    table_data = [["Rank", "Roll No", "Name", "Total", "Present", "Absent", "%"]]

    for s in student_rows:
        table_data.append([  # type: ignore
            s["rank"],
            s["roll"],
            s["name"],
            s["total_classes"],
            s["present"],
            s["absent"],
            f"{s['percentage']}%"
        ])

    table = Table(
        table_data,
        repeatRows=1,
        colWidths=[0.6*inch, 1.2*inch, 1.8*inch, 0.7*inch, 0.8*inch, 0.8*inch, 0.7*inch]
    )

    table.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#1F3A8A")),
        ("TEXTCOLOR", (0,0), (-1,0), colors.white),
        ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
        ("ALIGN", (0,0), (-1,-1), "CENTER"),
        ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
        ("GRID", (0,0), (-1,-1), 0.25, colors.grey)
    ]))

    elements.append(table)

    # ===== WATERMARK + BORDER + FOOTER =====

    def add_layout(canvas_obj, doc_obj):

        # Watermark (lighter)
        canvas_obj.saveState()
        canvas_obj.setFont("Helvetica-Bold", 80)
        canvas_obj.setFillColorRGB(0.96, 0.96, 0.96)
        canvas_obj.translate(A4[0]/2, A4[1]/2)
        canvas_obj.rotate(45)
        canvas_obj.drawCentredString(0, 0, "GVP-MAAA")
        canvas_obj.restoreState()

        # Footer (move higher)
        canvas_obj.setFont("Helvetica", 9)
        canvas_obj.drawString(40, 35,
            f"Generated on: {datetime.now().strftime('%d-%m-%Y %H:%M')}"
        )
        canvas_obj.drawRightString(A4[0]-40, 35,
            f"Page {doc_obj.page}"
        )

        # Border
        canvas_obj.setLineWidth(1)
        canvas_obj.rect(25, 25, A4[0]-50, A4[1]-50)

    doc.build(elements, onFirstPage=add_layout, onLaterPages=add_layout)

    return FileResponse(
        file_path,
        media_type="application/pdf",
        filename=file_path
    )

# =========================
# FACULTY – CLASS ATTENDANCE SUMMARY
# =========================
@app.get("/faculty/attendance/class-summary")
def get_class_attendance_summary(
    subject_id: int,
    department: str,
    year: int,
    section: str,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "faculty":
        raise HTTPException(status_code=403, detail="Faculty only")

    # convert department to id
    department_id = None
    for key, value in DEPARTMENT_MAP.items():
        if value == department:
            department_id = key

    subject = db.query(Subject).filter(
        Subject.subject_id == subject_id
    ).first()

    department_id = subject.department_id

    # 🔒 Validate assignment first
    assignment = db.query(FacultySubject).filter(
        FacultySubject.faculty_id == current_user["user_id"],
        FacultySubject.subject_id == subject_id,
        FacultySubject.year == year,
        FacultySubject.section == section,
        FacultySubject.is_active == True
    ).first()

    if not assignment:
        raise HTTPException(status_code=403, detail="Not assigned")

    subject = db.query(Subject).filter(
        Subject.subject_id == subject_id
    ).first()

    department_id = subject.department_id

    students = (
        db.query(Student, User)
        .join(User, Student.student_id == User.user_id)
        .filter(
            Student.year == year,
            Student.section == section,
            User.department_id == department_id,
            User.is_deleted == False
        )
        .all()
    )

    results = []

    for student, user in students:

        total = db.query(Attendance).filter(
            Attendance.student_id == student.student_id,
            Attendance.subject_id == subject_id
        ).count()

        present = db.query(Attendance).filter(
            Attendance.student_id == student.student_id,
            Attendance.subject_id == subject_id,
            Attendance.status == True
        ).count()

        percentage = (present / total * 100) if total > 0 else 0

        results.append({
            "student_id": student.student_id,
            "name": user.name,
            "percentage": round(percentage, 2)
        })

    return results


# =========================
# STUDENT – GET ATTENDANCE (SEMESTER BASED)
# =========================
@app.get("/student/attendance")
def get_student_attendance(
    semester: int,
    subject_id: Optional[int] = None,
    student_id: Optional[int] = None,
    view: str = "daily",
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):

    if current_user["role"] != "student":
        raise HTTPException(status_code=403, detail="Student only")

    selected_student_id = student_id or current_user["user_id"]
    if student_id and student_id != current_user["user_id"]:
        raise HTTPException(status_code=403, detail="Cannot access another student's attendance")

    # Get student record
    student = db.query(Student).filter(
        Student.student_id == selected_student_id
    ).first()

    if not student:
        raise HTTPException(status_code=404, detail="Student not found")

    # Get subjects for this semester & department
    subjects = db.query(Subject).filter(
        Subject.semester == semester,
        Subject.department_id == current_user["department_id"]
    ).all()

    result = []

    for subject in subjects:

        # Filter by subject if selected
        if subject_id and subject.subject_id != subject_id:
            continue

        total = db.query(Attendance).filter(
            Attendance.student_id == student.student_id,
            Attendance.subject_id == subject.subject_id
        ).count()

        present = db.query(Attendance).filter(
            Attendance.student_id == student.student_id,
            Attendance.subject_id == subject.subject_id,
            Attendance.status == True
        ).count()

        percentage = round((present / total) * 100, 2) if total > 0 else 0

        # Last 5 classes
        last_5_records = db.query(Attendance).filter(
            Attendance.student_id == student.student_id,
            Attendance.subject_id == subject.subject_id
        ).order_by(Attendance.attendance_date.desc()).limit(5).all()

        last_5 = [
            {
                "date": r.attendance_date,
                "status": r.status
            }
            for r in last_5_records
        ]

        attendance_records = db.query(Attendance).filter(
            Attendance.student_id == student.student_id,
            Attendance.subject_id == subject.subject_id
        ).order_by(Attendance.attendance_date.asc()).all()

        trend = build_attendance_trend(attendance_records, view)

        result.append({
            "subject_id": subject.subject_id,
            "subject_name": subject.subject_name,
            "conducted": total,
            "attended": present,
            "percentage": percentage,
            "last_5": last_5,
            "trend": trend,
        })

    return result



# =========================
# STUDENT – GET ATTENDANCE ( MONTHLY BASED)
# =========================
@app.get("/student/attendance/monthly")
def get_monthly_attendance(
    semester: int,
    month: int,
    year: int,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):

    if current_user["role"] != "student":
        raise HTTPException(status_code=403, detail="Student only")

    student = db.query(Student).filter(
        Student.student_id == current_user["user_id"]
    ).first()

    if not student:
        raise HTTPException(status_code=404, detail="Student not found")

    # Get subjects for this semester
    subjects = db.query(Subject).filter(
        Subject.semester == semester,
        Subject.department_id == current_user["department_id"]
    ).all()

    from calendar import monthrange
    from datetime import date

    total_days = monthrange(year, month)[1]

    response = []

    for day in range(1, total_days + 1):

        current_date = date(year, month, day)

        day_data = {
            "date": current_date,
            "subjects": []
        }

        for subject in subjects:

            # Check if class conducted for this subject on this date
            class_exists = db.query(Attendance).filter(
                Attendance.subject_id == subject.subject_id,
                Attendance.attendance_date == current_date
            ).first()

            if class_exists:

                student_record = db.query(Attendance).filter(
                    Attendance.student_id == student.student_id,
                    Attendance.subject_id == subject.subject_id,
                    Attendance.attendance_date == current_date
                ).first()

                status = student_record.status if student_record else False

                day_data["subjects"].append({  # type: ignore
                    "subject": subject.subject_name,
                    "working_day": True,
                    "status": status
                })

            else:
                day_data["subjects"].append({  # type: ignore
                    "subject": subject.subject_name,
                    "working_day": False,
                    "status": None
                })

        response.append(day_data)

    return response





# =========================
# STUDENT – ATTENDANCE ANALYTICS
# =========================
@app.get(
    "/student/attendance/analytics",
    response_model=AttendanceAnalyticsResponse
)
def get_attendance_analytics(
    semester: int,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):

    if current_user["role"] != "student":
        raise HTTPException(status_code=403, detail="Student only")

    student = db.query(Student).filter(
        Student.student_id == current_user["user_id"]
    ).first()

    if not student:
        raise HTTPException(status_code=404, detail="Student not found")

    # 🔹 Get subjects for semester
    subjects = db.query(Subject).filter(
        Subject.semester == semester,
        Subject.department_id == current_user["department_id"]
    ).all()

    trend_data = []
    subject_comparison = []

    all_dates = set()

    # Collect all attendance dates
    for subject in subjects:
        records = db.query(Attendance).filter(
            Attendance.student_id == student.student_id,
            Attendance.subject_id == subject.subject_id
        ).all()

        for r in records:
            all_dates.add(r.attendance_date)

    sorted_dates = sorted(list(all_dates))

    total_present = 0
    total_count = 0

    # 🔹 Build cumulative trend
    for d in sorted_dates:

        daily_records = db.query(Attendance).filter(
            Attendance.student_id == student.student_id,
            Attendance.attendance_date <= d
        ).all()

        present = len([r for r in daily_records if r.status])
        total = len(daily_records)

        percentage = round((present / total) * 100, 2) if total > 0 else 0  # type: ignore

        trend_data.append({
            "date": d,
            "percentage": percentage
        })

    # 🔹 Subject comparison
    for subject in subjects:

        total = db.query(Attendance).filter(
            Attendance.student_id == student.student_id,
            Attendance.subject_id == subject.subject_id
        ).count()

        present = db.query(Attendance).filter(
            Attendance.student_id == student.student_id,
            Attendance.subject_id == subject.subject_id,
            Attendance.status == True
        ).count()

        percentage = round((present / total) * 100, 2) if total > 0 else 0

        subject_comparison.append({
            "subject": subject.subject_name,
            "percentage": percentage
        })

        total_present += present
        total_count += total

    # 🔹 Simple Projection (Linear)
    current_percentage = round(  # type: ignore
        (total_present / total_count) * 100, 2  # type: ignore
    ) if total_count > 0 else 0

    projected_percentage = min(
        round(current_percentage + 2.5, 2),  # type: ignore
        100
    )

    confidence = "high" if total_count > 20 else "moderate"

    prediction = {
        "projected_percentage": projected_percentage,
        "confidence": confidence
    }

    return {
        "trend": trend_data,
        "subject_comparison": subject_comparison,
        "prediction": prediction
    }




# -------------------------
# ADMIN LOGIN (JWT BASED)
# -------------------------
@app.post("/login/admin")
def admin_login(data: AdminLoginRequest, db: Session = Depends(get_db)):
    print("\n=== ADMIN LOGIN ATTEMPT ===")
    print("Incoming email:", data.email)
    print("Incoming password length:", len(data.password))

    # 1️⃣ Validate admin access key
    if data.access_key != os.getenv("ADMIN_ACCESS_KEY"):
        raise HTTPException(status_code=403, detail="Invalid admin access key")

    # 2️⃣ Get admin user
    admin = db.query(User).filter(
        User.email == data.email,
        User.role == "admin"
    ).first()

    if not admin:
        raise HTTPException(status_code=404, detail="Admin not found")

    # 3️⃣ Verify password with backward compatibility
    password_result = verify_password(data.password, admin.password)
    print("Password verification result:", password_result)
    
    if password_result == "upgrade":
        # 🔄 AUTO-UPGRADE: Old password format detected
        print("[!] Upgrading password hash for admin:", admin.email)
        try:
            admin.password = hash_password(data.password)
            db.commit()
            print("[✓] Admin password upgraded successfully")
        except Exception as e:
            print(f"[!] Password upgrade error: {e}")
            db.rollback()
    elif password_result != True:
        raise HTTPException(status_code=401, detail="Invalid password")

    # 4️⃣ Create JWT
    access_token = create_access_token(
        data={
            "user_id": admin.user_id,
            "role": admin.role,
            "department_id": admin.department_id
        }
    )

    return {
        "access_token": access_token,
        "token_type": "bearer",
        "user_id": admin.user_id,
        "name": admin.name,
        "role": admin.role
    }

# =========================
# ADMIN – PROMOTE STUDENTS
# =========================
@app.put("/admin/students/promote")
def promote_students(
    current_year: int,
    new_year: int,
    section: str | None = None,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    # 🔐 Admin only
    if current_user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Admin only")

    query = db.query(Student).filter(Student.year == current_year)

    if section:
        query = query.filter(Student.section == section)

    students = query.all()

    if not students:
        raise HTTPException(status_code=404, detail="No students found")

    for s in students:
        s.year = new_year
        s.semester = new_year * 2 - 1  # semester logic

    db.commit()

    return {
        "message": f"{len(students)} students promoted successfully"
    }

# =========================
# ADMIN – GET ALL STUDENTS
# =========================
@app.get("/admin/students")
def get_all_students(
    year: Optional[int] = None,
    semester: Optional[int] = None,
    section: Optional[str] = None,
    department: Optional[str] = None,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Admin only")

    query = (
        db.query(Student, User)
        .join(User, Student.student_id == User.user_id)
        .filter(Student.is_deleted == False)
    )

    # -------------------------
    # APPLY FILTERS
    # -------------------------
    if year:
        query = query.filter(Student.year == year)

    if semester:
        query = query.filter(Student.semester == semester)

    if section:
        query = query.filter(Student.section == section)

    if department:
        dept_id = None
        for key, value in DEPARTMENT_MAP.items():
            if value == department:
                dept_id = key

        if dept_id:
            query = query.filter(User.department_id == dept_id)

    students = query.all()

    result = []

    for student, user in students:

        total = db.query(Attendance).filter(
            Attendance.student_id == student.student_id
        ).count()

        present = db.query(Attendance).filter(
            Attendance.student_id == student.student_id,
            Attendance.status == True
        ).count()

        percentage = round((present / total) * 100, 2) if total > 0 else 0

        # -------------------------
        # RISK CLASSIFICATION
        # -------------------------
        attendance_threshold = get_setting("attendance_threshold") or 75
        cgpa_threshold = get_setting("cgpa_threshold") or 6.5
        print("Using attendance threshold:", attendance_threshold, "CGPA threshold:", cgpa_threshold)

        student_cgpa = float(student.cgpa) if student.cgpa is not None else 0.0

        if percentage < 60:
            risk = "Critical"
        elif percentage < attendance_threshold or student_cgpa < cgpa_threshold:
            risk = "Warning"
        else:
            risk = "Safe"

        result.append({
            "id": student.student_id,
            "roll": student.roll_no,
            "name": user.name,
            "year": student.year,
            "semester": student.semester,
            "section": student.section,
            "department": DEPARTMENT_MAP.get(user.department_id, "UNKNOWN"),
            "attendance": percentage,
            "cgpa": student_cgpa,
            "risk": risk
        })

    return result

# =========================
# ADMIN – BULK PROMOTE STUDENTS
# =========================
@app.put("/admin/students/bulk-promote")
def bulk_promote_students(
    payload: StudentPromotionRequest,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Admin only")

    students = db.query(Student).filter(
        Student.student_id.in_(payload.student_ids)
    ).all()

    if not students:
        raise HTTPException(status_code=404, detail="No students found")

    for student in students:

        # ✅ Update semester if provided
        if payload.new_semester is not None:
            student.semester = payload.new_semester
            student.year = (payload.new_semester + 1) // 2

        # ✅ Update section if provided
        if payload.new_section is not None:
            student.section = payload.new_section

    db.commit()

    return {
        "message": "Students updated successfully",
        "updated_count": len(students)
    }


# =========================
# ADMIN – UPDATE SINGLE STUDENT
# =========================
@app.put("/admin/students/{student_id}")
def update_student(
    student_id: int,
    data: dict,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    # 🔐 Admin only
    if current_user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Admin only")

    student = db.query(Student).filter(Student.student_id == student_id).first()

    if not student:
        raise HTTPException(status_code=404, detail="Student not found")

    # 🔄 Update only provided fields
    for key, value in data.items():
        if hasattr(student, key):
            setattr(student, key, value)

    db.commit()
    db.refresh(student)

    return {
        "message": "Student updated successfully",
        "student": student
    }


# =========================
# ADMIN – DELETE STUDENTS (SINGLE + BULK)
# =========================
@app.delete("/admin/students")
def delete_students(
    payload: StudentDeleteRequest,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Admin only")

    if not payload.student_ids:
        raise HTTPException(status_code=400, detail="No students selected")

    students = db.query(Student).filter(
        Student.student_id.in_(payload.student_ids),
        Student.is_deleted == False
    ).all()

    if not students:
        raise HTTPException(status_code=404, detail="Students not found")

    for student in students:
        student.is_deleted = True
        student.deleted_at = datetime.utcnow()

    db.commit()

    return {
        "message": "Students marked as deleted",
        "deleted_count": len(students)
    }


# =========================
# ADMIN – DOWNLOAD RISK REPORT PDF
# =========================
@app.post("/admin/students/risk-report")
def download_risk_report(
    filters: dict = Body(...),
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):

    # 🔒 Admin check
    if current_user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Admin only")

    # 🔽 Extract filters from frontend JSON
    year = filters.get("year")
    section = filters.get("section")
    search = filters.get("search")

    # 🔽 Base query
    query = db.query(Student, User).join(
        User, Student.student_id == User.user_id
    ).filter(
        Student.is_deleted == False,
        User.is_deleted == False
    )

    # 🔽 Apply filters safely
    if year and year != "All":
        query = query.filter(Student.year == int(year))

    if section and section != "All":
        query = query.filter(Student.section == section)

    if search:
        query = query.filter(
            User.name.ilike(f"%{search}%") |
            Student.roll_no.ilike(f"%{search}%")
        )

    students = query.all()

    # =========================
    # PDF GENERATION
    # =========================

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    elements = []

    styles = getSampleStyleSheet()
    elements.append(Paragraph("Risk Student Report", styles["Heading1"]))
    elements.append(Spacer(1, 20))

    data = [["Roll No", "Name", "Year", "Section", "Attendance %", "Risk"]]

    for student, user in students:

        total = db.query(Attendance).filter(
            Attendance.student_id == student.student_id
        ).count()

        present = db.query(Attendance).filter(
            Attendance.student_id == student.student_id,
            Attendance.status == True
        ).count()

        percentage = round((present / total) * 100, 2) if total > 0 else 0

        if percentage < 60:
            risk = "Critical"
        elif percentage < (get_setting("attendance_threshold") or 75):
            risk = "Warning"
        else:
            continue  # Only include risk students

        data.append([
            student.roll_no,
            user.name,
            student.year,
            student.section,
            f"{percentage}%",
            risk
        ])

    # If no risk students
    if len(data) == 1:
        data.append(["-", "No Risk Students Found", "-", "-", "-", "-"])

    table = Table(data, repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
        ("ALIGN", (4, 1), (-1, -1), "CENTER"),
    ]))

    elements.append(table)

    doc.build(elements)

    buffer.seek(0)

    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={
            "Content-Disposition": "attachment; filename=risk_students.pdf"
        }
    )

# -------------------------
# ADMIN PROTECTED
# -------------------------
@app.get("/admin/protected")
def admin_protected(user=Depends(get_current_user)):
    if user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Not authorized")

    return {
        "message": "JWT works! Admin access granted",
        "user": user
    }


# =========================
# ADMIN – UPLOAD TIMETABLE
# =========================
@app.post("/admin/timetable/upload", response_model=TimetableResponse)
def upload_timetable(
    title: str = Form(...),
    timetable_type: str = Form(...),
    faculty_id: int = Form(None),


    department: str = Form(None),
    year: str = Form(None),
    section: str = Form(None),
    semester: str = Form(None),

    audience: str = Form("students"),
    file: UploadFile = File(...),

    user=Depends(get_current_user),
    db: Session = Depends(get_db)
):
    # 🔐 Only admin
    if user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Admin only")

    # 📁 Ensure directory exists
    upload_dir = "uploads/timetables"
    os.makedirs(upload_dir, exist_ok=True)

    # 🔹 Unique filename
    timestamp = int(datetime.utcnow().timestamp())
    filename = f"{timestamp}_{file.filename}"
    file_path = os.path.join(upload_dir, filename)

    # 💾 Save file
    with open(file_path, "wb") as f:
        f.write(file.file.read())

    # 📄 File type
    file_type = file.filename.split(".")[-1].lower()

   

    # 🔥 STEP 2: Create new timetable
    timetable = Timetable(
        title=title,
        timetable_type=timetable_type,
        faculty_id=faculty_id,


        department=department,
        year=year,
        section=section,
        semester=semester,

        file_name=file.filename,
        file_url=f"/uploads/timetables/{filename}",
        file_type=file_type,

        audience=audience,
        uploaded_by=user["user_id"],
        is_active=True
    )

    db.add(timetable)
    db.commit()
    db.refresh(timetable)

    # =========================
    # AUTO CREATE ALERT
    # =========================

    # CASE 1: Specific faculty selected
    if audience == "faculty" and faculty_id:

        new_alert = Alert(
            title="New Timetable Uploaded",
            message=f"{title} has been uploaded. Please check the timetable section.",
           type=timetable_type.lower(),
            target_role="faculty",
            target_type="individual",
            faculty_id=faculty_id
        )

        db.add(new_alert)
        db.commit()
        db.refresh(new_alert)

        recipient = AlertRecipient(
            alert_id=new_alert.id,
            user_id=faculty_id,
            is_read=False
        )

        db.add(recipient)
        db.commit()


    # CASE 2: Broadcast logic
    else:

        if audience == "students":
            roles = ["student"]

        elif audience == "faculty":
            roles = ["faculty"]

        elif audience == "both":
            roles = ["student", "faculty"]

        elif audience == "all":
            roles = ["student", "faculty", "admin"]

        else:
            roles = []

        for role in roles:

            new_alert = Alert(
                title="New Timetable Uploaded",
                message=f"{title} has been uploaded. Please check the timetable section.",
                type="timetable",
                target_role=role,
                target_type="all"
            )

            db.add(new_alert)
            db.commit()
            db.refresh(new_alert)

            query = db.query(User).filter(
                User.role == role,
                User.is_deleted == False
            )

            # 🔥 Filter by department if provided
            if department:
                department_id = None
                for key, value in DEPARTMENT_MAP.items():
                    if value == department:
                        department_id = key

                if department_id:
                    query = query.filter(User.department_id == department_id)

            users = query.all()


            for user_obj in users:
                recipient = AlertRecipient(
                    alert_id=new_alert.id,
                    user_id=user_obj.user_id,
                    is_read=False
                )
                db.add(recipient)

            db.commit()


    return timetable

# =========================
# GET PUBLISHED TIMETABLES
# =========================
@app.get("/timetables", response_model=list[TimetableResponse])
def get_timetables(
    faculty_id: Optional[int] = None,
    timetable_type: Optional[str] = None,
    department: Optional[str] = None,
    year: Optional[str] = None,
    semester: Optional[str] = None,
    section: Optional[str] = None,
    audience: Optional[str] = None,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):


    query = db.query(Timetable).filter(Timetable.is_active == True)

    # =============================
    # ROLE BASED SECURITY FILTER
    # =============================

    user_role = current_user["role"]
    user_department_id = current_user["department_id"]
    user_department = None
    if user_department_id:
        user_department = DEPARTMENT_MAP.get(user_department_id)

    if user_role == "student":
        query = query.filter(
            Timetable.department == user_department
        ).filter(
            Timetable.audience.in_(["students", "both", "all"])
        )

        if timetable_type:
            query = query.filter(
                Timetable.timetable_type.ilike(f"%{timetable_type}%")
            )


    elif user_role == "faculty":
        query = query.filter(
            (Timetable.department == user_department) |
            (Timetable.faculty_id == current_user["user_id"])
        ).filter(
            Timetable.audience.in_(["faculty", "both", "all"])
        )

    elif user_role == "admin":
        # admin can apply filters manually
        if department:
            query = query.filter(Timetable.department == department)
        if year:
            query = query.filter(Timetable.year == year)
        if semester:
            query = query.filter(Timetable.semester == semester)
        if section:
            query = query.filter(Timetable.section == section)
        if audience:
            query = query.filter(Timetable.audience == audience)

        if timetable_type:
            query = query.filter(
               Timetable.timetable_type.ilike(f"%{timetable_type}%")
            )


    return query.order_by(Timetable.uploaded_at.desc()).all()



# =========================
# DELETE TIMETABLES
# =========================
@app.delete("/admin/timetables/{timetable_id}")
def delete_timetable(
    timetable_id: int,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Not authorized")

    timetable = db.query(Timetable).filter(
        Timetable.id == timetable_id
    ).first()

    if not timetable:
        raise HTTPException(status_code=404, detail="Timetable not found")

    timetable.is_active = False
    db.commit()

    return {"message": "Timetable deleted successfully"}


# =========================
# ADMIN – GET ALL TEACHERS
# =========================
@app.get("/admin/teachers")
def get_all_teachers(
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Admin only")

    teachers = (
        db.query(Faculty, User)
        .join(User, Faculty.faculty_id == User.user_id)
        .filter(User.is_deleted == False)
        .all()
    )

    result = []

    for faculty, user in teachers:
        assignments = (
            db.query(FacultySubject, Subject)
            .join(Subject, FacultySubject.subject_id == Subject.subject_id)
            .filter(
                FacultySubject.faculty_id == user.user_id,
                FacultySubject.is_active == True
            )
            .all()
        )

        assigned_subjects = [
            {
                "assignment_id": fs.id,
                "subject_name": subject.subject_name,
                "year": fs.year,
                "section": fs.section,
                "semester": subject.semester
            }
            for fs, subject in assignments
        ]
        result.append({
            "id": user.user_id,
            "name": user.name,
            "department": DEPARTMENT_MAP.get(user.department_id, "UNKNOWN"),
            "designation": faculty.designation,
            "experience": faculty.experience,
            "email": user.email,
            "phone": faculty.phone,
            "subjects": faculty.expertise.split(",") if faculty.expertise else [],
            "alertsSent": 0,
            "classes": json.loads(faculty.classes) if faculty.classes else [],
            "assigned_subjects": assigned_subjects 
        })

    return result



# =========================
# ADMIN – UPDATE TEACHER
# =========================
@app.put("/admin/teachers/{teacher_id}")
def update_teacher(
    teacher_id: int,
    data: TeacherAdminUpdate,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Admin only")

    faculty = db.query(Faculty).filter(
        Faculty.faculty_id == teacher_id
    ).first()

    user = db.query(User).filter(
        User.user_id == teacher_id
    ).first()

    if not faculty or not user:
        raise HTTPException(status_code=404, detail="Teacher not found")

    # Update designation (Faculty table)
    if data.designation is not None:
     faculty.designation = data.designation


    # Update department (User table)
    if data.department_id is not None:
        user.department_id = data.department_id

    db.commit()

    return {"message": "Teacher updated successfully"}

# =========================
# ADMIN – DELETE TEACHERS (BULK + SOFT DELETE)
# =========================
@app.delete("/admin/teachers")
def delete_teachers(
    payload: TeacherDeleteRequest,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Admin only")

    teachers = db.query(User).filter(
        User.user_id.in_(payload.teacher_ids),
        User.role == "faculty",
        User.is_deleted == False
    ).all()

    if not teachers:
        raise HTTPException(status_code=404, detail="Teachers not found")

    for teacher in teachers:
        teacher.is_deleted = True
        teacher.deleted_at = datetime.utcnow()

    db.commit()

    return {
        "message": "Teachers deleted successfully",
        "deleted_count": len(teachers)
    }

# =========================
# ADMIN – CREATE ALERT
# ======================== 
@app.post("/admin/alerts")
def create_alert(
    title: str = Form(...),
    message: str = Form(...),
    type: str = Form(...),
    target_role: str = Form(...),
    target_type: str = Form(...),

    department: str = Form(None),
    faculty_id: int = Form(None),
    student_id: int = Form(None),

    file: UploadFile = File(None),

    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Admin only")

    # -------------------------
    # VALIDATION
    # -------------------------
    if target_type == "individual":
        if target_role == "faculty" and not faculty_id:
            raise HTTPException(status_code=400, detail="Faculty ID required")
        if target_role == "student" and not student_id:
            raise HTTPException(status_code=400, detail="Student ID required")

    if target_type == "department" and not department:
        raise HTTPException(status_code=400, detail="Department required")

    # -------------------------
    # FILE HANDLING
    # -------------------------
    file_name = None
    file_path = None
    file_type = None

    if file:
        upload_dir = "uploads/alerts"
        os.makedirs(upload_dir, exist_ok=True)

        unique_filename = f"{uuid.uuid4()}_{file.filename}"
        file_path = os.path.join(upload_dir, unique_filename)

        with open(file_path, "wb") as f:
            f.write(file.file.read())

        file_name = file.filename
        file_type = file.filename.split(".")[-1]

    # -------------------------
    # CREATE ALERT
    # -------------------------
    new_alert = Alert(
        title=title,
        message=message,
        type=type.lower(),
        target_role=target_role,
        target_type=target_type,
        department=department,
        faculty_id=faculty_id,
        student_id=student_id,
        file_name=file_name,
        file_path=file_path,
        file_type=file_type
    )

    db.add(new_alert)
    db.commit()
    db.refresh(new_alert)

    # -------------------------
    # CREATE RECIPIENTS
    # -------------------------
    users = []

    if target_type == "all":
        users = db.query(User).filter(
            User.role == target_role,
            User.is_deleted == False
        ).all()

    elif target_type == "individual":

        if target_role == "faculty" and faculty_id:
            users = db.query(User).filter(
                User.user_id == faculty_id,
                User.role == "faculty",
                User.is_deleted == False
            ).all()

        elif target_role == "student" and student_id:
            users = db.query(User).filter(
                User.user_id == student_id,
                User.role == "student",
                User.is_deleted == False
            ).all()

        else:
            users = []


    elif target_type == "department":
        department_id = None
        for key, value in DEPARTMENT_MAP.items():
            if value == department:
                department_id = key

        users = db.query(User).filter(
            User.role == target_role,
            User.department_id == department_id,
            User.is_deleted == False
        ).all()

    for user in users:
        recipient = AlertRecipient(
            alert_id=new_alert.id,
            user_id=user.user_id,
            is_read=False
        )
        print("TARGET ROLE:", target_role)
        print("TARGET TYPE:", target_type)
        print("STUDENT ID:", student_id)
        print("USERS FOUND:", users)

        db.add(recipient)

    db.commit()

    return {"message": "Alert created successfully"}


# =========================
# ADMIN – GET ALL ALERTS
# =========================
@app.get("/admin/alerts")
def get_all_alerts(
    role: Optional[str] = None,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Admin only")

    query = db.query(Alert)

    # ❌ Exclude timetable auto alerts
    query = query.filter(Alert.type != "timetable")

    # ✅ Filter by role (faculty or student)
    if role:
        query = query.filter(Alert.target_role == role)

    alerts = query.order_by(Alert.created_at.desc()).all()

    return [
        {
            "id": a.id,
            "title": a.title,
            "message": a.message,
            "type": a.type,
            "target_role": a.target_role,
            "target_type": a.target_type,
            "department": a.department,
            "faculty_id": a.faculty_id,
            "student_id": a.student_id,
            "created_at": a.created_at,
        }
        for a in alerts
    ]


# =========================
# ADMIN – DELETE ALERT
# =========================
@app.delete("/admin/alerts/{alert_id}")
def delete_alert(
    alert_id: int,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Admin only")

    alert = db.query(Alert).filter(Alert.id == alert_id).first()

    if not alert:
        raise HTTPException(status_code=404, detail="Alert not found")

    # delete recipients first
    db.query(AlertRecipient).filter(
        AlertRecipient.alert_id == alert_id
    ).delete()

    # delete alert
    db.delete(alert)
    db.commit()

    return {"message": "Alert deleted successfully"}


# =========================
# FACULTY – CREATE ALERT
# ======================== 
@app.post("/faculty/alerts")
def create_faculty_alert(
    title: str = Form(...),
    message: str = Form(...),
    type: str = Form(...),
    target_role: str = Form(...),
    target_type: str = Form(...),

    department: str = Form(None),
    faculty_id: int = Form(None),
    student_id: int = Form(None),

    file: UploadFile = File(None),

    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "faculty" and current_user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Not authorized")

    # -------------------------
    # VALIDATION
    # -------------------------
    if target_type == "individual":
        if target_role == "student" and not student_id:
            raise HTTPException(status_code=400, detail="Student ID required")

    if target_type == "department" and not department:
        raise HTTPException(status_code=400, detail="Department required")

    # -------------------------
    # FILE HANDLING
    # -------------------------
    file_name = None
    file_path = None
    file_type = None

    if file:
        upload_dir = "uploads/alerts"
        os.makedirs(upload_dir, exist_ok=True)

        unique_filename = f"{uuid.uuid4()}_{file.filename}"
        file_path = os.path.join(upload_dir, unique_filename)

        with open(file_path, "wb") as f:
            f.write(file.file.read())

        file_name = file.filename
        file_type = file.filename.split(".")[-1]

    # -------------------------
    # CREATE ALERT
    # -------------------------
    new_alert = Alert(
        title=title,
        message=message,
        type=type.lower(),
        target_role=target_role,
        target_type=target_type,
        department=department,
        faculty_id=current_user["user_id"],
        student_id=student_id,
        file_name=file_name,
        file_path=file_path,
        file_type=file_type
    )

    db.add(new_alert)
    db.commit()
    db.refresh(new_alert)

    # -------------------------
    # CREATE RECIPIENTS
    # -------------------------
    users = []

    if target_type == "all":
        users = db.query(User).filter(
            User.role == target_role,
            User.is_deleted == False
        ).all()

    elif target_type == "individual":
        if target_role == "student" and student_id:
            users = db.query(User).filter(
                User.user_id == student_id,
                User.role == "student",
                User.is_deleted == False
            ).all()

    elif target_type == "department":
        department_id = None
        for key, value in DEPARTMENT_MAP.items():
            if value == department:
                department_id = key

        users = db.query(User).filter(
            User.role == target_role,
            User.department_id == department_id,
            User.is_deleted == False
        ).all()

    for user in users:
        recipient = AlertRecipient(
            alert_id=new_alert.id,
            user_id=user.user_id,
            is_read=False
        )
        db.add(recipient)

    db.commit()

    return {"message": "Alert created successfully"}

# =========================
# FACULTY – GET ALERTS
# =========================
@app.get("/faculty/alerts")
def get_faculty_alerts(
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "faculty":
        raise HTTPException(status_code=403, detail="Faculty only")

    alerts = (
        db.query(Alert, AlertRecipient)
        .join(AlertRecipient, Alert.id == AlertRecipient.alert_id)
        .filter(AlertRecipient.user_id == current_user["user_id"])
        .order_by(Alert.created_at.desc())
        .all()
    )

    result = []

    for alert, recipient in alerts:
        result.append({
            "id": alert.id,
            "title": alert.title,
            "message": alert.message,
            "type": alert.type,
            "created_at": alert.created_at,
            "is_read": recipient.is_read,
            "file_name": alert.file_name,
            "file_path": alert.file_path,
            "file_type": alert.file_type
        })


    return result


# =========================
# FACULTY – GET MY ASSIGNED SUBJECTS
# =========================
@app.get("/faculty/my-subjects")
def get_my_subjects(
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "faculty":
        raise HTTPException(status_code=403, detail="Faculty only")

    assignments = (
        db.query(FacultySubject, Subject)
        .join(Subject, FacultySubject.subject_id == Subject.subject_id)
        .filter(
            FacultySubject.faculty_id == current_user["user_id"],
            FacultySubject.is_active == True
        )
        .all()
    )

    return [
    {
        "subject_id": s.subject_id,
        "subject_name": s.subject_name,
        "year": fs.year,
        "section": fs.section,
        "semester": s.semester,
        "department": DEPARTMENT_MAP.get(
            db.query(User)
            .filter(User.user_id == current_user["user_id"])
            .first()
            .department_id
        )
    }
    for fs, s in assignments
 ]


# =========================
# ADMIN / DEBUG - GET MARKS
# =========================
@app.get("/debug-marks")
def debug_marks(db: Session = Depends(get_db)):
    from models import Mark  # type: ignore
    marks = db.query(Mark).all()
    return [
        {
            "student_id": m.student_id,
            "subject_id": m.subject_id,
            "exam": m.exam,
            "year": m.year,
            "section": m.section,
            "marks": m.marks
        } for m in marks
    ]

# =========================
# FACULTY – GET MARKS
# =========================
@app.get("/faculty/marks")
def get_faculty_marks(
    year: int = Query(...),
    section: str = Query(...),
    subject_id: int = Query(...),
    exam: str = Query(...),
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "faculty":
        raise HTTPException(status_code=403, detail="Faculty only")

    from sqlalchemy import func  # type: ignore

    print("FETCH EXAM:", exam)
    print("SUBJECT ID:", subject_id)

    if exam == "Total":
        results = db.query(Student, User, ScaledMark).join(  # type: ignore
            User, Student.student_id == User.user_id
        ).outerjoin(
            ScaledMark, (Student.student_id == ScaledMark.student_id) &   # type: ignore
                        (ScaledMark.subject_id == subject_id)  # type: ignore
        ).filter(
            Student.year == year,
            Student.section == section
        ).all()

        result = []
        values = []
        for student, user, sm in results:
            extra = {}
            m_val = None
            if sm:
                extra["Assignment"] = float(sm.assignment_scaled) if sm.assignment_scaled is not None else 0
                extra["Mid 1"] = float(sm.mid1_scaled) if sm.mid1_scaled is not None else 0
                extra["Mid 2"] = float(sm.mid2_scaled) if sm.mid2_scaled is not None else 0
                extra["Semester"] = float(sm.semester_marks) if sm.semester_marks is not None else 0
                extra["Internal"] = float(sm.internal_total) if sm.internal_total is not None else 0
                extra["Final"] = float(sm.final_total) if sm.final_total is not None else 0
                m_val = float(sm.final_total) if sm.final_total is not None else 0
                if m_val > 0 or sm.internal_total:
                    values.append(m_val)

            result.append({
                "student_id": student.student_id,
                "name": user.name,
                "roll_no": student.roll_no,
                "marks": m_val,
                "extra_data": extra
            })
    else:
        results = db.query(Student, User, Mark).join(
            User, Student.student_id == User.user_id
        ).outerjoin(
            Mark, 
            (Student.student_id == Mark.student_id) & 
            (Mark.subject_id == subject_id) & 
            (func.lower(func.trim(Mark.exam)) == func.lower(func.trim(exam))) &
            (Mark.year == year) &
            (Mark.section == section)
        ).filter(
            Student.year == year,
            Student.section == section
        ).order_by(Student.roll_no).all()

        result = []
        values = []
        for student, user, mark in results:
            m_val = float(mark.marks) if mark and mark.marks is not None else None
            extra = mark.extra_data if mark and mark.extra_data else {}
            result.append({
                "student_id": student.student_id,
                "name": user.name,
                "roll_no": student.roll_no,
                "marks": m_val,
                "extra_data": extra
            })
            if m_val is not None:
                values.append(m_val)

    FAIL_THRESHOLD = 10  # change if needed
    fail_count = len([v for v in values if v < FAIL_THRESHOLD])
    avg = round(sum(values) / len(values), 2) if values else 0  # type: ignore
    highest = max(values) if values else 0

    return {
        "students": result,
        "stats": {
            "average": avg,
            "highest": highest,
            "fail_count": fail_count,
            "total_students": len(result)
        },
        "available_columns": ["Marks"]
    }


# =========================
# REUSABLE MARKS SAVE FUNCTION
# =========================
def save_or_update_marks(db, student_id, subject_id, exam, data, value=None):
    """Reusable function to save or update marks"""
    print(f"Saving marks for student_id: {student_id}, subject_id: {subject_id}, exam: {exam}, value: {value}")

    existing = db.query(Mark).filter(
        Mark.student_id == student_id,
        Mark.subject_id == subject_id,
        Mark.exam == exam
    ).first()

    # Determine which field to update based on exam
    field_name = None
    if exam == "Mid-1":
        field_name = "mid1"
    elif exam == "Mid-2":
        field_name = "mid2"
    elif exam.startswith("Assignment"):
        field_name = "assignment_total"
    elif exam == "Semester":
        field_name = "semester"

    if existing:
        # Check if value actually changed
        current_value = getattr(existing, field_name, 0) or 0  # type: ignore
        if current_value != value:
            print(f"Updating existing mark for student {student_id}: {field_name} from {current_value} to {value}")
            setattr(existing, field_name, value)  # type: ignore
            return True  # Value changed
        else:
            print(f"No change needed for student {student_id}: {field_name} is already {value}")
            return False  # No change
    else:
        print(f"Creating new mark for student {student_id}")
        new_mark = Mark(
            student_id=student_id,
            subject_id=subject_id,
            exam=exam,
            assignment_total=0,
            mid1=0,
            mid2=0,
            semester=0,
            total=0,
            sgpa=0,
            cgpa=0,
            year=data["year"],
            section=data["section"],
            faculty_id=data.get("faculty_id")
        )
        # Set the specific field for new record
        if field_name:
            setattr(new_mark, field_name, value)
        db.add(new_mark)
        return True  # New record created

# =========================
# FACULTY – PREVIEW MARKS EXCEL
# =========================
@app.post("/faculty/marks/preview")
async def preview_marks_excel(
    file: UploadFile = File(...),
    year: int = Query(...),
    section: str = Query(...),
    subject_id: int = Query(...),
    exam: str = Query(...),
    overwrite: str = Form("false"),
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "faculty":
        raise HTTPException(status_code=403, detail="Faculty only")

    try:
        import pandas as pd  # type: ignore
        from io import BytesIO
    except ImportError:
        return {"error": "pandas not installed"}

    try:
        contents = await file.read()
        df = pd.read_excel(BytesIO(contents))
        df.columns = df.columns.str.strip()
        df.columns = df.columns.str.replace("  ", " ")
        print("Exam:", exam)
        columns = df.columns.tolist()

        is_scaled = "Final Total" in columns or "Assignment Total (Scaled to 10)" in columns
        column_name = None

        if is_scaled:
            column_name = "Final Total" if "Final Total" in columns else "Assignment Total (Scaled to 10)"
        else:
            for col in columns:
                if col.strip().lower() == exam.strip().lower():
                    column_name = col
                    break
            if not column_name:
                exam_clean = exam.replace("-", " ").strip().lower()
                for col in columns:
                    if col.replace("-", " ").strip().lower() == exam_clean:
                        column_name = col
                        break
                        
            if not column_name:
                return {"error": f"Column for given exam '{exam}' not found in Excel"}
        print("Column:", column_name)
        print("Excel Columns:", df.columns.tolist())
    except Exception as e:
        return {"error": f"Excel read failed: {str(e)}"}

    required = ["Register Number", "Student Name"]
    for col in required:
        if col not in df.columns:
            return {"error": f"Missing column: {col}"}
    if column_name not in df.columns:
        return {"error": f"{column_name} column not found in Excel"}
    if df.empty:
        return {"error": "Excel file is empty"}

    students = db.query(Student).filter(
        Student.year == year,
        Student.section == section
    ).all()

    def normalize_reg(val):
        if pd.isna(val):
            return None
        val = str(val).strip()
        if val.endswith(".0"):
            val = val[:-2]  # type: ignore
        return val

    valid_regs = set(
        normalize_reg(s.roll_no)
        for s in students if s.roll_no
    )

    excel_regs = df["Register Number"].apply(normalize_reg)
    invalid_rows = df[~excel_regs.isin(valid_regs)]
    if not invalid_rows.empty:
        return {
            "error": "File contains students not matching selected Year/Section"
        }

    preview = []
    for idx, row in df.iterrows():
        reg = normalize_reg(row["Register Number"])
        value = row.get(column_name)
        
        student = db.query(Student).filter(
            Student.roll_no == reg,
            Student.year == year,
            Student.section == section
        ).first()
        
        if not student:
            preview.append({
                "register_number": row["Register Number"],
                "name": row["Student Name"],
                "marks": "" if pd.isna(value) else value,
                "status": "invalid"
            })
            continue
        
        record = db.query(Mark).filter_by(
            student_id=student.student_id,
            subject_id=subject_id,
            exam=exam
        ).first()
        
        existing = False
        if record and record.marks is not None:
            existing = True
        
        preview.append({
            "register_number": row["Register Number"],
            "name": row["Student Name"],
            "marks": "" if pd.isna(value) else value,
            "status": "exists" if existing else "new"
        })

    return {"preview": preview}

from pydantic import BaseModel  # type: ignore
class ValidateStudentsRequest(BaseModel):
    register_numbers: list[str]
    year: int
    section: str

@app.post("/faculty/marks/validate-students")
def validate_students(req: ValidateStudentsRequest, current_user=Depends(get_current_user), db: Session=Depends(get_db)):
    if current_user["role"] != "faculty":
        return {"success": False, "error": "Faculty only"}
    
    try:
        from models import Student  # type: ignore
        valid_students = db.query(Student.roll_no).filter(
            Student.year == req.year,
            Student.section == req.section
        ).all()
        valid_rolls = {s[0].strip().lower() for s in valid_students if s[0]}
        
        for r in req.register_numbers:
            if str(r).strip().lower() not in valid_rolls:
                return {"success": True, "valid": False, "message": f"Student data does not match selected Year/Section"}
        return {"success": True, "valid": True, "message": "All students valid"}
    except Exception as e:
        return {"success": False, "error": str(e)}


class AnalyzeUploadRequest(BaseModel):
    register_numbers: list[str]
    subject_id: int
    exam: str

@app.post("/faculty/marks/analyze-upload")
def analyze_upload(req: AnalyzeUploadRequest, current_user=Depends(get_current_user), db: Session=Depends(get_db)):
    if current_user["role"] != "faculty":
         return {"success": False, "error": "Faculty only"}

    try:
        from models import Mark, Student  # type: ignore
        existing_marks = db.query(Mark, Student).join(Student, Mark.student_id == Student.student_id).filter(
            Mark.subject_id == req.subject_id,
            Mark.exam == req.exam,
            Student.roll_no.in_(req.register_numbers)
        ).all()

        existing_rolls = {s.roll_no.strip().lower() for m, s in existing_marks}

        result = []
        for r in req.register_numbers:
            status = "existing" if str(r).strip().lower() in existing_rolls else "new"
            result.append({"register_number": r, "status": status})

        return {"success": True, "data": result}
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.post("/faculty/marks/upload")
async def upload_marks_excel(
    file: UploadFile = File(...),
    year: int = Query(...),
    section: str = Query(...),
    subject_id: int = Query(...),
    exam: str = Query(...),
    overwrite: str = Form("false"),
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "faculty":
        return {"success": False, "error": "Faculty only"}

    try:
        import pandas as pd  # type: ignore
        from io import BytesIO
        import hashlib
        from models import MarksUpload, ScalingLog, Mark, Student  # type: ignore

        contents = await file.read()
        file_hash = hashlib.sha256(contents).hexdigest()

        # NORMALIZE EXAM VALUE (CRITICAL)
        exam_raw = exam.strip().lower()
        if exam_raw in ["mid1", "mid-1", "mid 1"]:
            exam = "Mid-1"
        elif exam_raw in ["mid2", "mid-2", "mid 2"]:
            exam = "Mid-2"
        elif exam_raw in ["assignment1", "assignment-1", "assignment 1"]:
            exam = "Assignment-1"
        elif exam_raw in ["assignment2", "assignment-2", "assignment 2"]:
            exam = "Assignment-2"
        elif exam_raw in ["assignment3", "assignment-3", "assignment 3"]:
            exam = "Assignment-3"
        elif exam_raw in ["assignment4", "assignment-4", "assignment 4"]:
            exam = "Assignment-4"
        elif exam_raw in ["assignment5", "assignment-5", "assignment 5"]:
            exam = "Assignment-5"
        elif exam_raw in ["semester"]:
            exam = "Semester"
        elif exam_raw in ["total"]:
            exam = "Total"


        if overwrite.lower() != "true":
            duplicate = db.query(MarksUpload).filter(
                MarksUpload.file_hash == file_hash,
                MarksUpload.subject_id == subject_id,
                MarksUpload.year == year,
                MarksUpload.section == section,
                MarksUpload.exam == exam
            ).first()
            if duplicate:
                return {"duplicate": True, "message": "This file was already uploaded."}

        df = pd.read_excel(BytesIO(contents))
        df.columns = df.columns.str.strip()
        df.columns = df.columns.str.replace("  ", " ")

        columns = df.columns.tolist()
        
        headers = {str(h).strip().lower(): h for h in columns}

        def get_col(name):
            for key in headers:
                if name in key:
                    return headers[key]
            return None

        col_reg = get_col("register")
        col_mid1 = get_col("mid 1")
        col_mid2 = get_col("mid 2")

        assignment_cols = [
            col for col in df.columns
            if "assignment" in col.lower() and "total" not in col.lower()
        ]

        column_name = None
        for col in columns:
            if col.strip().lower() == exam.strip().lower():  # type: ignore
                column_name = col
                break
        if not column_name:
            exam_clean = exam.replace("-", " ").strip().lower()  # type: ignore
            for col in columns:
                if col.replace("-", " ").strip().lower() == exam_clean:
                    column_name = col
                    break
                    
        if not column_name:
            return {"success": False, "error": f"Column for given exam '{exam}' not found in Excel"}

        reserved_cols = {"register number", "student name"}
        if column_name:
            reserved_cols.add(column_name.lower())
        extra_columns_list = [c for c in columns if c.lower() not in reserved_cols]

        required = ["Register Number", "Student Name"]
        for col in required:
            if col not in df.columns:
                return {"success": False, "error": f"Missing column: {col}"}
        if df.empty:
            return {"success": False, "error": "Excel file is empty"}

        students = db.query(Student).filter(
            Student.year == year,
            Student.section == section
        ).all()

        def normalize_reg(val):
            if pd.isna(val): return None
            val = str(val).strip()
            if val.endswith(".0"): val = val[:-2]  # type: ignore
            return val

        valid_regs = set(normalize_reg(s.roll_no) for s in students if s.roll_no)

        def safe_get(row, col):
            return row[col] if col and col in row and not pd.isna(row[col]) else None

        overwrite_flag = overwrite.lower() == "true"
        updated = 0
        skipped = 0
        already_exists = 0

        # ATOMIC TRANSACTION BLOCK
        for _, row in df.iterrows():
            reg = str(row.get("Register Number") or row.get("register_number") or "").strip()
            if reg.endswith(".0"): reg = reg[:-2]  # type: ignore
            name = str(row.get("Student Name") or row.get("name") or "").strip()
            
            if not reg or reg not in valid_regs:
                skipped += 1
                continue

            student = db.query(Student).filter(
                Student.roll_no == reg,
                Student.year == year,
                Student.section == section
            ).first()

            if not student:
                skipped += 1
                continue

            # Removed the buggy explicit Mark object query here, replaced with proper UPSERT below.

            from models import ScaledMark  # type: ignore
            sm_record = db.query(ScaledMark).filter_by(
                student_id=student.student_id, subject_id=subject_id
            ).first()
            if not sm_record:
                sm_record = ScaledMark(
                    student_id=student.student_id,
                    subject_id=subject_id,
                    year=year,
                    section=section
                )
                db.add(sm_record)
            
            def safe(val):
                try:
                    return float(val)
                except:
                    return 0

            assignment_sum = sum([safe(row[c]) for c in assignment_cols])
            max_assignment_marks = len(assignment_cols) * 10
            assignment_total = (
                (assignment_sum / max_assignment_marks) * 10
                if max_assignment_marks > 0 else 0
            )

            mid1 = safe(row.get(col_mid1))
            mid2 = safe(row.get(col_mid2))

            mid1_scaled = (mid1 / 30) * 20 if mid1 > 0 else 0
            mid2_scaled = (mid2 / 30) * 20 if mid2 > 0 else 0

            valid_mids = []
            if mid1 > 0:
                valid_mids.append(mid1_scaled)
            if mid2 > 0:
                valid_mids.append(mid2_scaled)

            if valid_mids:
                mid_combined = (sum(valid_mids) / (len(valid_mids) * 20)) * 20  # type: ignore
            else:
                mid_combined = 0

            internal_total = assignment_total + mid_combined

            print({
                "a_total": assignment_total,
                "mid1_scaled": mid1_scaled,
                "mid2_scaled": mid2_scaled,
                "mid_combined": mid_combined,
                "internal": internal_total
            })

            sm_record.assignment_scaled = assignment_total
            sm_record.mid1_scaled = mid1_scaled
            sm_record.mid2_scaled = mid2_scaled
            sm_record.mid_combined = mid_combined
            sm_record.internal_total = internal_total
            
            # Removed the erroneous updated += 1 and continue here

            if column_name:
                val = safe_get(row, column_name)
                marks_value = None
                if val is not None and str(val).strip() not in ("", "-"):
                     try:
                         marks_value = float(val)
                     except:
                         marks_value = 0.0

                extra = {}
                for col in extra_columns_list:
                     v = row.get(col)
                     if not pd.isna(v) and str(v).strip() != "":
                         extra[col] = str(v)

                print("INSERT:", student.student_id, subject_id, exam, marks_value, year, section)

                existing = db.query(Mark).filter(
                    Mark.student_id == student.student_id,
                    Mark.subject_id == subject_id,
                    Mark.exam == exam
                ).first()

                if existing:
                    if not overwrite_flag and existing.marks is not None:
                        already_exists += 1  # type: ignore
                        continue
                    existing.marks = marks_value
                    existing.year = year
                    existing.section = section
                    existing.extra_data = extra
                else:
                    new_mark = Mark(
                        student_id=student.student_id,
                        subject_id=subject_id,
                        exam=exam,
                        marks=marks_value,
                        year=year,
                        section=section,
                        faculty_id=current_user["user_id"],  # type: ignore
                        extra_data=extra
                    )
                    db.add(new_mark)

            updated += 1

        new_upload = MarksUpload(
            faculty_id=current_user["user_id"],
            subject_id=subject_id,
            year=year,
            section=section,
            exam=exam,
            file_hash=file_hash
        )
        db.add(new_upload)
        
        db.commit()

        # Debug query to verify data
        saved = db.query(Mark).filter(
            Mark.year == year,
            Mark.section == section,
            Mark.subject_id == subject_id
        ).all()
        print("Saved count:", len(saved))

        # Decouple Logging (Non-Blocking)
        try:
            from models import ScalingLog  # type: ignore
            new_log = ScalingLog(
                faculty_id=current_user["user_id"],
                subject_id=subject_id,
                year=year,
                section=section,
                action_type="files_uploaded",
                file_name=file.filename
            )
            db.add(new_log)
            db.commit()
        except Exception as log_error:
            print("Logging failed:", log_error)
            db.rollback()

        marks_list = db.query(Mark).filter(
            Mark.subject_id == subject_id,
            Mark.year == year,
            Mark.section == section,
            Mark.exam == exam
        ).all()
        values = [m.marks for m in marks_list if m.marks is not None]
        avg = round(sum(values) / len(values), 2) if values else 0  # type: ignore
        highest = max(values) if values else 0

        available_columns = [column_name] + extra_columns_list  # type: ignore

        return {
            "success": True,
            "updated": updated,
            "skipped": skipped,
            "already_exists": already_exists,
            "message": f"{updated} students updated • {skipped} skipped • {already_exists} already existed",
            "available_columns": available_columns,
            "average": avg,
            "highest": highest
        }

    except Exception as e:
        db.rollback()
        return {"success": False, "error": str(e)}


# =========================
# MANUAL ENTRY REMOVED - Now using Excel upload only
# =========================


# =========================
# FACULTY – DOWNLOAD MARKS TEMPLATE
# =========================
@app.get("/faculty/marks/template")
def download_marks_template(
    year: int = Query(...),
    section: str = Query(...),
    subject_id: int = Query(...),
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "faculty":
        from fastapi.responses import JSONResponse  # type: ignore
        return JSONResponse(status_code=401, content={"detail": "Unauthorized"})

    print("Year:", year)
    print("Section:", section)

    # Get students in the class
    students = (
        db.query(Student, User)
        .join(User, Student.student_id == User.user_id)
        .filter(
            Student.year == year,
            Student.section == section
        )
        .all()
    )

    print("Students count:", len(students))

    if not students:
        return {"error": "No students found"}

    # Create DataFrame with simplified columns
    import pandas as pd  # type: ignore
    import io
    data = []
    for student, user in students:
        data.append({
            "Register Number": student.roll_no,
            "Student Name": user.name,
            "Marks": ""  # Single marks column
        })

    report_format = get_report_format("marks")
    if report_format == "excel":
        df = pd.DataFrame(data)
        output = io.BytesIO()
        df.to_excel(output, index=False, engine="openpyxl")
        output.seek(0)
        from fastapi.responses import StreamingResponse  # type: ignore
        return StreamingResponse(
            output,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": "attachment; filename=marks_template.xlsx"}
        )

    if report_format == "pdf":
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        styles = getSampleStyleSheet()
        elements = [Paragraph("Marks Template", styles["Heading1"]), Spacer(1, 0.2 * inch)]
        table_data = [["Register Number", "Student Name", "Marks"]] + [[row["Register Number"], row["Student Name"], ""] for row in data]
        table = Table(table_data, repeatRows=1)
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
        ]))
        elements.append(table)
        doc.build(elements)
        buffer.seek(0)
        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=marks_template.pdf"}
        )

    if report_format == "docx":
        try:
            from docx import Document
        except ImportError:
            raise HTTPException(status_code=500, detail="DOCX export support is unavailable")
        document = Document()
        document.add_heading("Marks Template", level=1)
        table_docx = document.add_table(rows=1, cols=3)
        hdr_cells = table_docx.rows[0].cells
        hdr_cells[0].text = "Register Number"
        hdr_cells[1].text = "Student Name"
        hdr_cells[2].text = "Marks"
        for row in data:
            cells = table_docx.add_row().cells
            cells[0].text = str(row["Register Number"])
            cells[1].text = str(row["Student Name"])
            cells[2].text = ""
        output = io.BytesIO()
        document.save(output)
        output.seek(0)
        return StreamingResponse(
            output,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=marks_template.docx"}
        )

    # fallback to excel
    df = pd.DataFrame(data)
    output = io.BytesIO()
    df.to_excel(output, index=False, engine="openpyxl")
    output.seek(0)
    from fastapi.responses import StreamingResponse  # type: ignore
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=marks_template.xlsx"}
    )


# =========================
# STUDENT – GET ALERTS
# =========================
@app.get("/student/alerts")
def get_student_alerts(
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "student":
        raise HTTPException(status_code=403, detail="Student only")

    alerts = (
        db.query(Alert, AlertRecipient)
        .join(AlertRecipient, Alert.id == AlertRecipient.alert_id)
        .filter(AlertRecipient.user_id == current_user["user_id"])
        .order_by(Alert.created_at.desc())
        .all()
    )

    risk_data = get_student_risk(
        student_id=current_user["user_id"],
        db=db,
        attendance_threshold=float(get_setting("attendance_threshold") or 75),
        cgpa_threshold=float(get_setting("cgpa_threshold") or 6.5),
    )

    result = []

    for alert, recipient in alerts:
        monitoring_type = alert.type in {"cgpa-monitor", "marks-monitor", "attendance-monitor"}
        if monitoring_type and not risk_data.get("has_valid_data"):
            continue

        if (
            alert.type == "cgpa-monitor"
            and "cgpa is 0.00" in str(alert.message or "").lower()
            and not risk_data.get("cgpa")
        ):
            continue

        result.append({
            "id": alert.id,
            "title": alert.title,
            "message": alert.message,
            "type": alert.type,
            "created_at": alert.created_at,
            "is_read": recipient.is_read
        })

    return result


# =========================
# STUDENT – GET MY MARKS
# =========================
@app.get("/student/my-marks")
def get_my_marks(
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    from models import Subject, Mark, ScaledMark, Student  # type: ignore
    if current_user["role"] != "student":
        raise HTTPException(status_code=403, detail="Student only")

    student_id = current_user["user_id"]
    student = db.query(Student).filter_by(student_id=student_id).first()
    if not student:
        return {"sgpa": None, "cgpa": None, "subjects": []}

    scaled_marks = db.query(ScaledMark).filter_by(student_id=student_id).all()
    raw_marks = db.query(Mark).filter_by(student_id=student_id).all()

    subject_ids = set([sm.subject_id for sm in scaled_marks] + [rm.subject_id for rm in raw_marks])
    subjects_info = db.query(Subject).filter(Subject.subject_id.in_(list(subject_ids))).all()
    subject_map = {s.subject_id: s.subject_name for s in subjects_info}

    cgpa_raw = float(student.cgpa) if student.cgpa is not None else None
    cgpa_val = cgpa_raw if (cgpa_raw is not None and cgpa_raw > 0) else None
    sgpas = [float(m.sgpa) for m in raw_marks if m.sgpa]
    sgpa_val = sum(sgpas)/len(sgpas) if sgpas else cgpa_val

    result_subjects = []
    for sub_id in subject_ids:
        sub_name = subject_map.get(sub_id, "Unknown")
        sm = next((x for x in scaled_marks if x.subject_id == sub_id), None)
        rms = [x for x in raw_marks if x.subject_id == sub_id]

        def get_raw(exam_name):
            rec = next((x for x in rms if x.exam == exam_name), None)
            if not rec: return "-"
            val = rec.marks
            if val is None:
                if rec.extra_data and exam_name in rec.extra_data:
                    try: return float(rec.extra_data[exam_name])
                    except: return "-"
                return "-"
            return float(val)

        def get_scl(attr):
            if not sm: return "-"
            val = getattr(sm, attr, None)
            return float(val) if val is not None else "-"

        result_subjects.append({
            "subject": sub_name,
            "assignments": {
                "A1": get_raw("Assignment-1"),
                "A2": get_raw("Assignment-2"),
                "A3": get_raw("Assignment-3"),
                "A4": get_raw("Assignment-4"),
                "A5": get_raw("Assignment-5")
            },
            "mid1": get_raw("Mid-1"),
            "mid2": get_raw("Mid-2"),
            "scaled": {
                "assignment_scaled": get_scl("assignment_scaled"),
                "mid_combined": get_scl("mid_combined"),
                "internal_total": get_scl("internal_total")
            },
            "semester": get_scl("semester_marks") if get_scl("semester_marks") != "-" else get_raw("Semester"),
            "final_total": get_scl("final_total") if get_scl("final_total") != "-" else get_raw("Total")
        })

    return {
        "sgpa": round(sgpa_val, 2) if sgpa_val is not None else None,
        "cgpa": round(cgpa_val, 2) if cgpa_val is not None else None,
        "subjects": result_subjects
    }


# =========================
# STUDENT – INSIGHTS (DATA-DRIVEN)
# =========================
@app.get("/student/insights")
def get_student_insights(
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "student":
        raise HTTPException(status_code=403, detail="Student only")

    student_id = current_user["user_id"]
    student_profile = db.query(Student).filter(Student.student_id == student_id).first()

    if not student_profile:
        raise HTTPException(status_code=404, detail="Student profile not found")

    attendance_threshold = float(get_setting("attendance_threshold") or 75)
    cgpa_threshold = float(get_setting("cgpa_threshold") or 6.5)
    risk_data = get_student_risk(
        student_id=student_id,
        db=db,
        attendance_threshold=attendance_threshold,
        cgpa_threshold=cgpa_threshold,
    )

    if not risk_data.get("has_valid_data"):
        return {
            "student_id": student_id,
            "has_valid_data": False,
            "no_data_message": NO_DATA_MESSAGE,
            "attendance": None,
            "attendance_trend": [],
            "mid1": None,
            "mid2": None,
            "cgpa": None,
            "risk_level": "INSUFFICIENT DATA",
            "primary_issue": "Insufficient data to analyze",
            "prediction": "Insufficient data to analyze",
            "actions": [],
            "daily_tasks": [],
            "warnings": [],
            "patterns": [],
            "early_warning_signals": [],
            "subject_intelligence": {"subjects": [], "weakest_subjects": [], "strongest_subject": None, "priority_subject": None},
            "placement_analysis": {
                "status": "INSUFFICIENT DATA",
                "reasons": ["Insufficient data to analyze"],
                "gaps": [],
                "roadmap": {"weekly": [], "monthly": []},
                "timeline": "Insufficient data to analyze",
                "risk_if_ignored": [],
            },
            "placement_readiness": "INSUFFICIENT DATA",
            "placement": {
                "readiness": None,
                "aptitude": None,
                "consistency": None,
                "consistency_interpretation": "Insufficient data",
            },
        }

    attendance_rows = (
        db.query(Attendance.subject_id, Attendance.attendance_date, Attendance.status)
        .filter(Attendance.student_id == student_id)
        .order_by(Attendance.attendance_date.asc())
        .all()
    )

    total_attendance = len(attendance_rows)
    present_attendance = sum(1 for row in attendance_rows if bool(row.status))
    attendance_percent = round((present_attendance / total_attendance) * 100, 2) if total_attendance > 0 else None

    attendance_by_day = {}
    for row in attendance_rows:
        day = row.attendance_date.isoformat() if row.attendance_date else None
        if not day:
            continue
        if day not in attendance_by_day:
            attendance_by_day[day] = {"present": 0, "total": 0}

        attendance_by_day[day]["total"] += 1
        if bool(row.status):
            attendance_by_day[day]["present"] += 1

    ordered_days = sorted(attendance_by_day.keys())
    attendance_trend_all = []
    for day in ordered_days:
        day_total = attendance_by_day[day]["total"]
        day_present = attendance_by_day[day]["present"]
        if day_total > 0:
            attendance_trend_all.append(round((day_present / day_total) * 100, 2))

    attendance_trend = attendance_trend_all[-7:]
    previous_attendance_window = attendance_trend_all[-14:-7]

    subject_attendance_map = {}
    for row in attendance_rows:
        subject_id = row.subject_id
        if subject_id is None:
            continue

        if subject_id not in subject_attendance_map:
            subject_attendance_map[subject_id] = {"present": 0, "total": 0}

        subject_attendance_map[subject_id]["total"] += 1
        if bool(row.status):
            subject_attendance_map[subject_id]["present"] += 1

    mark_rows = (
        db.query(Mark)
        .filter(Mark.student_id == student_id)
        .order_by(Mark.created_at.asc())
        .all()
    )

    def _normalize_exam_name(exam_name):
        if not exam_name:
            return ""
        return str(exam_name).strip().lower().replace("-", "").replace(" ", "")

    def _to_float(value):
        try:
            if value is None:
                return None
            return float(value)
        except Exception:
            return None

    mid1_values = []
    mid2_values = []
    marks_trend_values = []

    for row in mark_rows:
        exam_key = _normalize_exam_name(row.exam)
        marks_value = _to_float(row.marks)

        if exam_key in ["mid1", "mid01"] and marks_value is not None:
            mid1_values.append(marks_value)
            marks_trend_values.append(marks_value)
        elif exam_key in ["mid2", "mid02"] and marks_value is not None:
            mid2_values.append(marks_value)
            marks_trend_values.append(marks_value)

    if not mid1_values:
        for row in mark_rows:
            val = _to_float(row.mid1)
            if val is not None and val > 0:
                mid1_values.append(val)

    if not mid2_values:
        for row in mark_rows:
            val = _to_float(row.mid2)
            if val is not None and val > 0:
                mid2_values.append(val)

    mid1 = round(sum(mid1_values) / len(mid1_values), 2) if mid1_values else None
    mid2 = round(sum(mid2_values) / len(mid2_values), 2) if mid2_values else None

    # -----------------------------
    # Dynamic CGPA Prediction Engine
    # -----------------------------
    assignments = []
    excluded_assignments_without_max = 0

    def _extract_assignment_max(extra_data, exam_name, exam_key):
        if not isinstance(extra_data, dict):
            return None

        lowered = {str(k).strip().lower(): v for k, v in extra_data.items()}
        direct_candidates = [
            "max",
            "max_marks",
            "max_score",
            "out_of",
            f"{exam_name}_max".lower() if exam_name else None,
            f"{exam_key}_max".lower() if exam_key else None,
            f"{exam_key}max".lower() if exam_key else None,
        ]

        for candidate in direct_candidates:
            if candidate and candidate in lowered:
                val = _to_float(lowered[candidate])
                if val is not None and val > 0:
                    return val

        return None

    for row in mark_rows:
        exam_key = _normalize_exam_name(row.exam)
        if not exam_key.startswith("assignment"):
            continue

        score_value = _to_float(row.marks)
        if score_value is None:
            continue

        max_value = _extract_assignment_max(row.extra_data, row.exam, exam_key)
        if max_value is None:
            cohort_max_raw = (
                db.query(func.max(Mark.marks))
                .filter(
                    Mark.subject_id == row.subject_id,
                    Mark.exam == row.exam,
                    Mark.marks.isnot(None),
                )
                .scalar()
            )
            cohort_max = _to_float(cohort_max_raw)
            if cohort_max is not None and cohort_max > 0:
                max_value = cohort_max

        if max_value is None or max_value <= 0:
            excluded_assignments_without_max += 1
            continue

        assignments.append({"score": score_value, "max": max_value})

    if not assignments:
        for row in mark_rows:
            if not isinstance(row.extra_data, dict):
                continue
            lowered = {str(k).strip().lower(): v for k, v in row.extra_data.items()}
            for key, raw_value in lowered.items():
                if "assignment" not in key:
                    continue
                if "max" in key or "total" in key or "scaled" in key:
                    continue

                score_value = _to_float(raw_value)
                if score_value is None:
                    continue

                max_key = f"{key}_max"
                max_value = _to_float(lowered.get(max_key))
                if max_value is None or max_value <= 0:
                    excluded_assignments_without_max += 1
                    continue

                assignments.append({"score": score_value, "max": max_value})

    assignment_total = round(sum(item["score"] for item in assignments), 2)
    assignment_max_total = round(sum(item["max"] for item in assignments), 2)

    if assignment_max_total == 0:
        scaled_assignment = 0.0
    else:
        scaled_assignment = round((assignment_total / assignment_max_total) * 10, 2)

    assignment_confidence = "HIGH"
    if excluded_assignments_without_max > 0 and assignments:
        assignment_confidence = "MEDIUM"
    elif excluded_assignments_without_max > 0 and not assignments:
        assignment_confidence = "LOW"

    if mid1 is not None and mid2 is not None:
        mid_avg = (mid1 + mid2) / 2
    elif mid1 is not None:
        mid_avg = mid1
    elif mid2 is not None:
        mid_avg = mid2
    else:
        mid_avg = None

    external_candidates = []
    for row in mark_rows:
        sem_val = _to_float(row.semester)
        if sem_val is not None and sem_val > 0:
            external_candidates.append(sem_val)

        exam_key = _normalize_exam_name(row.exam)
        if exam_key == "semester":
            exam_mark = _to_float(row.marks)
            if exam_mark is not None and exam_mark > 0:
                external_candidates.append(exam_mark)

    external_estimate = round(sum(external_candidates) / len(external_candidates), 2) if external_candidates else 50.0

    cgpa = None
    internal = None
    scaled_mid = None
    cgpa_prediction_note = None
    required_mid2_targets = []
    simulation = []

    def _clamp(value, min_value, max_value):
        return max(min_value, min(max_value, value))

    if mid_avg is None:
        cgpa_prediction_note = "Not enough data"
    else:
        scaled_mid = round((mid_avg / 30) * 20, 2)
        internal = round(scaled_assignment + scaled_mid, 2)

        final_score = _clamp(round(internal + external_estimate, 2), 0, 100)
        cgpa = round(final_score / 10, 2)

        targets = [6.5, 7.5, 8.5]
        for target in targets:
            target_score = target * 10
            required_internal = target_score - external_estimate
            required_scaled_mid = required_internal - scaled_assignment
            required_mid_avg = (required_scaled_mid / 20) * 30

            if mid1 is not None:
                required_mid2_raw = (2 * required_mid_avg) - mid1
            else:
                required_mid2_raw = required_mid_avg

            if required_mid2_raw <= 30:
                status = "ACHIEVABLE"
            elif required_mid2_raw <= 35:
                status = "STRETCH"
            else:
                status = "IMPOSSIBLE"

            required_mid2 = round(_clamp(required_mid2_raw, 0, 30), 2)
            required_mid2_targets.append(
                {
                    "target": target,
                    "required_mid2_raw": round(required_mid2_raw, 2),
                    "required_mid2": required_mid2,
                    "status": status,
                }
            )

        for simulated_mid2 in [10, 15, 20, 25, 30]:
            sim_mid_avg = (mid1 + simulated_mid2) / 2 if mid1 is not None else simulated_mid2
            sim_scaled_mid = (sim_mid_avg / 30) * 20
            sim_internal = scaled_assignment + sim_scaled_mid
            sim_final_score = _clamp(sim_internal + external_estimate, 0, 100)
            sim_cgpa = round(sim_final_score / 10, 2)

            simulation.append(
                {
                    "mid2": simulated_mid2,
                    "cgpa": sim_cgpa,
                }
            )

    def _compute_placement_status(cgpa_value, attendance_value):
        if cgpa_value is None or attendance_value is None:
            return "INSUFFICIENT DATA"
        if cgpa_value >= 8 and attendance_value >= 80:
            return "READY"
        if cgpa_value >= 7 and attendance_value >= 70:
            return "BORDERLINE"
        return "NOT READY"

    placement_readiness = _compute_placement_status(cgpa, attendance_percent)

    # -----------------------------
    # Subject-Level Intelligence
    # -----------------------------
    subject_ids = set()
    for row in mark_rows:
        if row.subject_id is not None:
            subject_ids.add(row.subject_id)
    for sid in subject_attendance_map.keys():
        subject_ids.add(sid)

    subject_name_map = {}
    if subject_ids:
        subjects = db.query(Subject).filter(Subject.subject_id.in_(list(subject_ids))).all()
        subject_name_map = {sub.subject_id: sub.subject_name for sub in subjects}

    marks_by_subject = {}
    for row in mark_rows:
        subject_id = row.subject_id
        if subject_id is None:
            continue
        marks_by_subject.setdefault(subject_id, []).append(row)

    def _risk_from_cgpa(subject_cgpa):
        if subject_cgpa is None:
            return "INSUFFICIENT DATA"
        if subject_cgpa < 6:
            return "HIGH"
        if subject_cgpa < 7.5:
            return "MEDIUM"
        return "LOW"

    def _required_mid2_status(required_mid2_raw):
        if required_mid2_raw <= 30:
            return "ACHIEVABLE"
        if required_mid2_raw <= 35:
            return "STRETCH"
        return "IMPOSSIBLE"

    subject_items = []
    for subject_id in sorted(subject_ids):
        subject_rows = marks_by_subject.get(subject_id, [])
        subject_name = subject_name_map.get(subject_id, f"Subject {subject_id}")

        subject_mid1_values = []
        subject_mid2_values = []
        subject_assignments = []

        for row in subject_rows:
            exam_key = _normalize_exam_name(row.exam)
            marks_value = _to_float(row.marks)

            if exam_key in ["mid1", "mid01"] and marks_value is not None:
                subject_mid1_values.append(marks_value)
            elif exam_key in ["mid2", "mid02"] and marks_value is not None:
                subject_mid2_values.append(marks_value)

            if exam_key.startswith("assignment") and marks_value is not None:
                max_value = _extract_assignment_max(row.extra_data, row.exam, exam_key)
                if max_value is None:
                    cohort_max_raw = (
                        db.query(func.max(Mark.marks))
                        .filter(
                            Mark.subject_id == row.subject_id,
                            Mark.exam == row.exam,
                            Mark.marks.isnot(None),
                        )
                        .scalar()
                    )
                    cohort_max = _to_float(cohort_max_raw)
                    if cohort_max is not None and cohort_max > 0:
                        max_value = cohort_max

                if max_value is None or max_value <= 0:
                    max_value = marks_value

                subject_assignments.append({"score": marks_value, "max": max_value})

        if not subject_mid1_values:
            for row in subject_rows:
                val = _to_float(row.mid1)
                if val is not None and val > 0:
                    subject_mid1_values.append(val)

        if not subject_mid2_values:
            for row in subject_rows:
                val = _to_float(row.mid2)
                if val is not None and val > 0:
                    subject_mid2_values.append(val)

        if not subject_assignments:
            for row in subject_rows:
                if not isinstance(row.extra_data, dict):
                    continue

                lowered = {str(k).strip().lower(): v for k, v in row.extra_data.items()}
                for key, raw_value in lowered.items():
                    if "assignment" not in key:
                        continue
                    if "max" in key or "total" in key or "scaled" in key:
                        continue

                    score_value = _to_float(raw_value)
                    if score_value is None:
                        continue

                    max_value = _to_float(lowered.get(f"{key}_max"))
                    if max_value is None or max_value <= 0:
                        continue

                    subject_assignments.append({"score": score_value, "max": max_value})

        subject_mid1 = round(sum(subject_mid1_values) / len(subject_mid1_values), 2) if subject_mid1_values else None
        subject_mid2 = round(sum(subject_mid2_values) / len(subject_mid2_values), 2) if subject_mid2_values else None

        subject_assignment_total = round(sum(item["score"] for item in subject_assignments), 2)
        subject_assignment_max_total = round(sum(item["max"] for item in subject_assignments), 2)
        subject_scaled_assignment = round((subject_assignment_total / subject_assignment_max_total) * 10, 2) if subject_assignment_max_total > 0 else 0.0

        subject_attendance_obj = subject_attendance_map.get(subject_id, {"present": 0, "total": 0})
        subject_attendance = round((subject_attendance_obj["present"] / subject_attendance_obj["total"]) * 100, 2) if subject_attendance_obj["total"] > 0 else None

        if subject_mid1 is not None and subject_mid2 is not None:
            subject_mid_avg = (subject_mid1 + subject_mid2) / 2
        elif subject_mid1 is not None:
            subject_mid_avg = subject_mid1
        elif subject_mid2 is not None:
            subject_mid_avg = subject_mid2
        else:
            subject_mid_avg = None

        subject_external_candidates = []
        for row in subject_rows:
            sem_val = _to_float(row.semester)
            if sem_val is not None and sem_val > 0:
                subject_external_candidates.append(sem_val)

            exam_key = _normalize_exam_name(row.exam)
            if exam_key == "semester":
                exam_mark = _to_float(row.marks)
                if exam_mark is not None and exam_mark > 0:
                    subject_external_candidates.append(exam_mark)

        subject_external_estimate = (
            round(sum(subject_external_candidates) / len(subject_external_candidates), 2)
            if subject_external_candidates
            else 50.0
        )
        subject_internal = None
        subject_scaled_mid = None
        subject_cgpa = None
        required_mid2 = None
        required_mid2_raw = None
        required_mid2_for_8 = None
        required_mid2_for_8_raw = None
        status = "INSUFFICIENT DATA"

        if subject_mid_avg is not None:
            subject_scaled_mid = round((subject_mid_avg / 30) * 20, 2)
            subject_internal = round(subject_scaled_assignment + subject_scaled_mid, 2)
            subject_final_score = _clamp(subject_internal + subject_external_estimate, 0, 100)
            subject_cgpa = round(subject_final_score / 10, 2)

            target_cgpa = 7.5
            target_score = target_cgpa * 10
            required_internal = target_score - subject_external_estimate
            required_scaled_mid = required_internal - subject_scaled_assignment
            required_mid_avg = (required_scaled_mid / 20) * 30

            if subject_mid1 is not None:
                required_mid2_raw = (2 * required_mid_avg) - subject_mid1
            else:
                required_mid2_raw = required_mid_avg

            status = _required_mid2_status(required_mid2_raw)
            required_mid2_raw = round(required_mid2_raw, 2)
            required_mid2 = round(_clamp(required_mid2_raw, 0, 30), 2)

            target_cgpa_8 = 8.0
            target_score_8 = target_cgpa_8 * 10
            required_internal_8 = target_score_8 - subject_external_estimate
            required_scaled_mid_8 = required_internal_8 - subject_scaled_assignment
            required_mid_avg_8 = (required_scaled_mid_8 / 20) * 30

            if subject_mid1 is not None:
                required_mid2_raw_8 = (2 * required_mid_avg_8) - subject_mid1
            else:
                required_mid2_raw_8 = required_mid_avg_8

            required_mid2_for_8_raw = round(required_mid2_raw_8, 2)
            required_mid2_for_8 = round(_clamp(required_mid2_raw_8, 0, 30), 2)

        reasons = []
        if subject_attendance is not None and subject_attendance < 75:
            reasons.append("Low attendance")
        if subject_mid1 is not None and subject_mid1 < 15:
            reasons.append("Low marks")
        if subject_mid2 is not None and subject_mid2 < 15:
            reasons.append("Low marks")
        if subject_mid1 is not None and subject_mid2 is not None and subject_mid2 < subject_mid1:
            reasons.append("Declining mid performance")

        if not reasons and subject_cgpa is None:
            reasons.append("Insufficient data")
        elif not reasons:
            reasons.append("Stable performance")

        reason_text = ", ".join(dict.fromkeys(reasons))
        risk = _risk_from_cgpa(subject_cgpa)

        impact = None
        if subject_cgpa is not None:
            target_for_impact = 8.0
            gap = round(max(0, target_for_impact - subject_cgpa), 2)
            if gap > 0:
                impact = f"Improving this subject by +1 CGPA can increase overall CGPA significantly (gap to 8.0: {gap})"
            else:
                impact = "This subject already supports a strong overall CGPA"

        cgpa_gap_to_8 = round(max(0, 8.0 - subject_cgpa), 2) if subject_cgpa is not None else None

        if subject_cgpa is None:
            priority_score = 0
        else:
            attendance_gap = max(0, 80 - subject_attendance) if subject_attendance is not None else 0
            mark_gap = 0
            if subject_mid1 is not None:
                mark_gap += max(0, 15 - subject_mid1)
            if subject_mid2 is not None:
                mark_gap += max(0, 15 - subject_mid2)

            priority_score = round(
                _clamp((cgpa_gap_to_8 * 35) + (attendance_gap * 0.6) + (mark_gap * 1.5), 0, 100),
                2,
            )

        priority_breakdown = {
            "cgpa_gap_component": round((cgpa_gap_to_8 or 0) * 35, 2) if cgpa_gap_to_8 is not None else 0,
            "attendance_gap_component": round((max(0, 80 - subject_attendance) if subject_attendance is not None else 0) * 0.6, 2),
            "marks_gap_component": round(
                (
                    (max(0, 15 - subject_mid1) if subject_mid1 is not None else 0)
                    + (max(0, 15 - subject_mid2) if subject_mid2 is not None else 0)
                )
                * 1.5,
                2,
            ),
            "formula": "(cgpa_gap*35) + (attendance_gap*0.6) + (marks_gap*1.5)",
        }

        time_allocation = "Time allocation will be normalized after scoring"

        if subject_cgpa is None:
            faculty_feedback = {
                "summary": "Not enough subject data",
                "action": "Collect more marks and attendance records before final feedback",
                "source": "SYSTEM_GENERATED",
            }
        elif risk == "HIGH":
            faculty_feedback = {
                "summary": "Subject is below safe academic level",
                "action": "Teacher should review basics, check attendance, and assign focused practice",
                "source": "SYSTEM_GENERATED",
            }
        elif risk == "MEDIUM":
            faculty_feedback = {
                "summary": "Subject is near the warning zone",
                "action": "Teacher should reinforce weak units and monitor the next assessment closely",
                "source": "SYSTEM_GENERATED",
            }
        else:
            faculty_feedback = {
                "summary": "Subject is performing well",
                "action": "Teacher should maintain current pace and give advanced practice",
                "source": "SYSTEM_GENERATED",
            }

        subject_items.append(
            {
                "name": subject_name,
                "subject_id": subject_id,
                "attendance": subject_attendance,
                "mid1": subject_mid1,
                "mid2": subject_mid2,
                "assignment_total": subject_assignment_total,
                "assignment_max_total": subject_assignment_max_total,
                "scaled_assignment": subject_scaled_assignment,
                "scaled_mid": subject_scaled_mid,
                "internal": subject_internal,
                "external_estimate": subject_external_estimate,
                "cgpa": subject_cgpa,
                "risk": risk,
                "priority_score": priority_score,
                "cgpa_gap_to_8": cgpa_gap_to_8,
                "time_allocation": time_allocation,
                "faculty_feedback": faculty_feedback,
                "required_mid2": required_mid2,
                "required_mid2_raw": required_mid2_raw,
                "required_mid2_for_8": required_mid2_for_8,
                "required_mid2_for_8_raw": required_mid2_for_8_raw,
                "status": status,
                "reason": reason_text,
                "impact": impact,
                "priority_breakdown": priority_breakdown,
            }
        )

    # Normalize time allocation percentages so recommendations are bounded and auditable.
    if subject_items:
        positive_weights = [max(0.0, float(item.get("priority_score") or 0.0)) for item in subject_items]
        total_weight = sum(positive_weights)

        if total_weight <= 0:
            base_percent = round(100 / len(subject_items))
            allocated = [base_percent for _ in subject_items]
        else:
            allocated = [int(round((weight / total_weight) * 100)) for weight in positive_weights]

        drift = 100 - sum(allocated)
        if allocated:
            allocated[0] += drift

        for index, item in enumerate(subject_items):
            allocation_percent = max(0, allocated[index])
            item["time_allocation_percent"] = allocation_percent
            item["time_allocation"] = f"Spend {allocation_percent}% time on {item['name']}"

    subjects_with_cgpa = [item for item in subject_items if item.get("cgpa") is not None]
    subjects_with_cgpa.sort(key=lambda item: item["cgpa"])  # type: ignore

    weakest_subjects = [
        {
            "name": item["name"],
            "cgpa": item["cgpa"],
            "risk": item["risk"],
            "reason": item["reason"],
            "required_mid2": item["required_mid2"],
            "status": item["status"],
            "impact": item["impact"],
            "priority_score": item["priority_score"],
            "time_allocation": item["time_allocation"],
            "faculty_feedback": item["faculty_feedback"],
        }
        for item in subjects_with_cgpa[:2]
    ]

    strongest_subject = None
    if subjects_with_cgpa:
        strongest_item = subjects_with_cgpa[-1]
        strongest_subject = {
            "name": strongest_item["name"],
            "cgpa": strongest_item["cgpa"],
            "risk": strongest_item["risk"],
            "reason": strongest_item["reason"],
            "priority_score": strongest_item["priority_score"],
        }

    priority_subject = None
    if subject_items:
        priority_subject_item = max(subject_items, key=lambda item: item.get("priority_score") or 0)
        priority_subject = {
            "name": priority_subject_item["name"],
            "priority_score": priority_subject_item["priority_score"],
            "cgpa": priority_subject_item["cgpa"],
            "reason": priority_subject_item["reason"],
            "time_allocation": priority_subject_item["time_allocation"],
            "time_allocation_percent": priority_subject_item.get("time_allocation_percent", 0),
            "priority_breakdown": priority_subject_item.get("priority_breakdown", {}),
        }

    subject_intelligence = {
        "subjects": subject_items,
        "weakest_subjects": weakest_subjects,
        "strongest_subject": strongest_subject,
        "priority_subject": priority_subject,
    }

    marks_trend = marks_trend_values[-7:]
    if len(marks_trend) < 2:
        marks_trend = []
        if mid1 is not None:
            marks_trend.append(mid1)
        if mid2 is not None:
            marks_trend.append(mid2)

    engine_risk = str(risk_data.get("overall_risk") or "NO_DATA").upper()
    risk_level = "INSUFFICIENT DATA" if engine_risk == "NO_DATA" else engine_risk

    if risk_data.get("attendance_status") == "LOW":
        primary_issue = "Low attendance"
    elif risk_data.get("marks_status") == "LOW":
        primary_issue = "Low marks"
    elif risk_data.get("cgpa_status") == "LOW":
        primary_issue = "Low CGPA"
    else:
        primary_issue = "No major issues"

    attendance_decreasing = False
    if len(attendance_trend) >= 2:
        first_span = attendance_trend[:2]
        last_span = attendance_trend[-2:]
        first_avg = sum(first_span) / len(first_span)
        last_avg = sum(last_span) / len(last_span)
        attendance_decreasing = last_avg < first_avg

    marks_decreasing = (mid1 is not None and mid2 is not None and mid2 < mid1)

    attendance_previous = round(sum(previous_attendance_window) / len(previous_attendance_window), 2) if previous_attendance_window else None
    attendance_change = round(attendance_percent - attendance_previous, 2) if attendance_percent is not None and attendance_previous is not None else None
    marks_progress = None
    if mid1 is not None and mid2 is not None:
        if mid2 > mid1:
            marks_progress = "improved"
        elif mid2 < mid1:
            marks_progress = "declined"
        else:
            marks_progress = "stable"

    focus_now = "Attendance" if attendance_percent is not None and attendance_percent < 75 else "Marks" if ((mid1 is not None and mid1 < 15) or (mid2 is not None and mid2 < 15) or marks_decreasing) else "Consistency"
    focus_reason = (
        "Improving attendance will have the highest impact on CGPA"
        if focus_now == "Attendance"
        else "Improving marks will have the highest impact on final performance"
        if focus_now == "Marks"
        else "Consistency will protect both CGPA and placement readiness"
    )

    weekly_goal = []
    if attendance_percent is not None and attendance_percent < 75:
        weekly_goal.append("Attend all classes this week")
    if (mid1 is not None and mid1 < 15) or (mid2 is not None and mid2 < 15):
        weekly_goal.append("Revise weak subjects")
    if marks_decreasing:
        weekly_goal.append("Practice previous exam questions")
    if not weekly_goal:
        weekly_goal = ["Keep attendance above 80% this week", "Solve one revision set before the next class"]
    weekly_goal = weekly_goal[:3]

    consequences = []
    if attendance_percent is not None and attendance_percent < 75:
        consequences.append("May lose internal marks")
    if (mid1 is not None and mid1 < 15) or (mid2 is not None and mid2 < 15):
        consequences.append("Risk of low CGPA")
    if marks_decreasing:
        consequences.append("Performance may worsen in finals")
    if not consequences:
        consequences.append("Readiness margin will shrink if attendance or marks drop")

    if attendance_decreasing:
        prediction = "Attendance likely to drop below 75%"
    elif marks_decreasing:
        prediction = "Performance may decline in next exams"
    else:
        prediction = "Performance is stable"

    actions = list(risk_data.get("actions") or [])
    if not actions:
        actions = ["Maintain current attendance and marks consistency"]

    average_mid = None
    if mid1 is not None and mid2 is not None:
        average_mid = (mid1 + mid2) / 2
    elif mid1 is not None:
        average_mid = mid1
    elif mid2 is not None:
        average_mid = mid2

    readiness = round((attendance_percent + average_mid) / 2, 2) if attendance_percent is not None and average_mid is not None else None
    aptitude = round((mid1 / 30) * 100, 2) if mid1 is not None else None

    # Enhanced consistency score with interpretation
    consistency = None
    consistency_interpretation = "Insufficient data"
    if len(attendance_trend) >= 2:
        trend_mean = sum(attendance_trend) / len(attendance_trend)
        variance = sum((value - trend_mean) ** 2 for value in attendance_trend) / len(attendance_trend)
        std_dev = variance ** 0.5
        consistency = round(max(0, min(100, 100 - (std_dev * 5))), 2)  # Scale std dev impact
        
        if consistency >= 80:
            consistency_interpretation = "Very consistent attendance pattern"
        elif consistency >= 60:
            consistency_interpretation = "Moderately consistent"
        else:
            consistency_interpretation = "Inconsistent - shows irregular attendance"

    # WHAT-IF SIMULATION ENGINE
    what_if = {
        "attendance_improvement": "",
        "marks_improvement": "",
        "combined_impact": ""
    }
    
    if attendance_percent is not None:
        # Simulate attending next 5 classes (assume 5 more classes total)
        simulated_new_attendance = round(((present_attendance + 5) / (total_attendance + 5)) * 100, 2)
        attendance_boost = simulated_new_attendance - attendance_percent
        
        if attendance_boost > 0:
            what_if["attendance_improvement"] = f"If you attend next 5 classes: {attendance_percent}% → {simulated_new_attendance}% (+{attendance_boost}%)"
        else:
            what_if["attendance_improvement"] = f"Attendance is already high at {attendance_percent}%"
    
    if mid1 is not None:
        # Simulate marks improvement by +5 points
        simulated_mid2_improved = min(30, (mid2 or mid1) + 5)
        marks_boost = simulated_mid2_improved - (mid2 or mid1)
        original_avg = (mid1 + (mid2 or mid1)) / 2
        simulated_avg = (mid1 + simulated_mid2_improved) / 2
        
        if marks_boost > 0:
            what_if["marks_improvement"] = f"If marks improve by +5: Average {original_avg:.1f} → {simulated_avg:.1f}"
        else:
            what_if["marks_improvement"] = f"Marks are already strong at {mid1} (Mid1)"
    
    # Combined impact
    if what_if["attendance_improvement"] and what_if["marks_improvement"]:
        what_if["combined_impact"] = "Both attendance and marks improvement will compound your readiness gain"

    # EARLY WARNING SYSTEM
    warnings = []
    
    if attendance_percent is not None and attendance_percent < 75 and total_attendance > 0:
        classes_allowed_to_miss = max(0, int((present_attendance / 0.75) - total_attendance))
        if classes_allowed_to_miss <= 0:
            warnings.append("Critical: Missing even 1 more class can keep you below the 75% safe level")
        elif classes_allowed_to_miss == 1:
            warnings.append("Critical: You can miss only 1 more class before dropping below 75%")
        else:
            warnings.append(
                f"Critical: You can miss only {classes_allowed_to_miss} more classes before dropping below 75%"
            )
    
    if attendance_percent is not None and attendance_percent < 65:
        warnings.append(f"URGENT: Attendance at {attendance_percent}% is critically low")
    
    if mid1 is not None and mid2 is not None and mid2 < mid1 - 3:
        warnings.append(f"Performance declining: Mid2 ({mid2}) vs Mid1 ({mid1}) - focus on weak subjects")
    
    if (mid1 is not None and mid1 < 12) or (mid2 is not None and mid2 < 12):
        warnings.append("Marks are below safe threshold - high final exam risk")
    
    if len(attendance_trend) >= 3:
        recent_avg = sum(attendance_trend[-3:]) / 3
        older_avg = sum(attendance_trend[:3]) / 3 if len(attendance_trend) >= 3 else attendance_percent
        if recent_avg < older_avg - 5:
            warnings.append("Attendance trend is declining sharply")

    # DAILY TASK GENERATION (strictly from real data only)
    now_dt = datetime.utcnow()
    today_date = now_dt.date()
    daily_tasks = []

    if attendance_percent is not None and attendance_percent < 75:
        daily_tasks.append({
            "id": "attendance_safe_threshold",
            "title": "Attend all classes this week",
            "task": "Attend all classes this week",
            "reason": f"Attendance is {attendance_percent}% and below the 75% safe threshold",
            "priority": "HIGH",
            "status": False,
        })

    assignment_candidates = (
        db.query(Assignment)
        .filter(
            Assignment.year == student_profile.year,
            Assignment.section == student_profile.section,
            Assignment.is_active == True,
        )
        .order_by(Assignment.due_date.asc())
        .all()
    )

    assignment_ids = [item.id for item in assignment_candidates]
    submission_map = {}
    if assignment_ids:
        submission_rows = (
            db.query(AssignmentSubmission)
            .filter(
                AssignmentSubmission.student_id == student_id,
                AssignmentSubmission.assignment_id.in_(assignment_ids),
            )
            .all()
        )
        submission_map = {row.assignment_id: row for row in submission_rows}

    pending_assignment = None
    for assignment in assignment_candidates:
        submission = submission_map.get(assignment.id)
        is_submitted = bool(submission and submission.is_submitted)
        if is_submitted:
            continue
        if assignment.due_date and assignment.due_date < now_dt:
            continue
        pending_assignment = assignment
        break

    if pending_assignment is not None:
        due_text = pending_assignment.due_date.date().isoformat() if pending_assignment.due_date else "upcoming"
        daily_tasks.append({
            "id": "assignment_pending",
            "title": "Complete pending assignment",
            "task": "Complete pending assignment",
            "reason": f"Pending assignment '{pending_assignment.title}' due on {due_text}",
            "priority": "HIGH",
            "status": False,
        })

    if (mid1 is not None and mid1 < 15) or (mid2 is not None and mid2 < 15):
        daily_tasks.append({
            "id": "marks_low",
            "title": "Prepare for next internal exam",
            "task": "Prepare for next internal exam",
            "reason": "Mid marks are below the safe threshold",
            "priority": "HIGH",
            "status": False,
        })

    # Keep output compact and deterministic.
    priority_rank = {"HIGH": 0, "MEDIUM": 1, "LOW": 2}
    daily_tasks = sorted(daily_tasks, key=lambda item: priority_rank.get(item["priority"], 9))[:3]

    # BEHAVIOR PATTERN DETECTION
    patterns = []
    
    if len(attendance_trend) >= 5:
        # Check for mid-week drop
        first_half = attendance_trend[:len(attendance_trend)//2]
        second_half = attendance_trend[len(attendance_trend)//2:]
        
        if first_half and second_half:
            first_avg = sum(first_half) / len(first_half)
            second_avg = sum(second_half) / len(second_half)
            
            if second_avg < first_avg - 10:
                patterns.append("📉 Attendance drops later in the week - plan weekend revisions")
            elif second_avg > first_avg + 10:
                patterns.append("📈 Attendance improves later in the week - maintain momentum")
    
    if marks_decreasing and mid1 is not None and mid2 is not None:
        decline_rate = mid1 - mid2
        if decline_rate > 5:
            patterns.append(f"📉 Performance declining rapidly ({decline_rate} points) - needs immediate focus")
        else:
            patterns.append(f"📉 Performance declining slightly - adjust study approach")
    
    if marks_trend and len(marks_trend) >= 3:
        increasing = marks_trend[-1] > marks_trend[0]
        if increasing:
            patterns.append("📈 Marks showing upward trend - strategy is working")
        else:
            patterns.append("📉 Marks trend declining - current study method needs revision")
    
    if attendance_percent and mid1 and mid2:
        if attendance_percent >= 85 and mid2 >= mid1:
            patterns.append("✅ High attendance + improving marks = strong trajectory")
        elif attendance_percent < 70 and mid2 < mid1:
            patterns.append("⚠️ Both attendance and marks declining - urgent intervention needed")

    # EARLY WARNING SIGNALS (critical flags)
    early_warning_signals = []
    
    if attendance_percent is not None and attendance_percent < 60:
        early_warning_signals.append({
            "level": "CRITICAL",
            "signal": "Attendance critically low",
            "action": "Meet academic advisor immediately"
        })
    
    if (mid1 is not None and mid1 < 10) or (mid2 is not None and mid2 < 10):
        early_warning_signals.append({
            "level": "CRITICAL",
            "signal": "Marks are dangerously low",
            "action": "Seek tutor assistance or peer study group"
        })
    
    if risk_level == "HIGH" and len(early_warning_signals) == 0:
        early_warning_signals.append({
            "level": "HIGH",
            "signal": "Overall risk level is HIGH",
            "action": "Schedule meeting with mentor"
        })

    if placement_readiness == "INSUFFICIENT DATA":
        placement_status = "INSUFFICIENT DATA"
        placement_reasons = ["Insufficient data to analyze"]
        placement_gaps = []
        placement_weekly = []
        placement_monthly = []
        placement_timeline = "Insufficient data to analyze"
        placement_risk_if_ignored = []
    else:
        placement_status = placement_readiness

        placement_reasons = []
        if attendance_percent < 75:
            placement_reasons.append("Attendance below required level")
        if cgpa is not None and cgpa < 7:
            placement_reasons.append("Predicted CGPA is below safe placement range")
        if mid1 is not None and mid2 is not None and mid2 < mid1:
            placement_reasons.append("Declining performance trend")
        if not placement_reasons:
            placement_reasons.append("CGPA and attendance are within placement thresholds")

        placement_gaps = []
        if attendance_percent < 80:
            placement_gaps.append({"metric": "Attendance", "current": attendance_percent, "target": 80})
        if cgpa is not None and cgpa < 8:
            placement_gaps.append({"metric": "CGPA", "current": cgpa, "target": 8.0})

        placement_weekly = []
        if attendance_percent < 75:
            placement_weekly.append("Attend all classes this week")
        if cgpa is not None and cgpa < 8:
            placement_weekly.append("Revise weak subjects")
        if mid1 is not None and mid2 is not None and mid2 < mid1:
            placement_weekly.append("Focus on understanding concepts")
        if not placement_weekly:
            placement_weekly.append("Maintain current attendance and revision rhythm")

        placement_monthly = ["Improve marks by +5", "Maintain attendance above 80%"]

        if placement_status == "NOT READY":
            placement_timeline = "4–6 weeks to reach placement readiness"
        elif placement_status == "BORDERLINE":
            placement_timeline = "2–3 weeks to become fully ready"
        else:
            placement_timeline = "Currently ready"

        placement_risk_if_ignored = []
        if attendance_percent < 75:
            placement_risk_if_ignored.append("May lose internal marks")
        if cgpa is not None and cgpa < 7:
            placement_risk_if_ignored.append("Low CGPA risk")
        if not placement_risk_if_ignored:
            placement_risk_if_ignored.append("Could lose readiness margin if attendance or marks drop")

    return {
        "student_id": student_id,
        "has_valid_data": True,
        "no_data_message": None,
        "attendance": attendance_percent,
        "attendance_trend": attendance_trend,
        "mid1": mid1,
        "mid2": mid2,
        "external_estimate": external_estimate,
        "cgpa": cgpa,
        "internal": internal,
        "scaled_mid": scaled_mid,
        "scaled_assignment": scaled_assignment,
        "assignment_total": assignment_total,
        "assignment_max_total": assignment_max_total,
        "assignment_confidence": assignment_confidence,
        "excluded_assignments_without_max": excluded_assignments_without_max,
        "required_mid2_targets": required_mid2_targets,
        "simulation": simulation,
        "placement_readiness": placement_readiness,
        "cgpa_prediction_note": cgpa_prediction_note,
        "subject_intelligence": subject_intelligence,
        "marks_trend": marks_trend,
        "risk_level": risk_level,
        "risk_reasons": list(risk_data.get("reasons") or []),
        "primary_issue": primary_issue,
        "prediction": prediction,
        "weekly_goal": weekly_goal,
        "progress_this_week": {
            "attendance_current": attendance_percent,
            "attendance_previous": attendance_previous,
            "attendance_change": attendance_change,
            "marks_status": marks_progress,
        },
        "focus_now": {
            "metric": focus_now,
            "reason": focus_reason,
        },
        "consequences": consequences,
        "actions": actions,
        "what_if": what_if,
        "warnings": warnings,
        "daily_tasks": daily_tasks,
        "patterns": patterns,
        "early_warning_signals": early_warning_signals,
        "placement_analysis": {
            "status": placement_status,
            "reasons": placement_reasons,
            "gaps": placement_gaps,
            "roadmap": {
                "weekly": placement_weekly,
                "monthly": placement_monthly,
            },
            "timeline": placement_timeline,
            "risk_if_ignored": placement_risk_if_ignored,
        },
        "placement": {
            "readiness": readiness,
            "aptitude": aptitude,
            "consistency": consistency,
            "consistency_interpretation": consistency_interpretation,
        }
    }


# =========================
# MARK ALERT AS READ
# =========================
@app.patch("/alerts/{alert_id}/read")
@app.put("/alerts/{alert_id}/read")
def mark_read(alert_id: int, current_user=Depends(get_current_user), db: Session = Depends(get_db)):
    recipient = db.query(AlertRecipient).filter(
        AlertRecipient.alert_id == alert_id,
        AlertRecipient.user_id == current_user["user_id"]
    ).first()

    if not recipient:
        raise HTTPException(status_code=404, detail="Not found")

    recipient.is_read = True
    db.commit()

    return {"message": "Marked as read"}


# =========================
# ADMIN – GET ALL SUBJECTS
# =========================
@app.get("/admin/subjects")
def get_all_subjects(
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Admin only")

    subjects = db.query(Subject).all()

    return [
        {
            "subject_id": s.subject_id,
            "subject_name": s.subject_name,
            "semester": s.semester,
            "department_id": s.department_id
        }
        for s in subjects
    ]


# =========================
# =========================
# ADMIN – SUBJECT PERFORMANCE
@app.get("/admin/subject-performance")
def get_subject_performance(
    department: Optional[str] = Query(None),
    year: Optional[str] = Query(None),
    semester: Optional[str] = Query(None),
    assessment_type: Optional[str] = Query(None),
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Admin only")

    if not department or department == "All" or department.strip() == "" or \
       not year or str(year) == "All" or str(year).strip() == "" or \
       not semester or str(semester) == "All" or str(semester).strip() == "" or \
       not assessment_type or assessment_type == "All" or assessment_type.strip() == "":
        return {
            "subjects": [],
            "kpis": None,
            "students": []
        }

    query = db.query(Subject)

    department_key = department.strip().upper()
    department_lookup = {
        "CSE": 11, "CSM": 12, "ECE": 14, "MECH": 15, "CIVIL": 1,
    }
    department_id = department_lookup.get(department_key)
    if department_id is not None:
        query = query.filter(Subject.department_id == department_id)

    yr = int(year)
    year_to_semesters = {1: [1, 2], 2: [3, 4], 3: [5, 6], 4: [7, 8]}
    semesters = year_to_semesters.get(yr)
    if semesters:
        query = query.filter(Subject.semester.in_(semesters))

    query = query.filter(Subject.semester == int(semester))

    pass_threshold = 40
    mapping = {
        "Mid 1": "Mid-1",
        "Mid 2": "Mid-2",
        "Semester": "Semester"
    }
    db_exam = mapping.get(assessment_type, assessment_type)

    if db_exam in ["Mid-1", "Mid-2"]:
        pass_threshold = 15
    elif db_exam == "Semester":
        pass_threshold = 40

    filtered_subjects = query.subquery()

    base_marks_query = db.query(
        Mark.student_id,
        Mark.subject_id,
        Mark.marks,
        Mark.exam
    )

    if db_exam:
        base_marks_query = base_marks_query.filter(Mark.exam == db_exam)
            
    marks_subq = base_marks_query.subquery()

    subject_query = (
        db.query(
            filtered_subjects.c.subject_name.label("subject_name"),
            func.avg(marks_subq.c.marks).label("avg_marks"),
            func.count(func.distinct(marks_subq.c.student_id)).label("total_students"),
            func.sum(case((marks_subq.c.marks < pass_threshold, 1), else_=0)).label("fail_count"),
            func.sum(case((marks_subq.c.marks >= pass_threshold, 1), else_=0)).label("pass_count")
        )
        .outerjoin(marks_subq, marks_subq.c.subject_id == filtered_subjects.c.subject_id)
        .group_by(filtered_subjects.c.subject_name)
        .order_by(filtered_subjects.c.subject_name)
    )

    subjects_data = subject_query.all()

    final_subjects = []
    total_global_students = 0
    total_global_passed = 0

    for row in subjects_data:
        total_st = int(row.total_students or 0)
        
        # We only want to include subjects that actually have data context
        if total_st == 0:
            continue
            
        pass_ct = int(row.pass_count or 0)
        pass_rt = round((pass_ct / total_st) * 100, 1)

        total_global_students += total_st
        total_global_passed += pass_ct

        final_subjects.append({
            "subject_name": row.subject_name,
            "total_students": total_st,
            "avg_marks": round(float(row.avg_marks or 0), 1),
            "failure_count": int(row.fail_count or 0),
            "pass_rate": pass_rt
        })

    if total_global_students == 0:
        return {
            "subjects": [],
            "kpis": None,
            "students": []
        }

    # Calculate actual KPIs
    overall_pass_rate = round((total_global_passed / total_global_students) * 100, 1)

    avg_marks_result = db.query(func.avg(marks_subq.c.marks)).join(
        filtered_subjects, marks_subq.c.subject_id == filtered_subjects.c.subject_id
    ).scalar()
    avg_marks = round(float(avg_marks_result), 1) if avg_marks_result else None

    at_risk_count = db.query(func.count(func.distinct(marks_subq.c.student_id))).join(
        filtered_subjects, marks_subq.c.subject_id == filtered_subjects.c.subject_id
    ).filter(marks_subq.c.marks < pass_threshold).scalar()

    kpis = {
        "avgMarks": avg_marks,
        "atRisk": at_risk_count or 0,
        "passRate": overall_pass_rate
    }

    # Fetch Top At-Risk Students specifically
    at_risk_query = (
        db.query(
            User.name.label("student_name"),
            filtered_subjects.c.subject_name.label("subject"),
            func.avg(marks_subq.c.marks).label("marks")
        )
        .join(filtered_subjects, marks_subq.c.subject_id == filtered_subjects.c.subject_id)
        .join(Student, marks_subq.c.student_id == Student.student_id)
        .join(User, Student.student_id == User.user_id)
        .filter(marks_subq.c.marks < pass_threshold)
        .group_by(User.name, filtered_subjects.c.subject_name)
        .order_by(func.avg(marks_subq.c.marks))
        .limit(5)
    )
    at_risk_results = at_risk_query.all()
    at_risk_students = [
        {"student_name": r.student_name, "subject": r.subject, "marks": round(r.marks, 1) if r.marks else 0}
        for r in at_risk_results
    ]

    return {
        "subjects": sorted(final_subjects, key=lambda x: (-x["failure_count"], x["avg_marks"])),
        "kpis": kpis,
        "students": at_risk_students
    }

# =========================
# ADMIN – CREATE SUBJECT
@app.post("/admin/subjects")
def create_subject(
    subject: SubjectCreate,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Admin only")

    new_subject = Subject(
        subject_code=subject.subject_code,
        subject_name=subject.subject_name,
        semester=subject.semester,
        credits=subject.credits,
        department_id=subject.department_id
    )

    db.add(new_subject)
    db.commit()
    db.refresh(new_subject)

    return {"message": "Subject created successfully"}


# =========================
# ADMIN – DELETE SUBJECT
# =========================
@app.delete("/admin/subjects/{subject_id}")
def delete_subject(
    subject_id: int,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Admin only")

    subject = db.query(Subject).filter(
        Subject.subject_id == subject_id
    ).first()

    if not subject:
        raise HTTPException(status_code=404, detail="Not found")

    db.delete(subject)
    db.commit()

    return {"message": "Deleted successfully"}



# =========================
# ADMIN – ASSIGN SUBJECT TO FACULTY
# =========================
@app.post("/admin/assign-subject")
def assign_subject(
    data: AssignSubjectRequest,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Admin only")

    existing = db.query(FacultySubject).filter(
        FacultySubject.faculty_id == data.faculty_id,
        FacultySubject.subject_id == data.subject_id,
        FacultySubject.year == data.year,
        FacultySubject.section == data.section,
        FacultySubject.is_active == True
    ).first()

    if existing:
        raise HTTPException(status_code=400, detail="Already assigned")

    new_assignment = FacultySubject(
        faculty_id=data.faculty_id,
        subject_id=data.subject_id,
        year=data.year,
        section=data.section
    )

    db.add(new_assignment)
    db.commit()

    return {"message": "Subject assigned successfully"}


# =========================
# ADMIN – GET SUBJECTS ASSIGNED TO A FACULTY
# =========================
@app.get("/admin/faculty/{faculty_id}/subjects")
def get_faculty_subjects(
    faculty_id: int,
    db: Session = Depends(get_db)
):
    assignments = (
        db.query(FacultySubject, Subject)
        .join(Subject, FacultySubject.subject_id == Subject.subject_id)
        .filter(
            FacultySubject.faculty_id == faculty_id,
            FacultySubject.is_active == True
        )
        .all()
    )

    return [
    {
        "id": fs.id,
        "subject_name": s.subject_name,
        "semester": s.semester,   # ✅ from Subject table
        "year": fs.year,
        "section": fs.section,
        "assigned_at": fs.assigned_at
    }
    for fs, s in assignments
    ]




# =========================
# ADMIN – GET FACULTY-SUBJECT ASSIGNMENT HISTORY
# =========================
@app.get("/admin/faculty/{faculty_id}/subjects/history")
def get_faculty_subject_history(
    faculty_id: int,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Admin only")

    assignments = (
        db.query(FacultySubject, Subject)
        .join(Subject, FacultySubject.subject_id == Subject.subject_id)
        .filter(FacultySubject.faculty_id == faculty_id)
        .order_by(FacultySubject.assigned_at.desc())
        .all()
    )

    return [
        {
            "assignment_id": fs.id,
            "subject_name": subject.subject_name,
            "year": fs.year,
            "section": fs.section,
            "is_active": fs.is_active,
            "assigned_at": fs.assigned_at
        }
        for fs, subject in assignments
    ]


# =========================
# ADMIN – DELETE FACULTY-SUBJECT ASSIGNMENT
# =========================
@app.delete("/admin/remove-assignment/{assignment_id}")
def remove_assignment(
    assignment_id: int,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Admin only")

    assignment = db.query(FacultySubject).filter(
        FacultySubject.id == assignment_id
    ).first()

    if not assignment:
        raise HTTPException(status_code=404, detail="Not found")

    assignment.is_active = False
    db.commit()

    return {"message": "Assignment removed (soft delete)"}

# =========================
# ADMIN – UPDATE FACULTY-SUBJECT ASSIGNMENT
# =========================
@app.put("/admin/update-assignment/{assignment_id}")
def update_assignment(
    assignment_id: int,
    year: int,
    section: str,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Admin only")

    assignment = db.query(FacultySubject).filter(
        FacultySubject.id == assignment_id
    ).first()

    if not assignment:
        raise HTTPException(status_code=404, detail="Assignment not found")

    assignment.year = year
    assignment.section = section

    db.commit()

    return {"message": "Assignment updated successfully"}

# =========================
# FACULTY – DOWNLOAD WEEKLY PDF (RANK + DEFAULTER)
# =========================
@app.get("/faculty/attendance/weekly/{subject_id}/download")
def download_weekly_pdf(
    subject_id: int,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):

    if current_user["role"] != "faculty":
        raise HTTPException(status_code=403, detail="Faculty only")

    from datetime import date, timedelta

    today = date.today()
    start_week = today - timedelta(days=today.weekday())

    # 🔒 Validate assignment
    assignment = db.query(FacultySubject).filter(
        FacultySubject.faculty_id == current_user["user_id"],
        FacultySubject.subject_id == subject_id,
        FacultySubject.is_active == True
    ).first()

    if not assignment:
        raise HTTPException(status_code=403, detail="Not assigned to this subject")

    subject = db.query(Subject).filter(
        Subject.subject_id == subject_id
    ).first()

    unique_classes = db.query(Attendance.attendance_date).filter(
        Attendance.subject_id == subject_id,
        Attendance.attendance_date >= start_week,
        Attendance.attendance_date <= today
    ).distinct().count()

   
    # Get subject department
    department_id = subject.department_id

    students = (
        db.query(Student, User)
        .join(User, Student.student_id == User.user_id)
        .filter(
            Student.year == assignment.year,
            Student.section == assignment.section,
            User.department_id == department_id,
            User.is_deleted == False
        )
        .all()
    )

    # -----------------------------
    # Collect student performance
    # -----------------------------
    student_rows = []
    total_records = 0
    total_present = 0

    for student, user in students:

        records = db.query(Attendance).filter(
            Attendance.student_id == student.student_id,
            Attendance.subject_id == subject_id,
            Attendance.attendance_date >= start_week,
            Attendance.attendance_date <= today
        ).all()

        total = len(records)
        present = len([r for r in records if r.status])
        absent = total - present
        percent = round((present / total) * 100, 2) if total > 0 else 0  # type: ignore

        total_records += total
        total_present += present

        student_rows.append({
            "roll": student.roll_no,
            "name": user.name,
            "total": total,
            "present": present,
            "absent": absent,
            "percentage": percent
        })

    # -----------------------------
    # Sort by percentage (DESC)
    # -----------------------------
    student_rows.sort(key=lambda x: x["percentage"], reverse=True)

    # -----------------------------
    # PDF Setup
    # -----------------------------
    file_path = f"weekly_report_{subject_id}.pdf"
    doc = SimpleDocTemplate(file_path)
    elements = []
    styles = getSampleStyleSheet()

    elements.append(Paragraph("GVP-MAAA College", styles["Title"]))
    elements.append(Paragraph("Weekly Attendance Report", styles["Heading2"]))
    elements.append(Spacer(1, 0.3 * inch))

    elements.append(Paragraph(f"Subject: {subject.subject_name}", styles["Normal"]))
    elements.append(Paragraph(f"Week Starting: {start_week}", styles["Normal"]))
    elements.append(Paragraph(f"Total Class Sessions: {unique_classes}", styles["Normal"]))
    elements.append(Paragraph(f"Total Students: {len(students)}", styles["Normal"]))
    elements.append(Spacer(1, 0.2 * inch))

    # Add Rank column
    table_data = [
        ["Rank", "Roll No", "Name", "Total", "Present", "Absent", "Percentage"]
    ]

    for index, row in enumerate(student_rows, start=1):
        table_data.append([  # type: ignore
            index,
            row["roll"],
            row["name"],
            row["total"],
            row["present"],
            row["absent"],
            f'{row["percentage"]}%'
        ])

    
    total_students = len(students)

    class_average = round(  # type: ignore
        (total_present / (unique_classes * total_students)) * 100, 2  # type: ignore
    ) if unique_classes > 0 and total_students > 0 else 0

    report_format = get_report_format("attendance")
    if report_format == "excel":
        output = BytesIO()
        df = pd.DataFrame(student_rows)
        df.to_excel(output, index=False, engine="openpyxl")
        output.seek(0)
        return StreamingResponse(
            output,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f"attachment; filename=weekly_report_{subject_id}.xlsx"}
        )

    if report_format == "docx":
        try:
            from docx import Document
        except ImportError:
            raise HTTPException(status_code=500, detail="DOCX export support is unavailable")

        document = Document()
        document.add_heading("Weekly Attendance Report", level=1)
        document.add_paragraph(f"Subject: {subject.subject_name}")
        document.add_paragraph(f"Week Starting: {start_week}")
        document.add_paragraph(f"Class Average: {class_average}%")

        table_docx = document.add_table(rows=1, cols=7)
        hdr_cells = table_docx.rows[0].cells
        hdr_cells[0].text = "Rank"
        hdr_cells[1].text = "Roll No"
        hdr_cells[2].text = "Name"
        hdr_cells[3].text = "Total"
        hdr_cells[4].text = "Present"
        hdr_cells[5].text = "Absent"
        hdr_cells[6].text = "Percentage"

        for index, row in enumerate(student_rows, start=1):
            cells = table_docx.add_row().cells
            cells[0].text = str(index)
            cells[1].text = str(row["roll"])
            cells[2].text = str(row["name"])
            cells[3].text = str(row["total"])
            cells[4].text = str(row["present"])
            cells[5].text = str(row["absent"])
            cells[6].text = f"{row['percentage']}%"

        output = BytesIO()
        document.save(output)
        output.seek(0)
        return StreamingResponse(
            output,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=weekly_report_{subject_id}.docx"}
        )

    table = Table(table_data, repeatRows=1)

    style = TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
        ("GRID", (0, 0), (-1, -1), 1, colors.grey),
        ("ALIGN", (3, 1), (-1, -1), "CENTER"),
    ])

    threshold = get_setting("attendance_threshold") or 75
    print("Using attendance threshold:", threshold)

    # -----------------------------
    # Highlight Defaulters (<threshold %)
    # -----------------------------
    for i, row in enumerate(student_rows, start=1):
        if row["percentage"] < threshold:
            style.add(
                "BACKGROUND",
                (0, i),      # from Rank column
                (-1, i),     # entire row
                colors.lightcoral
            )

    table.setStyle(style)

    elements.append(table)
    elements.append(Spacer(1, 0.3 * inch))
    elements.append(Paragraph(f"Class Average: {class_average}%", styles["Heading3"]))

    doc.build(elements)

    return FileResponse(
        file_path,
        media_type="application/pdf",
        filename=file_path
    )


# =========================
# FACULTY – DOWNLOAD MONTHLY PDF (RANK + DEFAULTER)
# =========================
@app.get("/faculty/attendance/monthly/{subject_id}/download")
def download_monthly_pdf(
    subject_id: int,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):

    if current_user["role"] != "faculty":
        raise HTTPException(status_code=403, detail="Faculty only")

    from datetime import date

    today = date.today()
    month = today.month
    year = today.year

    # 🔒 Validate assignment
    assignment = db.query(FacultySubject).filter(
        FacultySubject.faculty_id == current_user["user_id"],
        FacultySubject.subject_id == subject_id,
        FacultySubject.is_active == True
    ).first()

    if not assignment:
        raise HTTPException(status_code=403, detail="Not assigned to this subject")

    subject = db.query(Subject).filter(
        Subject.subject_id == subject_id
    ).first()

    unique_classes = db.query(Attendance.attendance_date).filter(
        Attendance.subject_id == subject_id,
        extract('month', Attendance.attendance_date) == month,
        extract('year', Attendance.attendance_date) == year
    ).distinct().count()

    # Get subject department
    department_id = subject.department_id

    students = (
        db.query(Student, User)
        .join(User, Student.student_id == User.user_id)
        .filter(
            Student.year == assignment.year,
            Student.section == assignment.section,
            User.department_id == department_id,
            User.is_deleted == False
        )
        .all()
    )

    # -----------------------------
    # Collect student performance
    # -----------------------------
    student_rows = []
    total_present = 0

    for student, user in students:

        records = db.query(Attendance).filter(
            Attendance.student_id == student.student_id,
            Attendance.subject_id == subject_id,
            extract('month', Attendance.attendance_date) == month,
            extract('year', Attendance.attendance_date) == year
        ).all()

        total = len(records)
        present = len([r for r in records if r.status])
        absent = total - present
        percent = round((present / total) * 100, 2) if total > 0 else 0  # type: ignore

        total_present += present

        student_rows.append({
            "roll": student.roll_no,
            "name": user.name,
            "total": total,
            "present": present,
            "absent": absent,
            "percentage": percent
        })

    # -----------------------------
    # Sort by percentage (DESC)
    # -----------------------------
    student_rows.sort(key=lambda x: x["percentage"], reverse=True)

    # -----------------------------
    # PDF Setup
    # -----------------------------
    file_path = f"monthly_report_{subject_id}.pdf"
    doc = SimpleDocTemplate(file_path)
    elements = []
    styles = getSampleStyleSheet()

    elements.append(Paragraph("GVP-MAAA College", styles["Title"]))
    elements.append(Paragraph("Monthly Attendance Report", styles["Heading2"]))
    elements.append(Spacer(1, 0.3 * inch))

    elements.append(Paragraph(f"Subject: {subject.subject_name}", styles["Normal"]))
    elements.append(Paragraph(f"Month: {month} / {year}", styles["Normal"]))
    elements.append(Spacer(1, 0.2 * inch))

    # Add Rank column
    table_data = [
        ["Rank", "Roll No", "Name", "Total", "Present", "Absent", "Percentage"]
    ]

    for index, row in enumerate(student_rows, start=1):
        table_data.append([  # type: ignore
            index,
            row["roll"],
            row["name"],
            row["total"],
            row["present"],
            row["absent"],
            f'{row["percentage"]}%'
        ])

    total_students = len(students)

    class_average = round(  # type: ignore
        (total_present / (unique_classes * total_students)) * 100, 2  # type: ignore
    ) if unique_classes > 0 and total_students > 0 else 0

    report_format = get_report_format("attendance")
    if report_format == "excel":
        output = BytesIO()
        df = pd.DataFrame(student_rows)
        df.to_excel(output, index=False, engine="openpyxl")
        output.seek(0)
        return StreamingResponse(
            output,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f"attachment; filename=monthly_report_{subject_id}.xlsx"}
        )

    if report_format == "docx":
        try:
            from docx import Document
        except ImportError:
            raise HTTPException(status_code=500, detail="DOCX export support is unavailable")

        document = Document()
        document.add_heading("Monthly Attendance Report", level=1)
        document.add_paragraph(f"Subject: {subject.subject_name}")
        document.add_paragraph(f"Month: {month}/{year}")
        document.add_paragraph(f"Class Average: {class_average}%")

        table_docx = document.add_table(rows=1, cols=7)
        hdr_cells = table_docx.rows[0].cells
        hdr_cells[0].text = "Rank"
        hdr_cells[1].text = "Roll No"
        hdr_cells[2].text = "Name"
        hdr_cells[3].text = "Total"
        hdr_cells[4].text = "Present"
        hdr_cells[5].text = "Absent"
        hdr_cells[6].text = "Percentage"

        for index, row in enumerate(student_rows, start=1):
            cells = table_docx.add_row().cells
            cells[0].text = str(index)
            cells[1].text = str(row["roll"])
            cells[2].text = str(row["name"])
            cells[3].text = str(row["total"])
            cells[4].text = str(row["present"])
            cells[5].text = str(row["absent"])
            cells[6].text = f"{row['percentage']}%"

        output = BytesIO()
        document.save(output)
        output.seek(0)
        return StreamingResponse(
            output,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=monthly_report_{subject_id}.docx"}
        )

    table = Table(
        table_data,
        colWidths=[0.7*inch, 1*inch, 1.5*inch, 0.7*inch, 0.7*inch, 0.7*inch, 0.8*inch],
        repeatRows=1
    )

    style = TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
        ("GRID", (0, 0), (-1, -1), 1, colors.grey),
        ("ALIGN", (3, 1), (-1, -1), "CENTER"),
    ])

    threshold = get_setting("attendance_threshold") or 75
    print("Using attendance threshold:", threshold)

    # -----------------------------
    # Highlight Defaulters (<threshold %)
    # -----------------------------
    for i, row in enumerate(student_rows, start=1):
        if row["percentage"] < threshold:
            style.add(
                "BACKGROUND",
                (0, i),
                (-1, i),
                colors.lightcoral
            )

    table.setStyle(style)

    elements.append(table)
    elements.append(Spacer(1, 0.3 * inch))
    elements.append(Paragraph(f"Class Average: {class_average}%", styles["Heading3"]))

    doc.build(elements)

    return FileResponse(
        file_path,
        media_type="application/pdf",
        filename=file_path
    )


# -------------------------
# FORGOT PASSWORD
# -------------------------
@app.post("/forgot-password")
def forgot_password(email: str, db: Session = Depends(get_db)):

    user = db.query(User).filter(User.email == email).first()

    # Security best practice
    if not user:
        return {"message": "If the email exists, a reset link has been sent"}

    reset_token = create_reset_token(email)
    reset_link = f"http://localhost:5173/reset-password?token={reset_token}"

    send_reset_email(email, reset_link)   # ✅ EMAIL SENT HERE

    return {"message": "Password reset link sent"}


# -------------------------
# RESET PASSWORD
# -------------------------
@app.post("/reset-password")
def reset_password(
    data: ResetPasswordRequest,
    db: Session = Depends(get_db)
):
    # Verify token
    email = verify_reset_token(data.token)

    user = db.query(User).filter(User.email == email).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    # bcrypt safety
    if len(data.new_password) > 72:
        raise HTTPException(
            status_code=400,
            detail="Password must be 72 characters or less"
        )

    user.password = hash_password(data.new_password)
    db.commit()

    return {"message": "Password reset successful"}


# -------------------------
# scheduler to check attendance thresholds and send alerts (students)
# -------------------------
def check_attendance_thresholds():
    db = SessionLocal()
    attendance_enabled = get_setting("attendance_alert_enabled")
    attendance_enabled = attendance_enabled if attendance_enabled is not None else True
    threshold = float(get_setting("attendance_threshold") or 75)
    cgpa_threshold = float(get_setting("cgpa_threshold") or 6.5)

    try:
        if not attendance_enabled:
            return

        students = db.query(Student).all()
        for student in students:
            alerts = generate_student_alerts(
                student_id=student.student_id,
                db=db,
                attendance_threshold=threshold,
                cgpa_threshold=cgpa_threshold,
            )

            for alert_payload in alerts:
                if alert_payload.get("type") not in {"attendance-monitor", "marks-monitor"}:
                    continue

                duplicate = db.query(Alert).filter(
                    Alert.student_id == student.student_id,
                    Alert.type == alert_payload.get("type"),
                    Alert.message == alert_payload.get("message"),
                    Alert.created_at >= datetime.utcnow() - timedelta(hours=24),
                ).first()
                if duplicate:
                    continue

                alert = Alert(
                    title=alert_payload.get("title") or "Alert",
                    message=alert_payload.get("message") or "",
                    type=alert_payload.get("type") or "academic-monitor",
                    target_role="student",
                    target_type="individual",
                    student_id=student.student_id,
                )
                db.add(alert)
                db.commit()
                db.refresh(alert)

                db.add(AlertRecipient(alert_id=alert.id, user_id=student.student_id, is_read=False))
                db.commit()

    except Exception as e:
        print("Attendance/Marks scheduler error:", e)

    finally:
        db.close()


def check_cgpa_thresholds():
    db = SessionLocal()
    cgpa_enabled = get_setting("cgpa_alert_enabled")
    cgpa_enabled = cgpa_enabled if cgpa_enabled is not None else True
    threshold = float(get_setting("cgpa_threshold") or 6.5)
    attendance_threshold = float(get_setting("attendance_threshold") or 75)

    try:
        if not cgpa_enabled:
            return

        students = db.query(Student).all()
        for student in students:
            alerts = generate_student_alerts(
                student_id=student.student_id,
                db=db,
                attendance_threshold=attendance_threshold,
                cgpa_threshold=threshold,
            )

            for alert_payload in alerts:
                if alert_payload.get("type") != "cgpa-monitor":
                    continue

                duplicate = db.query(Alert).filter(
                    Alert.student_id == student.student_id,
                    Alert.type == "cgpa-monitor",
                    Alert.message == alert_payload.get("message"),
                    Alert.created_at >= datetime.utcnow() - timedelta(hours=24),
                ).first()
                if duplicate:
                    continue

                alert = Alert(
                    title=alert_payload.get("title") or "⚠ CGPA Alert",
                    message=alert_payload.get("message") or NO_DATA_MESSAGE,
                    type="cgpa-monitor",
                    target_role="student",
                    target_type="individual",
                    student_id=student.student_id,
                )
                db.add(alert)
                db.commit()
                db.refresh(alert)

                db.add(AlertRecipient(alert_id=alert.id, user_id=student.student_id, is_read=False))
                db.commit()

    except Exception as e:
        print("CGPA Alert Scheduler Error:", e)

    finally:
        db.close()


# -------------------------
# scheduler to check attendance thresholds and send alerts (parents)
# -------------------------
def check_monthly_faculty_attendance():

    db = SessionLocal()

    try:
        today = datetime.utcnow()

        # 👉 Calculate previous month properly
        if today.month == 1:
            target_month = 12
            target_year = today.year - 1
        else:
            target_month = today.month - 1
            target_year = today.year

        assignments = db.query(FacultySubject).filter(
            FacultySubject.is_active == True
        ).all()

        for assignment in assignments:

            faculty_id = assignment.faculty_id
            subject_id = assignment.subject_id
            year_class = assignment.year
            section = assignment.section

            subject = db.query(Subject).filter(
                Subject.subject_id == subject_id
            ).first()

            students = (
                db.query(Student, User)
                .join(User, Student.student_id == User.user_id)
                .filter(
                    Student.year == year_class,
                    Student.section == section,
                    User.department_id == subject.department_id,
                    User.is_deleted == False
                )
                .all()
            )

            below_60_count = 0

            for student, user in students:

                records = db.query(Attendance).filter(
                    Attendance.student_id == student.student_id,
                    Attendance.subject_id == subject_id,
                    extract('month', Attendance.attendance_date) == target_month,
                    extract('year', Attendance.attendance_date) == target_year
                ).all()

                total = len(records)

                if total < 5:
                    continue  # avoid noise

                present = len([r for r in records if r.status])
                percentage = (present / total) * 100 if total > 0 else 0

                if percentage < 60:
                    below_60_count += 1  # type: ignore

            if below_60_count == 0:
                continue

            # 🔒 Prevent duplicate monthly alerts
            existing = db.query(FacultyMonthlyAttendanceAlert).filter(
                FacultyMonthlyAttendanceAlert.faculty_id == faculty_id,
                FacultyMonthlyAttendanceAlert.subject_id == subject_id,
                FacultyMonthlyAttendanceAlert.year == year_class,
                FacultyMonthlyAttendanceAlert.section == section,
                FacultyMonthlyAttendanceAlert.month == target_month,
                FacultyMonthlyAttendanceAlert.year_value == target_year
            ).first()

            if existing:
                continue

            # ✅ Create alert
            alert = Alert(
                title="📊 Monthly Attendance Risk Summary",
                message=(
                    f"{below_60_count} students in "
                    f"{year_class}-{section} "
                    f"({subject.subject_name}) "
                    f"were below 60% attendance in "
                    f"{target_month}/{target_year}."
                ),
                type="monthly-attendance-summary",
                target_role="faculty",
                target_type="individual",
                faculty_id=faculty_id
            )

            db.add(alert)
            db.commit()
            db.refresh(alert)

            recipient = AlertRecipient(
                alert_id=alert.id,
                user_id=faculty_id,
                is_read=False
            )

            db.add(recipient)

            tracking = FacultyMonthlyAttendanceAlert(
                faculty_id=faculty_id,
                subject_id=subject_id,
                year=year_class,
                section=section,
                month=target_month,
                year_value=target_year,
                last_sent=datetime.utcnow()
            )

            db.add(tracking)
            db.commit()

    except Exception as e:
        print("Monthly Faculty Scheduler Error:", e)

    finally:
        db.close()


@app.get("/teacher/my-subjects")
def get_teacher_subjects(
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):

    if current_user["role"] != "faculty":
        raise HTTPException(status_code=403, detail="Only teachers allowed")

    subjects = db.query(FacultySubject).filter(
        FacultySubject.faculty_id == current_user["user_id"],
        FacultySubject.is_active == True
    ).all()

    result = []

    for s in subjects:
        subject = db.query(Subject).filter(
            Subject.subject_id == s.subject_id
        ).first()

        result.append({
            "subject_id": s.subject_id,
            "subject_name": subject.subject_name if subject else "Unknown",
            "year": s.year,
            "section": s.section
        })

    return {"subjects": result}


# ========================
# ASSIGNMENT ENDPOINTS
# ========================

@app.post("/teacher/create-assignment")
def create_assignment(
    title: str = Form(...),
    description: Optional[str] = Form(None),
    subject_id: int = Form(...),
    year: int = Form(...),
    section: str = Form(...),
    due_date: str = Form(...),
    file: Optional[UploadFile] = File(None),
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    try:
        if current_user["role"] != "faculty":
            raise HTTPException(status_code=403, detail="Only teachers can create assignments")

        faculty = db.query(Faculty).filter(
            Faculty.faculty_id == current_user["user_id"]
        ).first()

        if not faculty:
            raise HTTPException(status_code=404, detail="Faculty not found")

        faculty_subject = db.query(FacultySubject).filter(
            FacultySubject.faculty_id == current_user["user_id"],
            FacultySubject.subject_id == subject_id,
            FacultySubject.year == year,
            FacultySubject.section == section
        ).first()

        if not faculty_subject:
            raise HTTPException(
                status_code=403,
                detail="You are not assigned to teach this class/subject"
            )

        # Parse due_date
        try:
            due_date_parsed = datetime.fromisoformat(due_date.replace("Z", "+00:00"))
        except Exception:
            due_date_parsed = datetime.strptime(due_date[:10], "%Y-%m-%d")  # type: ignore

        # Handle optional file upload
        file_name = None
        file_path = None
        if file and file.filename:
            upload_dir = "uploads/assignments"
            os.makedirs(upload_dir, exist_ok=True)
            unique_filename = f"{uuid.uuid4()}_{file.filename}"
            file_path = os.path.join(upload_dir, unique_filename)
            with open(file_path, "wb") as f:
                f.write(file.file.read())
            file_name = file.filename

        new_assignment = Assignment(
            title=title,
            description=description,
            faculty_id=current_user["user_id"],
            subject_id=subject_id,
            year=year,
            section=section,
            due_date=due_date_parsed,
            is_active=True,
            file_name=file_name,
            file_path=file_path
        )

        db.add(new_assignment)
        db.commit()
        db.refresh(new_assignment)

        # =========================
        # CREATE ALERT FOR STUDENTS
        # =========================

        subject = db.query(Subject).filter(
            Subject.subject_id == subject_id
        ).first()

        faculty_user = db.query(User).filter(
            User.user_id == current_user["user_id"]
        ).first()

        students = (
            db.query(Student, User)
            .join(User, Student.student_id == User.user_id)
            .filter(
                Student.year == year,
                Student.section == section,
                User.department_id == subject.department_id,
                User.is_deleted == False
            )
            .all()
        )

        alert = Alert(
            title="📚 New Assignment Posted",
            message=f"{faculty_user.name} posted '{title}' for {subject.subject_name}. Due: {due_date_parsed.date()}",
            type="assignment",
            target_role="student",
            target_type="class"
        )

        db.add(alert)
        db.commit()
        db.refresh(alert)

        for student, user in students:
            recipient = AlertRecipient(
                alert_id=alert.id,
                user_id=user.user_id,
                is_read=False
            )
            db.add(recipient)

        db.commit()

        

        return {
            "status": "success",
            "message": "Assignment created successfully",
            "assignment_id": new_assignment.id
        }

    except HTTPException as e:
        raise e
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/teacher/assignments/{year}/{section}")
def get_teacher_assignments(
    year: int,
    section: str,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):

    if current_user["role"] != "faculty":
        raise HTTPException(status_code=403, detail="Only teachers can access this")

    assignments = db.query(Assignment).filter(
        Assignment.faculty_id == current_user["user_id"],
        Assignment.year == year,
        Assignment.section == section
    ).order_by(Assignment.created_at.desc()).all()

    result = []

    for assignment in assignments:

        submissions = db.query(AssignmentSubmission).filter(
            AssignmentSubmission.assignment_id == assignment.id
        ).all()

        submitted_count = len([s for s in submissions if s.is_submitted])
        total_submitted = len(submissions)

        total_students = db.query(Student).filter(
            Student.year == assignment.year,
            Student.section == assignment.section,
            Student.is_deleted == False
        ).count()

        result.append({
            "id": assignment.id,
            "title": assignment.title,
            "description": assignment.description,
            "due_date": assignment.due_date,
            "created_at": assignment.created_at,
            "subject_id": assignment.subject_id,
            "submitted": submitted_count,
            "total_students": total_students,
            "pending": total_students - total_submitted,
            "status": "Active" if assignment.is_active else "Inactive"
        })

    return {"status": "success", "assignments": result}



@app.get("/teacher/assignment-details/{assignment_id}")
def get_assignment_details(
    assignment_id: int,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):

    if current_user["role"] != "faculty":
        raise HTTPException(status_code=403, detail="Only teachers can access this")

    assignment = db.query(Assignment).filter(
        Assignment.id == assignment_id,
        Assignment.faculty_id == current_user["user_id"]
    ).first()

    if not assignment:
        raise HTTPException(status_code=404, detail="Assignment not found")

    students = db.query(Student).filter(
        Student.year == assignment.year,
        Student.section == assignment.section,
        Student.is_deleted == False
    ).all()

    submissions = db.query(AssignmentSubmission).filter(
        AssignmentSubmission.assignment_id == assignment_id
    ).all()

    submitted_ids = {s.student_id for s in submissions if s.is_submitted}

    submitted_students = []
    pending_students = []

    for student in students:
        user = db.query(User).filter(User.user_id == student.student_id).first()

        student_info = {
            "name": user.name if user else "Unknown",
            "roll": student.roll_no,
            "student_id": student.student_id
        }

        # Find the specific submission for this student among those submitted
        submission = next((s for s in submissions if s.student_id == student.student_id and s.is_submitted), None)

        if submission:
            submitted_students.append({
                "submission_id": submission.id,
                "student_id": submission.student_id,
                "name": user.name if user else "Unknown",
                "roll": student.roll_no,
                "file_path": submission.file_path,
                "status": submission.status
            })
        else:
            pending_students.append(student_info)

    return {
        "status": "success",
        "assignment": {
            "id": assignment.id,
            "title": assignment.title,
            "description": assignment.description,
            "due_date": assignment.due_date,
            "created_at": assignment.created_at,
            "subject_id": assignment.subject_id,
            "year": assignment.year,
            "section": assignment.section
        },
        "submitted": submitted_students,
        "pending": pending_students
    }

@app.get("/student/assignments")
def get_student_assignments(
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Get all assignments for the logged-in student
    """
    try:
        if current_user["role"] != "student":
            raise HTTPException(status_code=403, detail="Only students can access this")

        student = db.query(Student).filter(
            Student.student_id == current_user["user_id"]
        ).first()

        if not student:
            raise HTTPException(status_code=404, detail="Student not found")

        # Get all assignments for student's class
        assignments = db.query(Assignment).filter(
            Assignment.year == student.year,
            Assignment.section == student.section,
            Assignment.is_active == True
        ).order_by(Assignment.due_date).all()

        # Check which ones are submitted
        submissions = db.query(AssignmentSubmission).filter(
            AssignmentSubmission.student_id == current_user["user_id"]
        ).all()

        submitted_ids = {s.assignment_id for s in submissions}

        result = []
        for assignment in assignments:
            subject = db.query(Subject).filter(
                Subject.subject_id == assignment.subject_id
            ).first()

            submission = next((s for s in submissions if s.assignment_id == assignment.id), None)
            is_late = submission.is_late if submission else False
            status = submission.status if submission else "pending"
            
            # If submitted but status is still 'pending' at backend, show as 'submitted' for frontend 
            if submission and status == "pending":
                status = "submitted"

            result.append({
                "id": assignment.id,
                "title": assignment.title,
                "description": assignment.description,
                "subject": subject.subject_name if subject else "Unknown",
                "due_date": assignment.due_date,
                "created_at": assignment.created_at,
                "status": status,
                "is_late": is_late
            })

        return {"status": "success", "assignments": result}

    except HTTPException as e:
        raise e
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        db.close()


@app.post("/student/submit-assignment/{assignment_id}")
async def submit_assignment(
    assignment_id: int,
    submission_text: Optional[str] = Form(None),
    file: UploadFile = File(None),
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Student submits an assignment
    """
    try:
        if current_user["role"] != "student":
            raise HTTPException(status_code=403, detail="Only students can submit")

        assignment = db.query(Assignment).filter(
            Assignment.id == assignment_id
        ).first()

        if not assignment:
            raise HTTPException(status_code=404, detail="Assignment not found")

        student = db.query(Student).filter(
            Student.student_id == current_user["user_id"]
        ).first()

        if not student:
            raise HTTPException(status_code=404, detail="Student not found")

        # Check if student belongs to this class
        if student.year != assignment.year or student.section != assignment.section:
            raise HTTPException(status_code=403, detail="This assignment is not for your class")

        # Check if already submitted
        existing_submission = db.query(AssignmentSubmission).filter(
            AssignmentSubmission.assignment_id == assignment_id,
            AssignmentSubmission.student_id == current_user["user_id"]
        ).first()

        if existing_submission:
            raise HTTPException(status_code=400, detail="You have already submitted this assignment")

        # Handle file upload
        file_name = None
        file_path = None
        file_type = None

        if file:
            try:
                # Create submissions folder if it doesn't exist
                os.makedirs("uploads/assignments", exist_ok=True)

                file_ext = file.filename.split(".")[-1]
                file_name = f"assignment_{assignment_id}_student_{current_user['user_id']}.{file_ext}"
                file_path = f"uploads/assignments/{file_name}"

                content = await file.read()
                with open(file_path, "wb") as f:
                    f.write(content)

                file_type = file_ext
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"File upload failed: {str(e)}")

        # Check if late
        is_late = datetime.utcnow() > assignment.due_date

        # Create submission
        submission = AssignmentSubmission(
            assignment_id=assignment_id,
            student_id=current_user["user_id"],
            file_name=file_name,
            file_path=file_path,
            file_type=file_type,
            submission_text=submission_text,
            submitted_at=datetime.utcnow(),
            is_late=is_late,
            is_submitted=True,
            status="pending"
        )

        db.add(submission)
        db.commit()
        db.refresh(submission)

        return {
            "status": "success",
            "message": "Assignment submitted successfully",
            "is_late": is_late,
            "submission_id": submission.id
        }

    except HTTPException as e:
        raise e
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/teacher/student-assignments-summary/{year}/{section}")
def get_student_assignments_summary(
    year: int,
    section: str,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    try:
        if current_user["role"] != "faculty":
            raise HTTPException(status_code=403, detail="Only teachers can access this")

        # Get the 5 most recent assignments for this class
        recent_assignments = db.query(Assignment).filter(
            Assignment.faculty_id == current_user["user_id"],
            Assignment.year == year,
            Assignment.section == section,
            Assignment.is_active == True
        ).order_by(Assignment.created_at.desc()).limit(5).all()

        assignment_ids = [a.id for a in recent_assignments]

        students_query = db.query(Student, User).join(
            User, Student.student_id == User.user_id
        ).filter(
            Student.year == year,
            Student.section == section,
            Student.is_deleted == False
        ).all()

        student_summaries = []

        for student, user in students_query:
            # Avoid SQLAlchemy crash when list is empty
            if assignment_ids:
                submissions = db.query(AssignmentSubmission).filter(
                    AssignmentSubmission.student_id == student.student_id,
                    AssignmentSubmission.assignment_id.in_(assignment_ids)
                ).all()
            else:
                submissions = []

            submission_map = {s.assignment_id: s for s in submissions}

            recent_assignment_dots = []
            for assignment in recent_assignments:
                now = datetime.utcnow()

                if assignment.id in submission_map:
                    sub = submission_map[assignment.id]
                    status = sub.status  # "pending", "approved", "rejected"
                else:
                    # Compare with timezone-naive datetime
                    due = assignment.due_date
                    if hasattr(due, 'tzinfo') and due.tzinfo is not None:
                        due = due.replace(tzinfo=None)
                    status = "future" if due > now else "not_submitted"

                recent_assignment_dots.append({
                    "assignment_id": assignment.id,
                    "title": assignment.title,
                    "status": status,
                    "due_date": assignment.due_date.isoformat() if assignment.due_date else None
                })

            student_summaries.append({
                "student_id": student.student_id,
                "name": user.name if user else "Unknown",
                "roll": student.roll_no if hasattr(student, 'roll_no') else "",
                "year": student.year,
                "section": student.section,
                "recent_assignments": recent_assignment_dots
            })

        return {"status": "success", "students": student_summaries}

    except HTTPException as e:
        raise e
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error fetching student summaries: {str(e)}")

@app.put("/teacher/assignment-submissions/{submission_id}/status")
def update_submission_status(
    submission_id: int,
    status_data: StatusUpdateRequest,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "faculty":
        raise HTTPException(status_code=403, detail="Only teachers can access this")

    submission = db.query(AssignmentSubmission).filter(
        AssignmentSubmission.id == submission_id
    ).first()

    if not submission:
        raise HTTPException(status_code=404, detail="Submission not found")

    # Verify teacher owns the assignment
    assignment = db.query(Assignment).filter(
        Assignment.id == submission.assignment_id,
        Assignment.faculty_id == current_user["user_id"]
    ).first()

    if not assignment:
        raise HTTPException(status_code=403, detail="You do not have permission to update this submission")

    if status_data.status not in ["approved", "rejected", "pending"]:
         raise HTTPException(status_code=400, detail="Invalid status. Must be approved, rejected, or pending")

    if submission.status == "approved" and status_data.status == "approved":
        raise HTTPException(
            status_code=400,
            detail="Assignment already approved"
        )

    submission.status = status_data.status
    db.commit()

    return {"status": "success", "message": f"Submission status updated to {status_data.status}"}



def check_assignment_deadlines():

    db = SessionLocal()

    try:
        today = datetime.utcnow().date()

        assignments = db.query(Assignment).filter(
            Assignment.is_active == True
        ).all()

        for assignment in assignments:

            due_date = assignment.due_date.date()
            days_left = (due_date - today).days

            # Send alerts only from 2 days before deadline until deadline day
            if days_left < 0 or days_left > 2:
                continue

            # -----------------------------
            # Determine message for teacher
            # -----------------------------
            if days_left == 2:
                deadline_text = "deadline in 2 days"
            elif days_left == 1:
                deadline_text = "deadline tomorrow"
            else:
                deadline_text = "deadline today"

            # -----------------------------
            # Get students in that class
            # -----------------------------
            students = db.query(Student).filter(
                Student.year == assignment.year,
                Student.section == assignment.section,
                Student.is_deleted == False
            ).all()

            total_students = len(students)

            # -----------------------------
            # Get submissions
            # -----------------------------
            submissions = db.query(AssignmentSubmission).filter(
                AssignmentSubmission.assignment_id == assignment.id,
                AssignmentSubmission.is_submitted == True
            ).all()

            submitted = len(submissions)
            pending = total_students - submitted

            # -----------------------------
            # Get subject info
            # -----------------------------
            subject = db.query(Subject).filter(
                Subject.subject_id == assignment.subject_id
            ).first()

            # -----------------------------
            # Create alert
            # -----------------------------
            alert = Alert(
                title="📌 Assignment Deadline Reminder",
                message=(
                    f"{subject.subject_name} - {assignment.title}\n"
                    f"{deadline_text}\n\n"
                    f"Total Students: {total_students} | "
                    f"Submitted: {submitted} | Pending: {pending}"
                ),
                type="assignment-reminder",
                target_role="faculty",
                target_type="individual",
                faculty_id=assignment.faculty_id
            )

            db.add(alert)
            db.commit()
            db.refresh(alert)

            # -----------------------------
            # Send to teacher alerts page
            # -----------------------------
            recipient = AlertRecipient(
                alert_id=alert.id,
                user_id=assignment.faculty_id,
                is_read=False
            )

            db.add(recipient)
            db.commit()

    except Exception as e:
        print("Assignment Deadline Scheduler Error:", e)

    finally:
        db.close()



scheduler = BackgroundScheduler()

scheduler.add_job(
    check_monthly_faculty_attendance,
    "cron",
    day=1,
    hour=9,
    minute=0
)

scheduler.add_job(
    check_assignment_deadlines,
    "cron",
    hour=18,
    minute=0
)

@app.on_event("startup")
def start_scheduler():
    scheduler.add_job(
        check_attendance_thresholds,
        "cron",
        hour=20,
        minute=0
    )
    scheduler.add_job(
        check_cgpa_thresholds,
        "cron",
        hour=20,
        minute=5
    )
    scheduler.add_job(
        process_event_reminders,
        "cron",
        hour=8,
        minute=0
    )
    scheduler.start()


# -----------------------------
# Upload resource
# -----------------------------     
@app.post("/faculty/upload-resource")
async def upload_resource(
    title: str = Form(...),
    description: str = Form(...),
    subject_id: int = Form(...),
    type: str = Form(...),
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):

    try:
        if current_user["role"] != "faculty":
            raise HTTPException(status_code=403)

        UPLOAD_DIR = "uploads/resources"
        os.makedirs(UPLOAD_DIR, exist_ok=True)

        unique_name = f"{uuid.uuid4()}_{file.filename}"
        file_location = os.path.join(UPLOAD_DIR, unique_name)

        with open(file_location, "wb") as buffer:
            buffer.write(await file.read())

        resource = Resource(
            title=title,
            description=description,
            subject_id=subject_id,
            faculty_id=current_user["user_id"],
            type=type,
            file_url=file_location,
            created_at=datetime.utcnow()
        )

        db.add(resource)
        db.commit()
        db.refresh(resource)

        # START NEW ALERT LOGIC
        subject = db.query(Subject).filter(Subject.subject_id == subject_id).first()
        if subject:
            # Find all assigned classes for this faculty and subject
            assigned_classes = db.query(FacultySubject).filter(
                FacultySubject.subject_id == subject_id,
                FacultySubject.faculty_id == current_user["user_id"],
                FacultySubject.is_active == True
            ).all()

            for ac in assigned_classes:
                # Find all students in this year/section
                students = db.query(Student).filter(
                    Student.year == ac.year,
                    Student.section == ac.section
                ).all()

                for st in students:
                    # Create alert for each student
                    new_alert = Alert(
                        title="New Resource Uploaded",
                        message=f"A new resource '{title}' ({type}) has been uploaded by your faculty for {subject.subject_name}.",
                        type="resource",
                        target_role="student",
                        target_type="individual",
                        student_id=st.student_id,
                        faculty_id=current_user["user_id"]
                    )
                    db.add(new_alert)
                    db.flush()

                    db.add(AlertRecipient(
                        alert_id=new_alert.id,
                        user_id=st.student_id,
                        is_read=False
                    ))
            
            db.commit()

        return {"message": "Resource uploaded successfully"}

    except Exception as e:
        db.rollback()
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/student/resources")
def get_student_resources(
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):

    # Student's year and section
    student = db.query(Student).filter(
        Student.student_id == current_user["user_id"]
    ).first()

    if not student:
        raise HTTPException(status_code=404, detail="Student not found")

    # Find subjects assigned to this student's year/section
    faculty_subjects = db.query(FacultySubject).filter(
        FacultySubject.year == student.year,
        FacultySubject.section == student.section,
        FacultySubject.is_active == True
    ).all()

    subject_ids = [fs.subject_id for fs in faculty_subjects]

    resources = db.query(Resource, Subject).join(
        Subject, Resource.subject_id == Subject.subject_id
    ).filter(
        Resource.subject_id.in_(subject_ids)
    ).all()

    result = []
    for r, s in resources:
        result.append({
            "id": r.id,
            "title": r.title,
            "description": r.description,
            "type": r.type,
            "file_url": r.file_url,
            "created_at": r.created_at,
            "subject": s.subject_name
        })

    return result

@app.post("/student/resource-access/{resource_id}")
def track_access(
    resource_id: int,
    payload: ResourceAccessRequest,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):

    existing = db.query(ResourceAccess).filter(
        ResourceAccess.resource_id == resource_id,
        ResourceAccess.student_id == current_user["user_id"],
        ResourceAccess.action_type == payload.action_type
    ).first()

    if not existing:
        access = ResourceAccess(
            resource_id=resource_id,
            student_id=current_user["user_id"],
            action_type=payload.action_type,
            accessed_at=datetime.utcnow()
        )
        db.add(access)
        db.commit()

    return {"message": "Access recorded"}


@app.get("/faculty/resources/{subject_id}")
def faculty_resources(
    subject_id: int,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):

    resources = db.query(Resource, Subject).join(
        Subject, Resource.subject_id == Subject.subject_id
    ).filter(
        Resource.subject_id == subject_id,
        Resource.faculty_id == current_user["user_id"]
    ).all()

    result = []

    for r, s in resources:

        from sqlalchemy import func  # type: ignore
        accessed = db.query(func.count(func.distinct(ResourceAccess.student_id))).filter(
            ResourceAccess.resource_id == r.id
        ).scalar()

        downloads = db.query(func.count(func.distinct(ResourceAccess.student_id))).filter(
            ResourceAccess.resource_id == r.id,
            ResourceAccess.action_type == "download"
        ).scalar()

        # Find how many students are in the batches assigned to this subject+faculty
        assigned_classes = db.query(FacultySubject).filter(
            FacultySubject.faculty_id == current_user["user_id"],
            FacultySubject.subject_id == subject_id,
            FacultySubject.is_active == True
        ).all()
        
        total_students = 0
        for ac in assigned_classes:
            count = db.query(Student).filter(
                Student.year == ac.year,
                Student.section == ac.section
            ).count()
            total_students += count

        result.append({
            "id": r.id,
            "title": r.title,
            "type": r.type,
            "subject": s.subject_name,
            "created_at": r.created_at,
            "accessed": accessed,
            "downloads": downloads,
            "total_students": total_students
        })
    return result

@app.get("/faculty/subjects")
def get_faculty_subjects(
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "faculty":
        raise HTTPException(status_code=403, detail="Faculty only")

    subjects = db.query(FacultySubject, Subject).join(
        Subject, FacultySubject.subject_id == Subject.subject_id
    ).filter(
        FacultySubject.faculty_id == current_user["user_id"],
        FacultySubject.is_active == True
    ).all()

    result = []
    for fs, s in subjects:
        dept_name = DEPARTMENT_MAP.get(s.department_id, str(s.department_id))
        result.append({
            "subject_id": s.subject_id,
            "subject_name": s.subject_name,
            "year": fs.year,
            "section": fs.section,
            "department": dept_name
        })

    return result

@app.get("/faculty/resource-access-details/{resource_id}")
def get_resource_access_details(
    resource_id: int,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "faculty":
        raise HTTPException(status_code=403, detail="Faculty only")
        
    # Verify the resource belongs to this faculty
    resource = db.query(Resource).filter(
        Resource.id == resource_id,
        Resource.faculty_id == current_user["user_id"]
    ).first()
    
    if not resource:
        raise HTTPException(status_code=404, detail="Resource not found")
        
    accesses = db.query(ResourceAccess, Student, User).join(
        Student, ResourceAccess.student_id == Student.student_id
    ).join(
        User, Student.student_id == User.user_id
    ).filter(
        ResourceAccess.resource_id == resource_id
    ).order_by(ResourceAccess.accessed_at.desc()).all()
    
    result = []
    for ra, st, u in accesses:
        result.append({
            "student_id": st.student_id,
            "name": u.name,
            "roll_no": st.roll_no,
            "action_type": ra.action_type,
            "accessed_at": ra.accessed_at
        })
        
    return result


# ==========================================
# ADVANCED ALERT SYSTEM (FACULTY)
# ==========================================

@app.get("/faculty/search-students", response_model=List[schemas.StudentSearchResponse])
def search_students(
    q: str = Query(..., min_length=2),
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_user)
):
    """Search for specific students by name or roll number."""
    if current_user["role"] != "faculty":
        raise HTTPException(status_code=403, detail="Not authorized")
    
    search_term = f"%{q.lower()}%"
    
    # Needs to match User.name or Student.roll_no
    students = db.query(Student, User.name).join(
        User, Student.student_id == User.user_id
    ).filter(
        or_(
            func.lower(User.name).like(search_term),
            func.lower(Student.roll_no).like(search_term)
        )
    ).limit(10).all()
    
    # Return formatted objects
    results = [
        {"student_id": st.Student.student_id, "name": st.name, "roll_no": st.Student.roll_no}
        for st in students
    ]
    return results


@app.post("/faculty/send-alert")
def send_alert(
    alert_req: schemas.AlertSendRequest,
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_user)
):
    if current_user["role"] != "faculty":
        raise HTTPException(status_code=403, detail="Not authorized")
    
    faculty_id = current_user["user_id"]
    target_students = set()
    
    # Mode 1: Whole Class (Single Subject)
    if alert_req.target == "class" and alert_req.subject_id:
        fs = db.query(FacultySubject).filter(
            FacultySubject.subject_id == alert_req.subject_id,
            FacultySubject.faculty_id == faculty_id,
            FacultySubject.is_active == True
        ).first()
        if fs:
            st_list = db.query(Student.student_id).filter(
                Student.year == fs.year,
                Student.section == fs.section
            ).all()
            target_students.update([s.student_id for s in st_list])
            
    # Mode 2: Multiple Classes
    elif alert_req.target == "multiple_classes" and alert_req.subject_ids:
        for sid in alert_req.subject_ids:
            fs = db.query(FacultySubject).filter(
                FacultySubject.subject_id == sid,
                FacultySubject.faculty_id == faculty_id,
                FacultySubject.is_active == True
            ).first()
            if fs:
                st_list = db.query(Student.student_id).filter(
                    Student.year == fs.year,
                    Student.section == fs.section
                ).all()
                target_students.update([s.student_id for s in st_list])
                
    # Mode 3: Specific Students
    elif alert_req.target == "students" and alert_req.student_ids:
        target_students.update(alert_req.student_ids)
        
    if not target_students:
        raise HTTPException(status_code=400, detail="No students found for given targets")
    
    new_alerts = []
    recipients = []
    
    title_mapping = {
        "Emergency": "Emergency Announcement",
        "Announcement": "New Announcement",
        "Info": "Information Alert",
        "Reminder": "Reminder"
    }
    
    alert_title = title_mapping.get(alert_req.type, "Alert")
    
    for sid in target_students:
        new_alert = Alert(
            title=alert_title,
            message=alert_req.message,
            type=alert_req.type.lower(),
            target_role="student",
            target_type="individual",
            student_id=sid,
            faculty_id=faculty_id
        )
        new_alerts.append(new_alert)
        
    db.add_all(new_alerts)
    db.flush() # assign IDs
    
    for alert in new_alerts:
        recipients.append(AlertRecipient(
            alert_id=alert.id,
            user_id=alert.student_id,
            is_read=False
        ))
        
    db.add_all(recipients)
    db.commit()
    
    return {"message": "Alert sent successfully", "students_targeted": len(target_students)}


@app.post("/faculty/send-resource-reminder/{resource_id}")
def send_resource_reminder(
    resource_id: int,
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_user)
):
    if current_user["role"] != "faculty":
        raise HTTPException(status_code=403, detail="Not authorized")
        
    resource = db.query(Resource).filter(
        Resource.id == resource_id,
        Resource.faculty_id == current_user["user_id"]
    ).first()
    
    if not resource:
        raise HTTPException(status_code=404, detail="Resource not found")
        
    # Find the assigned class for this resource
    fs = db.query(FacultySubject).filter(
        FacultySubject.subject_id == resource.subject_id,
        FacultySubject.faculty_id == current_user["user_id"]
    ).first()
    
    if not fs:
        raise HTTPException(status_code=400, detail="Faculty subject mapping not found")
        
    # Get all students in this class
    all_students_in_class_subq = db.query(Student.student_id).filter(
        Student.year == fs.year,
        Student.section == fs.section
    ).subquery()
    
    # Get students who HAVE accessed it
    accessed_students_subq = db.query(ResourceAccess.student_id).filter(
        ResourceAccess.resource_id == resource_id
    ).distinct().subquery()
    
    # Find students who are in the class but NOT in the accessed list
    unaccessed_students = db.query(Student.student_id).filter(
        Student.student_id.in_(all_students_in_class_subq),
        ~Student.student_id.in_(accessed_students_subq)
    ).all()
    
    target_ids = [s.student_id for s in unaccessed_students]
    
    if not target_ids:
        return {"message": "All students have already accessed this resource", "sent_count": 0}
        
    # Bulk create reminders
    new_alerts = []
    
    for sid in target_ids:
        new_alert = Alert(
            title="Resource Reminder",
            message=f"Reminder: Please check the latest study material '{resource.title}'.",
            type="reminder",
            target_role="student",
            target_type="individual",
            student_id=sid,
            faculty_id=current_user["user_id"]
        )
        new_alerts.append(new_alert)
        
    db.add_all(new_alerts)
    db.flush()
    
    recipients = [
        AlertRecipient(alert_id=al.id, user_id=al.student_id, is_read=False)
        for al in new_alerts
    ]
    
    db.add_all(recipients)
    db.commit()
    
    return {"message": "Reminders sent successfully", "sent_count": len(target_ids)}


# ==========================================
# EVENTS MANAGEMENT API
# ==========================================

@app.post("/faculty/events", response_model=EventResponse)
def create_event(
    payload: EventCreate,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "faculty":
        raise HTTPException(status_code=403, detail="Faculty only")

    # -------------------------
    # Create Event
    # -------------------------
    new_event = Event(
        title=payload.title,
        description=payload.description,
        event_type=payload.event_type,
        organizer=payload.organizer,
        venue=payload.venue,
        location=payload.venue,
        event_date=payload.event_date,
        max_participants=payload.max_participants,
        registration_deadline=payload.registration_deadline,
        external_registration_link=payload.external_registration_link,
        year=payload.year,
        section=payload.section,
        created_by=current_user["user_id"],
        status="upcoming"
    )

    db.add(new_event)
    db.commit()
    db.refresh(new_event)

    # -------------------------
    # Find Target Students
    # -------------------------
    query = (
        db.query(Student.student_id)
        .join(User, Student.student_id == User.user_id)
        .filter(or_(User.is_deleted == False, User.is_deleted == None))
    )

    if payload.year != "All":
        query = query.filter(Student.year == int(payload.year))

    if payload.section != "All":
        query = query.filter(Student.section.ilike(payload.section))

    students = query.all()

    # Debug output
    print("TARGET YEAR:", payload.year)
    print("TARGET SECTION:", payload.section)
    print("STUDENTS FOUND:", students)

    # -------------------------
    # Create Alerts
    # -------------------------
    title = f"New Event Created: {payload.title}"
    message = f"New Event Created: {payload.title} on {payload.event_date.strftime('%d %b %Y')} at {payload.venue}."

    new_alerts = []

    for (sid,) in students:
        new_alert = Alert(
            title=title,
            message=message,
            type="announcement",
            target_role="student",
            target_type="individual",
            student_id=sid,
            faculty_id=current_user["user_id"]
        )
        new_alerts.append(new_alert)

    if len(new_alerts) > 0:
        db.add_all(new_alerts)
        db.flush()

        recipients = [
            AlertRecipient(alert_id=al.id, user_id=al.student_id, is_read=False)
            for al in new_alerts
        ]

        db.add_all(recipients)
        db.commit()

    # -------------------------
    # Prepare Response
    # -------------------------
    response_data = EventResponse.from_orm(new_event)
    response_data.total_students = len(students)
    response_data.present_count = 0
    response_data.absent_count = 0

    today = date.today()

    if new_event.event_date > today:
        response_data.status = "Upcoming"
    elif new_event.event_date == today:
        response_data.status = "Ongoing"
    else:
        response_data.status = "Completed"

    return response_data

@app.get("/faculty/events", response_model=List[EventResponse])
def get_events(
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "faculty":
        raise HTTPException(status_code=403, detail="Faculty only")

    events = db.query(Event).filter(Event.created_by == current_user["user_id"]).order_by(Event.event_date.desc()).all()
    
    results = []
    today = date.today()
    for ev in events:
        resp = EventResponse.from_orm(ev)
        
        # Adjust status on the fly
        if ev.event_date > today:
            resp.status = "Upcoming"
        elif ev.event_date == today:
            resp.status = "Ongoing"
        else:
            resp.status = "Completed"

        # Compute counts
        if resp.status == "Upcoming":
            resp.total_students = db.query(EventRegistration).filter(EventRegistration.event_id == ev.id).count()
            resp.present_count = 0
            resp.absent_count = 0
        else:
            total = db.query(EventRegistration).filter(EventRegistration.event_id == ev.id).count()
            present = db.query(EventRegistration).filter(EventRegistration.event_id == ev.id, EventRegistration.attendance == "present").count()
            absent = db.query(EventRegistration).filter(EventRegistration.event_id == ev.id, EventRegistration.attendance == "absent").count()
            
            resp.total_students = total
            resp.present_count = present
            resp.absent_count = absent
        
        results.append(resp)
        
    return results

@app.get("/faculty/events/{event_id}/attendance", response_model=EventAttendanceResponse)
def get_event_attendance(
    event_id: int,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "faculty":
        raise HTTPException(status_code=403, detail="Faculty only")

    event = db.query(Event).filter(Event.id == event_id, Event.created_by == current_user["user_id"]).first()
    if not event:
        raise HTTPException(status_code=404, detail="Event not found")
        
    today = date.today()
    event_status = "upcoming"
    if event.event_date == today:
        event_status = "ongoing"
    elif event.event_date < today:
        event_status = "completed"

    # EXTERNAL EVENT Check
    if event.event_type == "External":
        return EventAttendanceResponse(
            event_id=event.id,
            title=event.title,
            date=event.event_date,
            location=event.location,
            students=[],
            message="Attendance tracking is not required for external events."
        )

    # UPCOMING Internal
    if event_status == "upcoming":
        return EventAttendanceResponse(
            event_id=event.id,
            title=event.title,
            date=event.event_date,
            location=event.location,
            students=[],
            message="Attendance will be available when the event starts."
        )

    # COMPLETED Internal (Show stats, but roster loading depends on specific requirement)
    # The requirement says "Attendance roster must only load when: internal AND ongoing"
    # However, for completed, it says "Show: Final attendance statistics". 
    # I'll return the students for ongoing AND completed, but the frontend will disable editing for completed.
    # WAIT, Section 9 says "*MUST ONLY LOAD* when ONGOING". I will stick to that to be safe.
    
    if event_status == "completed":
        # Check if we should still return students for COMPLETED to show "Final stats"
        # The prompt says "Show: Final attendance statistics" for COMPLETED. 
        # Usually statistics are calculated from the registration records.
        # If I don't return students, the frontend and backend counts still work.
        return EventAttendanceResponse(
            event_id=event.id,
            title=event.title,
            date=event.event_date,
            location=event.location,
            students=[],
            message="Event completed. Viewing final statistics."
        )

    # ONGOING Internal
    attendance_records = (
        db.query(EventRegistration, Student, User)
        .join(Student, EventRegistration.student_id == Student.student_id)
        .join(User, Student.student_id == User.user_id)
        .filter(EventRegistration.event_id == event_id)
        .order_by(User.name.asc())
        .all()
    )
    
    if not attendance_records:
         return EventAttendanceResponse(
            event_id=event.id,
            title=event.title,
            date=event.event_date,
            location=event.location,
            students=[],
            message="No students registered yet."
        )

    students = []
    for reg, st, usr in attendance_records:
        students.append(EventStudentDetail(  # type: ignore
            student_id=st.student_id,
            name=usr.name,
            roll_no=st.roll_no,
            attendance_status=reg.attendance,
            result=reg.result
        ))
        
    return EventAttendanceResponse(
        event_id=event.id,
        title=event.title,
        date=event.event_date,
        location=event.location,
        students=students
    )

@app.patch("/faculty/events/{event_id}/attendance")
def update_event_attendance(
    event_id: int,
    payload: EventAttendanceUpdate,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "faculty":
        raise HTTPException(status_code=403, detail="Faculty only")

    # verify ownership
    event = db.query(Event).filter(Event.id == event_id, Event.created_by == current_user["user_id"]).first()
    if not event:
        raise HTTPException(status_code=404, detail="Event not found")
        
    att = db.query(EventRegistration).filter(
        EventRegistration.event_id == event_id,
        EventRegistration.student_id == payload.student_id
    ).first()
    
    if not att:
        raise HTTPException(status_code=404, detail="Student attendance record not found for this event")
        
    att.attendance = payload.status
    db.commit()
    
    return {"message": "Attendance updated"}

@app.patch("/faculty/events/{event_id}/attendance/bulk")
def bulk_update_event_attendance(
    event_id: int,
    payload: BulkEventAttendanceUpdate,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "faculty":
        raise HTTPException(status_code=403, detail="Faculty only")

    event = db.query(Event).filter(Event.id == event_id, Event.created_by == current_user["user_id"]).first()
    if not event:
        raise HTTPException(status_code=404, detail="Event not found")

    for record in payload.students:
        att = db.query(EventRegistration).filter(
            EventRegistration.event_id == event_id,
            EventRegistration.student_id == record.student_id
        ).first()
        if att:
            att.attendance = record.status
            
    db.commit()
    return {"message": "Bulk attendance updated"}

@app.patch("/faculty/events/result")
def update_event_result(
    payload: EventResultUpdate,
    event_id: int = Query(...),
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "faculty":
        raise HTTPException(status_code=403, detail="Faculty only")

    event = db.query(Event).filter(Event.id == event_id, Event.created_by == current_user["user_id"]).first()
    if not event:
        raise HTTPException(status_code=404, detail="Event not found or unauthorized")

    att = db.query(EventRegistration).filter(
        EventRegistration.event_id == event_id,
        EventRegistration.student_id == payload.student_id
    ).first()

    if not att:
        raise HTTPException(status_code=404, detail="Student attendance record not found")
    
    if att.attendance != "present":
        raise HTTPException(status_code=400, detail="Cannot assign result to absent student")

    att.result = payload.result
    db.commit()

    return {"message": "Result updated successfully"}

# ==========================================
# STUDENT EVENTS API
# ==========================================

@app.get("/student/events", response_model=List[StudentEventResponse])
def get_student_events(
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "student":
        raise HTTPException(status_code=403, detail="Student only")

    student = db.query(Student).filter(Student.student_id == current_user["user_id"]).first()
    if not student:
        raise HTTPException(status_code=404, detail="Student record not found")

    user_data = db.query(User).filter(User.user_id == student.student_id).first()
    if not user_data:
        raise HTTPException(status_code=404, detail="User record not found")

    # Fetch events targeted to this student's year/section
    # Not filtered by department to allow global events
    events = (
        db.query(Event)
        .filter(
            or_(Event.year == "All", Event.year == str(student.year)),
            or_(Event.section == "All", Event.section == student.section)
        )
        .order_by(Event.event_date.desc())
        .all()
    )

    results = []
    today = date.today()
    for ev in events:
        resp = StudentEventResponse.from_orm(ev)
        
        # Adjust dynamic status accurately
        if ev.event_date > today:
            resp.status = "Upcoming"
        elif ev.event_date == today:
            resp.status = "Ongoing"
        else:
            resp.status = "Completed"

        # Check registration specifically for this student
        reg = db.query(EventRegistration).filter(
            EventRegistration.event_id == ev.id,
            EventRegistration.student_id == student.student_id
        ).first()
        
        resp.is_registered = bool(reg)
        if reg:
            resp.attendance_status = reg.attendance
            resp.result = reg.result

        results.append(resp)

    return results

@app.post("/student/events/register")
def register_student_event(
    payload: EventRegistrationRequest,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "student":
        raise HTTPException(status_code=403, detail="Student only")

    event = db.query(Event).filter(Event.id == payload.event_id).first()
    if not event:
        raise HTTPException(status_code=404, detail="Event not found")

    # Check for duplicate
    existing_reg = db.query(EventRegistration).filter(
        EventRegistration.event_id == event.id,
        EventRegistration.student_id == current_user["user_id"]
    ).first()

    if existing_reg:
        raise HTTPException(status_code=400, detail="Already registered")
        
    # Validation checks
    if event.registration_deadline and datetime.now() > event.registration_deadline:
        raise HTTPException(status_code=400, detail="Registration deadline has passed")
        
    if event.max_participants is not None:
        current_count = db.query(EventRegistration).filter(EventRegistration.event_id == event.id).count()
        if current_count >= event.max_participants:
            raise HTTPException(status_code=400, detail="Event has reached maximum capacity")

    new_reg = EventRegistration(
        event_id=event.id,
        student_id=current_user["user_id"]
    )
    db.add(new_reg)
    db.commit()

    return {"message": "Successfully registered for event"}

@app.post("/faculty/events/{event_id}/alert")
def send_event_alert(
    event_id: int,
    payload: EventAlertRequest,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "faculty":
        raise HTTPException(status_code=403, detail="Faculty only")

    event = db.query(Event).filter(Event.id == event_id, Event.created_by == current_user["user_id"]).first()
    if not event:
        raise HTTPException(status_code=404, detail="Event not found")

    query = db.query(EventRegistration.student_id).filter(EventRegistration.event_id == event_id)
    if payload.target == "present":
        query = query.filter(EventRegistration.attendance == "present")
    elif payload.target == "absent":
        query = query.filter(EventRegistration.attendance == "absent")
        
    students = query.all()
    target_ids = [s[0] for s in students]
    
    if not target_ids:
        return {"message": "No students found for the given target"}
        
    title = f"Alert: {event.title}"

    new_alerts = []
    for sid in target_ids:
        new_alert = Alert(
            title=title,
            message=payload.message,
            type=payload.type,
            target_role="student",
            target_type="individual",
            student_id=sid,
            faculty_id=current_user["user_id"]
        )
        new_alerts.append(new_alert)
        
    db.add_all(new_alerts)
    db.flush()
    
    recipients = [
        AlertRecipient(alert_id=al.id, user_id=al.student_id, is_read=False)
        for al in new_alerts
    ]
    
    db.add_all(recipients)
    db.commit()

    return {"message": "Alerts sent successfully", "sent_count": len(target_ids)}

@app.post("/faculty/events/{event_id}/reminder")
def remind_absent_students(
    event_id: int,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "faculty":
        raise HTTPException(status_code=403, detail="Faculty only")

    event = db.query(Event).filter(Event.id == event_id, Event.created_by == current_user["user_id"]).first()
    if not event:
        raise HTTPException(status_code=404, detail="Event not found")

    students = db.query(EventRegistration.student_id).filter(
        EventRegistration.event_id == event_id,
        EventRegistration.attendance == "absent"
    ).all()
    
    target_ids = [s[0] for s in students]
    if not target_ids:
        return {"message": "No absent students to remind"}

    title = f"Reminder: {event.title}"
    message = f"Reminder: Please make sure to attend the event '{event.title}' scheduled on {event.event_date}."

    new_alerts = []
    for sid in target_ids:
        new_alert = Alert(
            title=title,
            message=message,
            type="reminder",
            target_role="student",
            target_type="individual",
            student_id=sid,
            faculty_id=current_user["user_id"]
        )
        new_alerts.append(new_alert)
        
    db.add_all(new_alerts)
    db.flush()
    
    recipients = [
        AlertRecipient(alert_id=al.id, user_id=al.student_id, is_read=False)
        for al in new_alerts
    ]
    
    db.add_all(recipients)
    db.commit()

    return {"message": "Reminders sent successfully", "sent_count": len(target_ids)}

# ==========================================
# EXTERNAL EVENT SUBMISSIONS API
# ==========================================

# Directory for external event achievement uploads
UPLOAD_DIR_EXTERNAL = "uploads/external_events"
os.makedirs(UPLOAD_DIR_EXTERNAL, exist_ok=True)

@app.post("/student/events/external-submit", response_model=ExternalEventSubmissionResponse)
def submit_external_achievement(
    event_name: str = Form(...),
    organizer: Optional[str] = Form(None),
    event_date: date = Form(...),
    achievement_type: Optional[str] = Form(None),
    position: Optional[str] = Form(None),
    certificate_file: Optional[UploadFile] = File(None),
    proof_file: Optional[UploadFile] = File(None),
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "student":
        raise HTTPException(status_code=403, detail="Student only")
        
    cert_path = None
    if certificate_file:
        file_ext = certificate_file.filename.split(".")[-1]
        unique_name = f"cert_{uuid.uuid4()}.{file_ext}"
        cert_path = os.path.join(UPLOAD_DIR_EXTERNAL, unique_name)
        with open(cert_path, "wb") as buffer:
            shutil.copyfileobj(certificate_file.file, buffer)

    proof_path = None
    if proof_file:
        file_ext = proof_file.filename.split(".")[-1]
        unique_name = f"proof_{uuid.uuid4()}.{file_ext}"
        proof_path = os.path.join(UPLOAD_DIR_EXTERNAL, unique_name)
        with open(proof_path, "wb") as buffer:
            shutil.copyfileobj(proof_file.file, buffer)

    new_sub = ExternalEventSubmission(
        student_id=current_user["user_id"],
        event_name=event_name,
        organizer=organizer,
        event_date=event_date,
        achievement_type=achievement_type,
        position=position,
        certificate_file=cert_path,
        proof_file=proof_path,
        status="pending"
    )
    
    db.add(new_sub)
    db.commit()
    db.refresh(new_sub)
    
    return new_sub

@app.get("/faculty/external-submissions", response_model=List[FacultyExternalSubmissionDetail])
def get_external_submissions(
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "faculty":
        raise HTTPException(status_code=403, detail="Faculty only")
        
    # Faculty can see pending achievements of students
    submissions = (
        db.query(ExternalEventSubmission, User, Student)
        .join(Student, ExternalEventSubmission.student_id == Student.student_id)
        .join(User, Student.student_id == User.user_id)
        .filter(ExternalEventSubmission.status == "pending")
        .order_by(ExternalEventSubmission.submitted_at.desc())
        .all()
    )
    
    results = []
    for sub, usr, st in submissions:
        resp = FacultyExternalSubmissionDetail.from_orm(sub)
        resp.student_name = usr.name
        resp.student_roll_no = st.roll_no
        results.append(resp)
        
    return results

@app.patch("/faculty/external-submissions/{sub_id}/status")
def update_external_submission_status(
    sub_id: int,
    status: str = Query(...), # approved or rejected
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "faculty":
        raise HTTPException(status_code=403, detail="Faculty only")
        
    if status not in ["approved", "rejected"]:
        raise HTTPException(status_code=400, detail="Invalid status")

    sub = db.query(ExternalEventSubmission).filter(ExternalEventSubmission.id == sub_id).first()
    if not sub:
        raise HTTPException(status_code=404, detail="Submission not found")
        
    sub.status = status
    sub.faculty_reviewed_by = current_user["user_id"]
    
    if status == "approved":
        student_profile = db.query(Student).filter(Student.student_id == sub.student_id).first()
        if student_profile:
            # Update student certificate count or list
            cert_entry = f"{sub.event_name} ({sub.achievement_type or 'Achievement'})"
            if student_profile.certificates:
                student_profile.certificates += f", {cert_entry}"
            else:
                student_profile.certificates = cert_entry
                
    db.commit()
    
    return {"message": f"External submission {status} successfully"}


#==========================================
# FACULTY MARKS UPLOAD API
#==========================================

@app.post("/faculty/upload-marks")
def upload_marks(
    data: MarksUpload,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):

    for item in data.marks:

        existing = db.query(Mark).filter(
            Mark.student_id == item.student_id,
            Mark.subject == data.subject,
            Mark.exam_type == data.exam,
            Mark.year == data.year,
            Mark.section == data.section
        ).first()

        if existing:
            existing.marks = item.marks
        else:
            new_mark = Mark(
                student_id=item.student_id,
                subject=data.subject,
                exam_type=data.exam,
                marks=item.marks,
                year=data.year,
                section=data.section,
                faculty_id=current_user.id
            )

            db.add(new_mark)

    db.commit()

    return {"message": "Marks uploaded successfully"}


#==========================================
# GET MARKS FOR SELECTED SUBJECT & EXAM
#==========================================

@app.get("/faculty/marks")
def get_marks(
    year: str,
    section: str,
    subject_id: int,
    exam: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):

    students = db.query(Student).filter(
        Student.year == year,
        Student.section == section
    ).all()

    result = []

    for s in students:

        mark = db.query(Mark).filter(
            Mark.student_id == s.student_id,
            Mark.subject == str(subject_id),
            Mark.exam_type == exam
        ).first()

        result.append({
            "student_id": s.student_id,
            "name": db.query(User).filter(User.user_id == s.student_id).first().name,
            "roll_no": s.roll_no,
            "marks": mark.marks if mark else None
        })

    return result


# ==========================================
# DOWNLOAD MARKS TEMPLATE
# ==========================================

from fastapi.responses import FileResponse  # type: ignore
import pandas as pd  # type: ignore


@app.get("/faculty/marks/template")
def download_marks_template(
    year: int = Query(...),
    section: str = Query(...),
    subject_id: int = Query(...),
    db: Session = Depends(get_db)
):
    # Fetch students for the given year and section
    students = db.query(Student, User).join(User, Student.student_id == User.user_id).filter(
        Student.year == str(year),
        Student.section == section,
        User.is_deleted == False
    ).order_by(Student.roll_no.asc()).all()

    data = []
    for s, u in students:
            data.append({
                "Register Number": s.roll_no,
                "Student Name": u.name,
                "Assignment 1": "",
                "Assignment 2": "",
                "Assignment 3": "",
                "Assignment 4": "",
                "Assignment 5": "",
                "Assignment Total (Scaled to 10)": "",
                "Mid 1": "",
                "Mid 1 (Scaled to 20)": "",
                "Mid 2": "",
                "Mid 2 (Scaled to 20)": "",
                "Semester": "",
                "Semester (Scaled)": "",
                "Total Marks": "",
                "SGPA": "",
                "CGPA": ""
            })

    df = pd.DataFrame(data)
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False)
    output.seek(0)

    filename = f"marks_template_year{year}_section{section}.xlsx"
    return StreamingResponse(output, media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', headers={"Content-Disposition": f"attachment; filename={filename}"})

# ==========================================
# UPLOAD MARKS VIA EXCEL
# ==========================================

@app.post("/faculty/marks/upload-excel")
async def upload_marks_excel(
    file: UploadFile = File(...),
    subject: str = "",
    year: str = "",
    section: str = "",
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):

    df = pd.read_excel(file.file)

    for _, row in df.iterrows():

        roll_no = row["RollNo"]

        student = db.query(Student).filter(Student.roll_no == roll_no).first()

        if not student:
            continue

        student_id = student.student_id

        for column in df.columns:

            if column in ["RollNo", "StudentName"]:
                continue

            value = row[column]

            if pd.isna(value):
                continue

            new_mark = Mark(
                student_id=student_id,
                subject=subject,
                exam_type=column,
                marks=int(value),
                year=year,
                section=section,
                faculty_id=current_user.id
            )

            db.add(new_mark)

    db.commit()

    return {"message": "Marks uploaded successfully"}

from pydantic import BaseModel  # type: ignore

class ApplyScalingRequest(BaseModel):
    year: int
    section: str
    subject_id: int

def process_scaling(year, section, subject_id, faculty_id, db, action_type):
    from models import ScaledMark, ScalingLog, Student, Mark  # type: ignore
    students = db.query(Student).filter(
        Student.year == year,
        Student.section == section
    ).all()
    
    if not students:
        return {"error": "No students found"}

    current_scaled = db.query(ScaledMark).filter(
        ScaledMark.subject_id == subject_id,
        ScaledMark.year == year,
        ScaledMark.section == section
    ).all()

    snapshot_data = []
    for sm in current_scaled:
        snapshot_data.append({
            "student_id": sm.student_id,
            "assignment_scaled": float(sm.assignment_scaled) if sm.assignment_scaled is not None else None,
            "mid1_scaled": float(sm.mid1_scaled) if sm.mid1_scaled is not None else None,
            "mid2_scaled": float(sm.mid2_scaled) if sm.mid2_scaled is not None else None,
            "mid_combined": float(sm.mid_combined) if sm.mid_combined is not None else None,
            "internal_total": float(sm.internal_total) if sm.internal_total is not None else None,
            "semester_marks": float(sm.semester_marks) if sm.semester_marks is not None else None,
            "final_total": float(sm.final_total) if sm.final_total is not None else None
        })

    marks = db.query(Mark).filter(
        Mark.subject_id == subject_id,
        Mark.year == year,
        Mark.section == section
    ).all()
    
    student_marks = {}
    for m in marks:
        if m.student_id not in student_marks:
            student_marks[m.student_id] = []
        student_marks[m.student_id].append(m)

    for student in students:
        sid = student.student_id
        sm_record = db.query(ScaledMark).filter_by(
            student_id=sid, subject_id=subject_id
        ).first()

        if not sm_record:
            sm_record = ScaledMark(
                student_id=sid,
                subject_id=subject_id,
                year=year,
                section=section
            )
            db.add(sm_record)

        s_marks = student_marks.get(sid, [])
        assignments = []
        mid1 = None
        mid2 = None
        sem = None
        
        for m in s_marks:
            if m.marks is None: continue  # type: ignore
            v = float(m.marks)  # type: ignore
            if "Assignment" in m.exam: assignments.append(v)  # type: ignore
            elif m.exam == "Mid-1": mid1 = v  # type: ignore
            elif m.exam == "Mid-2": mid2 = v  # type: ignore
            elif m.exam == "Semester": sem = v  # type: ignore
            
        ass_total = sum(assignments)
        sm_record.assignment_scaled = round((ass_total / 50.0) * 10, 2)  # type: ignore
        sm_record.mid1_scaled = round((mid1 / 30.0) * 20, 2) if mid1 is not None else 0.0  # type: ignore
        sm_record.mid2_scaled = round((mid2 / 30.0) * 20, 2) if mid2 is not None else 0.0  # type: ignore
        
        m_c = round(((sm_record.mid1_scaled + sm_record.mid2_scaled) / 40.0) * 20, 2)  # type: ignore
        sm_record.mid_combined = m_c  # type: ignore
        
        i_t = round(sm_record.assignment_scaled + m_c, 2)  # type: ignore
        sm_record.internal_total = i_t  # type: ignore
        
        sm_record.semester_marks = sem if sem is not None else 0.0  # type: ignore
        sm_record.final_total = round(i_t + sm_record.semester_marks, 2)  # type: ignore

    db.commit()

    log = ScalingLog(
        faculty_id=faculty_id,
        action_type=action_type,
        year=year,
        section=section,
        subject_id=subject_id,
        snapshot_data=snapshot_data
    )
    db.add(log)
    db.commit()

    return {"message": f"Successfully executed '{action_type}' for scaling records."}

@app.get("/faculty/overview")
def get_faculty_overview(
    year: int = Query(...),
    section: str = Query(...),
    subject_id: int = Query(...),
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "faculty":
        raise HTTPException(status_code=403, detail="Faculty only")

    from models import Student, Mark, User, Attendance, Subject, FacultySubject  # type: ignore
    from sqlalchemy import or_, func  # type: ignore
    
    faculty_id = current_user["user_id"]
    
    # ===============================
    # STEP 1: GET ALL SUBJECTS OF FACULTY
    # ===============================
    assigned_subjects = db.query(FacultySubject).filter(
        FacultySubject.faculty_id == faculty_id,
        FacultySubject.is_active == True
    ).all()
    
    total_subjects = len({s.subject_id for s in assigned_subjects})
    
    # ===============================
    # STEP 2: TOTAL STUDENTS (ALL CLASSES)
    # ===============================
    # Count DISTINCT students across all faculty-assigned subjects
    total_students = 0
    if assigned_subjects:
        subject_ids = [s.subject_id for s in assigned_subjects]
        total_students = db.query(func.count(func.distinct(Mark.student_id))).filter(
            Mark.subject_id.in_(subject_ids),
            Mark.faculty_id == faculty_id
        ).scalar() or 0
    
    faculty_scope = {
        "total_students": total_students,
        "total_subjects": total_subjects
    }

    # ===============================
    # STEP 3: GET SUBJECT NAME
    # ===============================
    subject_obj = db.query(Subject).filter(Subject.subject_id == subject_id).first()
    subject_name = subject_obj.subject_name if subject_obj else "Subject"

    # ===============================
    # STEP 4: CURRENT CLASS STUDENTS (SELECTED SUBJECT)
    # ===============================
    students_query = db.query(Student.student_id, User.name).join(
        User, Student.student_id == User.user_id
    ).filter(
        Student.year == year,
        Student.section == section
    ).all()
    
    total_class_students = len(students_query)
    class_stats = {
        "current_class_students": total_class_students,
        "subject_name": subject_name
    }
    
    if not students_query:
        return {
            "class_stats": class_stats,
            "faculty_scope": faculty_scope,
            "kpis": {
                "class_avg": 0,
                "pass_rate": 0,
                "topper": None,
                "at_risk_count": 0
            },
            "metrics": { "mid1": [], "mid2": [], "total": [], "assignment": [] },
            "attendance": [],
            "marks_risk_students": [],
            "attendance_risk_students": []
        }
    
    student_ids = [s.student_id for s in students_query]
    student_names = {s.student_id: s.name for s in students_query}
    
    # ===============================
    # STEP 5: FETCH MARKS FOR SELECTED SUBJECT
    # ===============================
    marks_query = db.query(Mark.student_id, Mark.exam, Mark.marks).filter(
        Mark.student_id.in_(student_ids),
        Mark.subject_id == subject_id
    ).all()

    metrics = { "mid1": [], "mid2": [], "total": [], "assignment": [] }
    student_metrics = { sid: {"mid1": 0.0, "mid2": 0.0, "assignment": 0.0, "total": 0.0} for sid in student_ids}
    
    # Store individual exam marks for risk analysis
    exam_marks = { sid: {"mid1": None, "mid2": None, "semester": None} for sid in student_ids}
    
    for m in marks_query:
        if m.marks is None: continue
        val = float(m.marks)
        exam = str(m.exam).strip().lower()
        sid = m.student_id
        
        student_metrics[sid]["total"] += val
        
        if "mid-1" in exam or "mid 1" in exam or "mid1" in exam:
            student_metrics[sid]["mid1"] += val
            exam_marks[sid]["mid1"] = val
        elif "mid-2" in exam or "mid 2" in exam or "mid2" in exam:
            student_metrics[sid]["mid2"] += val
            exam_marks[sid]["mid2"] = val
        elif "semester" in exam:
            exam_marks[sid]["semester"] = val
        elif "assignment" in exam:
            student_metrics[sid]["assignment"] += val

    for sid, data in student_metrics.items():
        name = student_names[sid]
        metrics["mid1"].append({"name": name, "marks": round(data["mid1"], 2)})  # type: ignore
        metrics["mid2"].append({"name": name, "marks": round(data["mid2"], 2)})  # type: ignore
        metrics["assignment"].append({"name": name, "marks": round(data["assignment"], 2)})  # type: ignore
        metrics["total"].append({"name": name, "marks": round(data["total"], 2)})  # type: ignore
        
    if sum(d["mid1"] for d in student_metrics.values()) == 0: metrics["mid1"] = []
    if sum(d["mid2"] for d in student_metrics.values()) == 0: metrics["mid2"] = []
    if sum(d["assignment"] for d in student_metrics.values()) == 0: metrics["assignment"] = []
    if sum(d["total"] for d in student_metrics.values()) == 0: metrics["total"] = []

    totals_list = [d["total"] for d in student_metrics.values()]
    class_avg = round(sum(totals_list) / len(totals_list), 2) if totals_list else 0.0  # type: ignore
    passed = sum(1 for t in totals_list if t >= 12)
    pass_rate = round((passed / total_class_students) * 100, 2) if total_class_students > 0 else 0.0  # type: ignore
    at_risk_count = sum(1 for t in totals_list if t < 15)
    
    topper = None
    if student_metrics:
        max_sid = max(student_metrics.keys(), key=lambda k: student_metrics[k]["total"])
        if student_metrics[max_sid]["total"] > 0:
            topper = {"name": student_names[max_sid], "marks": round(student_metrics[max_sid]["total"], 2)}  # type: ignore

    # ===============================
    # STEP 6: MARKS RISK ANALYSIS
    # ===============================
    MARKS_RISK_THRESHOLD = 15
    marks_risk_students = []
    
    for sid in student_ids:
        marks_list = []
        for exam_name, mark in exam_marks[sid].items():
            if mark is not None and mark > 0:
                marks_list.append((exam_name.replace("_", "-").upper(), mark))
        
        if marks_list:
            exam_name, min_mark = min(marks_list, key=lambda x: x[1])
            if min_mark < MARKS_RISK_THRESHOLD:
                marks_risk_students.append({
                    "name": student_names[sid],
                    "value": round(min_mark, 2),
                    "exam": exam_name
                })
    
    # Sort by mark ascending (lowest first)
    marks_risk_students.sort(key=lambda x: x["value"])
    
    # ===============================
    # STEP 7: ATTENDANCE RISK ANALYSIS
    # ===============================
    att_query = db.query(Attendance.student_id, Attendance.status).filter(
        Attendance.student_id.in_(student_ids),
        Attendance.subject_id == subject_id
    ).all()
    
    att_counts = { sid: {"present": 0, "total": 0} for sid in student_ids }
    for a in att_query:
        att_counts[a.student_id]["total"] += 1
        if a.status:
            att_counts[a.student_id]["present"] += 1
            
    attendance_data = []
    threshold = get_setting("attendance_threshold") or 75
    print("Using attendance threshold:", threshold)
    attendance_risk_students = []
    
    for sid, counts in att_counts.items():
        pct = round((counts["present"]/counts["total"])*100, 2) if counts["total"] > 0 else 100.0  # type: ignore
        attendance_data.append({"name": student_names[sid], "percentage": pct})
        
        if pct < threshold:
            attendance_risk_students.append({
                "name": student_names[sid],
                "value": pct
            })
    
    # Sort by percentage ascending (lowest first)
    attendance_risk_students.sort(key=lambda x: x["value"])

    # ===============================
    # STEP 8: RETURN RESPONSE
    # ===============================
    return {
        "class_stats": class_stats,
        "faculty_scope": faculty_scope,
        "kpis": {
            "class_avg": class_avg,
            "pass_rate": pass_rate,
            "topper": topper,
            "at_risk_count": at_risk_count
        },
        "metrics": metrics,
        "attendance": attendance_data,
        "marks_risk_students": marks_risk_students,
        "attendance_risk_students": attendance_risk_students
    }

@app.get("/faculty/insights-data")
def get_faculty_insights_data(
    year: Optional[int] = Query(None),
    section: Optional[str] = Query(None),
    subject_id: Optional[int] = Query(None),
    timeRange: Optional[str] = Query(None),
    time_range: Optional[str] = Query(None),
    trend_view: Optional[str] = Query(None),
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "faculty":
        raise HTTPException(status_code=403, detail="Faculty only")

    from models import Student, User, FacultySubject, Subject, Attendance, Assignment, AssignmentSubmission, Mark  # type: ignore

    faculty_id = current_user["user_id"]

    faculty_filters = [FacultySubject.faculty_id == faculty_id, FacultySubject.is_active == True]
    if subject_id is not None:
        faculty_filters.append(FacultySubject.subject_id == subject_id)
    if year is not None:
        faculty_filters.append(FacultySubject.year == year)
    if section is not None:
        faculty_filters.append(FacultySubject.section == section)

    assigned_subjects = db.query(FacultySubject).filter(*faculty_filters).all()

    if not assigned_subjects:
        return {
            "filters": {"year": year, "section": section, "subject_id": subject_id},
            "faculty_scope": {"total_subjects": 0, "assigned_classes": 0, "total_students": 0},
            "students": [],
            "predictions": {
                "future_risk_students": 0,
                "expected_attendance": 0.0,
                "expected_avg_marks": 0.0,
                "confidence": "LOW",
                "trend_direction": "stable",
                "confidence_reason": "based on 0 available academic data points"
            },
            "attendance_trend": [],
            "trend_summary": {"direction": "stable", "change_percent": 0.0, "status": "below_threshold"},
            "attendance_chart_annotation": "",
            "mid_analysis": {
                "improved": 0, "declined": 0, "stable": 0, "total": 0,
                "avg_mid1": 0.0, "avg_mid2": 0.0, "trend": "stable"
            },
            "top_risks": [],
            "insights": [{
                "title": "No cohort assigned",
                "message": "No subjects are assigned to this faculty account for the current filters.",
                "priority": "medium",
                "action": "Contact admin to map your subjects, then reload this page.",
                "severity": "medium"
            }, {
                "title": "Next step",
                "message": "Once subjects appear, attendance and Mid trends will populate automatically.",
                "priority": "low",
                "action": "Select subject and time range after assignment is complete.",
                "severity": "low"
            }],
            "trendInsight": "No sufficient data available for selected filters.",
            "recommended_actions": [
                "Contact admin to map your subjects, then reload this page.",
                "Select subject and time range after assignment is complete."
            ],
            "priority_summary": {
                "count": 0,
                "breakdown": {"low_attendance": 0, "low_marks": 0, "declining": 0}
            },
            "attendance_summary": {
                "overall_percentage": 0.0, "present_count": 0, "absent_count": 0, "total_records": 0, "by_subject": [],
                "average": 0.0, "threshold": 75, "status": "critical", "trend": "stable",
                "message": "No attendance records available for selected filters."
            },
            "assignment_summary": {"total_assignments": 0, "submitted_count": 0, "late_submissions": 0, "submission_rate": 0.0},
            "marks_summary": {"avg_mid1": 0.0, "avg_mid2": 0.0, "avg_assignment": 0.0, "avg_semester": 0.0, "avg_total": 0.0},
            "risk_summary": {"attendance_risk_count": 0, "marks_risk_count": 0, "assignment_engagement_risk_count": 0},
            "top_students": [],
            "weakest_subject": {"name": "None", "trend": "stable", "reason": "", "reason_lines": []},
            "mid_comparison": [],
            "mid_comparison_summary": "",
            "marks_trend_summary": "",
            "risk_distribution": {"low": 0, "medium": 0, "high": 0},
            "attendance_trend_direction": "stable",
            "recommendations": [],
            "alerts": []
        }

    subject_ids = list({fs.subject_id for fs in assigned_subjects})
    class_pairs = {(fs.year, fs.section) for fs in assigned_subjects}
    year_options = sorted({fs.year for fs in assigned_subjects})
    section_options = sorted({fs.section for fs in assigned_subjects})

    student_query = db.query(Student.student_id, Student.year, Student.section, User.name).join(
        User, Student.student_id == User.user_id
    )

    # Subject-driven student selection to avoid empty attendance trends
    derived_student_filter = False
    if subject_id is not None:
        attendance_student_ids = db.query(Attendance.student_id).filter(
            Attendance.faculty_id == faculty_id,
            Attendance.subject_id == subject_id
        ).distinct().all()
        attendance_student_ids = [row[0] for row in attendance_student_ids if row and row[0] is not None]

        if attendance_student_ids:
            student_query = student_query.filter(Student.student_id.in_(attendance_student_ids))
            derived_student_filter = True

    # Fallback: preserve legacy behavior when we couldn't derive students from attendance
    if not derived_student_filter:
        if year is not None and section is not None:
            student_query = student_query.filter(Student.year == year, Student.section == section)
        elif year is not None:
            student_query = student_query.filter(Student.year == year, Student.section.in_(section_options))
        elif section is not None:
            student_query = student_query.filter(Student.section == section, Student.year.in_(year_options))
        else:
            student_query = student_query.filter(Student.year.in_(year_options), Student.section.in_(section_options))

    students = student_query.all()
    student_ids = [s.student_id for s in students]
    student_names = {s.student_id: s.name for s in students}
    total_students = len(student_ids)

    subject_map = {
        s.subject_id: s.subject_name
        for s in db.query(Subject).filter(Subject.subject_id.in_(subject_ids)).all()
    }
    selected_subject_name = None
    if subject_id is not None:
        selected_subject_name = subject_map.get(subject_id, "Unknown Subject")

    attendance_query = db.query(Attendance).filter(
        Attendance.faculty_id == faculty_id,
        Attendance.student_id.in_(student_ids)
    )
    if subject_id is not None:
        attendance_query = attendance_query.filter(Attendance.subject_id == subject_id)

    attendance_records = attendance_query.all()

    # Attendance trend for charts: actual points only {label, value} (no forecast line)
    attendance_trend = []
    attendance_values = []
    try:
        today = datetime.utcnow().date()
        tr = (timeRange or time_range or "").strip().lower()

        # Default to "semester" bucket if nothing provided
        if not tr:
            tr = "semester"

        # Build deterministic buckets with carry-forward values.
        view = (trend_view or "").strip().lower()
        if view not in {"days", "weeks", "months"}:
            view = "days"

        # Time window controls which real records are included in bucket values.
        if tr in {"last7", "last_7_days", "last 7 days"}:
            window_start = today - timedelta(days=6)
        elif tr in {"last30", "last_30_days", "last 30 days"}:
            window_start = today - timedelta(days=29)
        else:
            window_start = today - timedelta(days=83)

        def in_window(d):
            return window_start <= d <= today

        attendance_trend = []
        last_val = 75

        if view == "days":
            date_list = [today - timedelta(days=6) + timedelta(days=i) for i in range(7)]
            for d in date_list:
                present = 0
                total = 0
                for r in attendance_records:
                    rd = r.attendance_date
                    if rd == d and in_window(rd):
                        total += 1
                        if r.status:
                            present += 1
                if total > 0:
                    last_val = round((present / total) * 100, 2)
                attendance_trend.append({"label": d.strftime("%a"), "value": last_val})

        elif view == "weeks":
            week_start = today - timedelta(days=27)
            for i in range(4):
                start_d = week_start + timedelta(days=i * 7)
                end_d = start_d + timedelta(days=6)
                present = 0
                total = 0
                for r in attendance_records:
                    rd = r.attendance_date
                    if start_d <= rd <= end_d and in_window(rd):
                        total += 1
                        if r.status:
                            present += 1
                if total > 0:
                    last_val = round((present / total) * 100, 2)
                attendance_trend.append({"label": f"Week {i + 1}", "value": last_val})

        else:
            # months => last 3 calendar months ending current month
            def month_start(d, offset_months):
                y = d.year
                m = d.month + offset_months
                y += (m - 1) // 12
                m = (m - 1) % 12 + 1
                return datetime(y, m, 1).date()

            m0 = month_start(today, -2)
            m1 = month_start(today, -1)
            m2 = month_start(today, 0)
            bucket_starts = [m0, m1, m2]
            bucket_ends = [m1 - timedelta(days=1), m2 - timedelta(days=1), today]

            for i in range(3):
                start_d = bucket_starts[i]
                end_d = bucket_ends[i]
                present = 0
                total = 0
                for r in attendance_records:
                    rd = r.attendance_date
                    if start_d <= rd <= end_d and in_window(rd):
                        total += 1
                        if r.status:
                            present += 1
                if total > 0:
                    last_val = round((present / total) * 100, 2)
                # Keep month labels consistent for UI readability.
                month_labels = ["Jan", "Feb", "Mar"]
                attendance_trend.append({"label": month_labels[i], "value": last_val})

    except Exception:
        attendance_trend = []

    # Normalize trend to [{label, value}] and build attendance_values (no prediction / no "Next N" labels)
    try:
        print("SUBJECT:", subject_id)
        print("FETCHED STUDENTS:", len(students))

        attendance_values = []
        trend_labels = []
        for x in attendance_trend:
            if isinstance(x, dict):
                v = x.get("value")
                if v is not None:
                    try:
                        attendance_values.append(float(v))
                        trend_labels.append(x.get("label"))
                    except Exception:
                        pass

        if not attendance_values:
            trend_labels = []
            attendance_values = []

        attendance_trend = [
            {"label": trend_labels[i] if i < len(trend_labels) else f"P{i + 1}", "value": attendance_values[i]}
            for i in range(len(attendance_values))
        ]

        print("ATTENDANCE VALUES:", attendance_values)
        print("ATTENDANCE TREND (actual only):", attendance_trend)
    except Exception:
        attendance_values = []
        attendance_trend = []

    # Attendance direction + drop (absolute points + percent-of-start for summaries)
    attendance_trend_direction = "stable"
    attendance_drop_percentage = 0.0
    attendance_change_percent = 0.0
    chart_series_avg = round(sum(attendance_values) / len(attendance_values), 2) if attendance_values else 0.0
    if attendance_values:
        first_val = float(attendance_values[0])
        last_val = float(attendance_values[-1])
        if last_val > first_val:
            attendance_trend_direction = "improving"
            if first_val > 0:
                attendance_change_percent = round((last_val - first_val) / first_val * 100, 1)
        elif last_val < first_val:
            attendance_trend_direction = "declining"
            attendance_drop_percentage = round(first_val - last_val, 2)
            if first_val > 0:
                attendance_change_percent = round((first_val - last_val) / first_val * 100, 1)

    # Human-readable chart annotation (largest consecutive drop)
    attendance_chart_annotation = ""
    if attendance_values and len(attendance_values) >= 2:
        labels_for_trend = [x.get("label") for x in attendance_trend if isinstance(x, dict)]
        best_drop = 0.0
        best_after = None
        for i in range(1, len(attendance_values)):
            drop = float(attendance_values[i - 1]) - float(attendance_values[i])
            if drop > best_drop:
                best_drop = drop
                best_after = labels_for_trend[i - 1] if i - 1 < len(labels_for_trend) else None
        if best_drop >= 5 and best_after:
            attendance_chart_annotation = f"Attendance dropped sharply after {best_after}"
        elif attendance_trend_direction == "improving" and len(attendance_values) >= 2:
            attendance_chart_annotation = "Attendance is trending upward over this window"

    present_count = sum(1 for a in attendance_records if a.status)
    absent_count = len(attendance_records) - present_count
    total_attendance = len(attendance_records)
    overall_attendance_pct = round((present_count / total_attendance) * 100, 2) if total_attendance else 0.0

    attendance_by_student = {sid: {"present": 0, "total": 0} for sid in student_ids}
    attendance_by_subject = {}
    for a in attendance_records:
        attendance_by_student[a.student_id]["total"] += 1
        if a.status:
            attendance_by_student[a.student_id]["present"] += 1

        subject_stats = attendance_by_subject.setdefault(a.subject_id, {"present": 0, "total": 0})
        subject_stats["total"] += 1
        if a.status:
            subject_stats["present"] += 1

    attendance_subject_breakdown = [
        {
            "subject_id": sid,
            "subject_name": subject_map.get(sid, "Unknown Subject"),
            "present": data["present"],
            "total": data["total"],
            "percentage": round((data["present"] / data["total"]) * 100, 2) if data["total"] else 0.0
        }
        for sid, data in attendance_by_subject.items()
    ]

    attendance_threshold = get_setting("attendance_threshold") or 75
    cgpa_threshold = get_setting("cgpa_threshold") or 6.5
    try:
        _thr = float(attendance_threshold)
    except Exception:
        _thr = 75.0

    trend_summary_status = "below_threshold" if (
        chart_series_avg < _thr or (attendance_values and float(attendance_values[-1]) < _thr)
    ) else "ok"
    trend_summary = {
        "direction": attendance_trend_direction,
        "change_percent": attendance_change_percent,
        "status": trend_summary_status
    }

    attendance_risk_students = [
        {
            "student_id": sid,
            "name": student_names.get(sid, "Unknown"),
            "attendance_percentage": round((counts["present"] / counts["total"]) * 100, 2) if counts["total"] else 0.0
        }
        for sid, counts in attendance_by_student.items()
        if counts["total"] and ((counts["present"] / counts["total"]) * 100) < attendance_threshold
    ]

    assignment_query = db.query(Assignment).filter(
        Assignment.faculty_id == faculty_id,
        Assignment.is_active == True
    )
    if subject_id is not None:
        assignment_query = assignment_query.filter(Assignment.subject_id == subject_id)
    if year is not None:
        assignment_query = assignment_query.filter(Assignment.year == year)
    if section is not None:
        assignment_query = assignment_query.filter(Assignment.section == section)

    assignments = assignment_query.all()
    assignment_ids = [a.id for a in assignments]
    total_assignments = len(assignments)

    submission_count = 0
    late_submissions = 0
    submission_by_student = {sid: 0 for sid in student_ids}

    if assignment_ids:
        submission_records = db.query(AssignmentSubmission).filter(
            AssignmentSubmission.assignment_id.in_(assignment_ids),
            AssignmentSubmission.is_submitted == True
        ).all()
        submission_count = len(submission_records)
        late_submissions = sum(1 for s in submission_records if s.is_late)
        for sub in submission_records:
            if sub.student_id in submission_by_student:
                submission_by_student[sub.student_id] += 1

    total_possible_submissions = total_assignments * total_students
    submission_rate = round((submission_count / total_possible_submissions) * 100, 2) if total_possible_submissions else 0.0
    assignment_engagement_risk_count = sum(
        1 for sid in student_ids
        if total_assignments and submission_by_student.get(sid, 0) < total_assignments
    ) if total_assignments else 0

    marks_query = db.query(Mark).filter(
        Mark.faculty_id == faculty_id,
        Mark.student_id.in_(student_ids)
    )
    if subject_id is not None:
        marks_query = marks_query.filter(Mark.subject_id == subject_id)
    if year is not None:
        marks_query = marks_query.filter(Mark.year == year)
    if section is not None:
        marks_query = marks_query.filter(Mark.section == section)

    mark_records = marks_query.all()
    mark_totals = {"mid1": 0.0, "mid2": 0.0, "assignment": 0.0, "semester": 0.0, "total": 0.0}
    mark_counts = {"mid1": 0, "mid2": 0, "assignment": 0, "semester": 0, "total": 0}
    student_mark_aggregation = {sid: {"total": 0.0, "count": 0} for sid in student_ids}

    for m in mark_records:
        if m.marks is None:
            continue
        exam_name = str(m.exam or "").strip().lower()
        value = float(m.marks)

        if "mid-1" in exam_name or "mid1" in exam_name:
            mark_totals["mid1"] += value
            mark_counts["mid1"] += 1
        elif "mid-2" in exam_name or "mid2" in exam_name:
            mark_totals["mid2"] += value
            mark_counts["mid2"] += 1
        elif "assignment" in exam_name:
            mark_totals["assignment"] += value
            mark_counts["assignment"] += 1
        elif "semester" in exam_name:
            mark_totals["semester"] += value
            mark_counts["semester"] += 1

        mark_totals["total"] += value
        mark_counts["total"] += 1
        if m.student_id in student_mark_aggregation:
            student_mark_aggregation[m.student_id]["total"] += value
            student_mark_aggregation[m.student_id]["count"] += 1

    marks_summary = {
        "avg_mid1": round(mark_totals["mid1"] / mark_counts["mid1"], 2) if mark_counts["mid1"] else 0.0,
        "avg_mid2": round(mark_totals["mid2"] / mark_counts["mid2"], 2) if mark_counts["mid2"] else 0.0,
        "avg_assignment": round(mark_totals["assignment"] / mark_counts["assignment"], 2) if mark_counts["assignment"] else 0.0,
        "avg_semester": round(mark_totals["semester"] / mark_counts["semester"], 2) if mark_counts["semester"] else 0.0,
        "avg_total": round(mark_totals["total"] / mark_counts["total"], 2) if mark_counts["total"] else 0.0
    }

    # Mid 1 vs Mid 2 (per-student) for intelligence + charts
    student_mid_marks = {sid: {"mid1": None, "mid2": None} for sid in student_ids}
    for m in mark_records:
        if m.marks is None:
            continue
        sid = m.student_id
        if sid not in student_mid_marks:
            continue
        exam_name = str(m.exam or "").strip().lower()
        try:
            val = float(m.marks)
        except Exception:
            continue
        if "mid-1" in exam_name or "mid1" in exam_name:
            student_mid_marks[sid]["mid1"] = val
        elif "mid-2" in exam_name or "mid2" in exam_name:
            student_mid_marks[sid]["mid2"] = val

    mid_comparison = []
    improved_count = 0
    declined_count = 0
    same_count = 0
    mid_risk_count = 0
    for sid in student_ids:
        m1 = student_mid_marks.get(sid, {}).get("mid1")
        m2 = student_mid_marks.get(sid, {}).get("mid2")
        if isinstance(m1, (int, float)) and isinstance(m2, (int, float)):
            if m2 > m1:
                improved_count += 1
            elif m2 < m1:
                declined_count += 1
            else:
                same_count += 1
        if (isinstance(m1, (int, float)) and m1 < 15) or (isinstance(m2, (int, float)) and m2 < 15):
            mid_risk_count += 1
        mid_comparison.append({
            "name": student_names.get(sid, "Unknown"),
            "mid1": round(float(m1), 2) if isinstance(m1, (int, float)) else None,
            "mid2": round(float(m2), 2) if isinstance(m2, (int, float)) else None,
        })

    avg_mid1 = marks_summary.get("avg_mid1", 0.0)
    avg_mid2 = marks_summary.get("avg_mid2", 0.0)
    if avg_mid2 > avg_mid1:
        performance_trend = "improving"
    elif avg_mid2 < avg_mid1:
        performance_trend = "declining"
    else:
        performance_trend = "stable"

    mid_comparison_summary = f"Mid 2 performance improved for {improved_count} out of {max(improved_count + declined_count + same_count, 1)} students."
    marks_trend_summary = f"Mid performance trend is {performance_trend} (Avg Mid 1: {round(avg_mid1,1)}, Avg Mid 2: {round(avg_mid2,1)})."

    student_averages = []
    for sid, values in student_mark_aggregation.items():
        if values["count"]:
            student_averages.append({
                "student_id": sid,
                "name": student_names.get(sid, "Unknown"),
                "average_marks": round(values["total"] / values["count"], 2)
            })

    top_students = sorted(student_averages, key=lambda item: item["average_marks"], reverse=True)[:5]
    marks_risk_count = sum(1 for item in student_averages if item["average_marks"] < 15)

    # Build per-student attendance trend points (ordered recent percentages) for risk logic.
    student_attendance_daily = {sid: {} for sid in student_ids}
    for a in attendance_records:
        sid = a.student_id
        if sid not in student_attendance_daily:
            continue
        day_bucket = student_attendance_daily[sid].setdefault(a.attendance_date, {"present": 0, "total": 0})
        day_bucket["total"] += 1
        if a.status:
            day_bucket["present"] += 1

    student_attendance_trend = {}
    for sid, daily in student_attendance_daily.items():
        ordered_days = sorted(daily.keys())
        trend_vals = []
        for d in ordered_days:
            totals = daily[d]
            if totals["total"] > 0:
                trend_vals.append(round((totals["present"] / totals["total"]) * 100, 2))
        # Keep latest 5 points to represent recent trend.
        student_attendance_trend[sid] = trend_vals[-5:]

    student_rows = {}
    if student_ids:
        student_rows = {
            row.student_id: row
            for row in db.query(Student).filter(Student.student_id.in_(student_ids)).all()
        }

    # === ML PREDICTION & RISK BLOCK ===
    students_list = []
    
    for sid in student_ids:
        student_row = student_rows.get(sid)
        att_data = attendance_by_student.get(sid, {"present": 0, "total": 0})
        att_pct = round((att_data["present"] / att_data["total"]) * 100, 2) if att_data["total"] else 0.0
        marks_data = student_mark_aggregation.get(sid, {"total": 0.0, "count": 0})
        marks_avg = round(marks_data["total"] / marks_data["count"], 2) if marks_data["count"] else 0.0
        subs = submission_by_student.get(sid, 0)

        attendance_history = get_student_attendance_trend({"attendance_history": student_attendance_trend.get(sid, [])})
        
        mm = student_mid_marks.get(sid, {"mid1": None, "mid2": None})
        stu_obj = {
            "student_id": sid,
            "name": student_names.get(sid, "Unknown"),
            "attendance": att_pct,
            "attendance_trend": attendance_history,
            "attendance_history": attendance_history,
            "marks": marks_avg,
            "assignments": subs,
            "mid1": mm.get("mid1"),
            "mid2": mm.get("mid2"),
        }
        risk_payload = get_student_risk(
            student_id=sid,
            db=db,
            attendance_threshold=float(attendance_threshold),
            cgpa_threshold=float(cgpa_threshold),
        )
        risk_level = str(risk_payload.get("overall_risk") or "NO_DATA").upper()
        if risk_level == "NO_DATA":
            risk_level = "LOW"

        risk_score = 0.0
        if risk_payload.get("overall_risk") == "HIGH":
            risk_score = 85.0
        elif risk_payload.get("overall_risk") == "MEDIUM":
            risk_score = 60.0
        elif risk_payload.get("overall_risk") == "LOW":
            risk_score = 25.0

        stu_obj["risk"] = {
            "level": risk_level,
            "risk_score": risk_score,
            "reasons": risk_payload.get("reasons", []),
            "actions": risk_payload.get("actions", []),
            "has_valid_data": bool(risk_payload.get("has_valid_data")),
        }
        stu_obj["risk_level"] = risk_level
        stu_obj["risk_score"] = risk_score
        stu_obj["risk_breakdown"] = []
        stu_obj["reasons"] = risk_payload.get("reasons", [])
        stu_obj["actions"] = risk_payload.get("actions", [])

        previous_risk_score = None
        if student_row is not None:
            stored_previous = getattr(student_row, "previous_risk_score", None)
            if stored_previous is not None:
                try:
                    previous_risk_score = float(stored_previous)
                except (TypeError, ValueError):
                    previous_risk_score = None

            if previous_risk_score is None:
                numeric_history = [float(value) for value in attendance_history if isinstance(value, (int, float))]
                if len(numeric_history) >= 2:
                    previous_snapshot = dict(stu_obj)
                    previous_snapshot["attendance"] = numeric_history[-2]
                    try:
                        previous_risk_score = float(calculate_risk_score(previous_snapshot)[0])
                    except Exception:
                        previous_risk_score = None

        stu_obj["previous_risk_score"] = previous_risk_score
        stu_obj["risk_movement"] = calculate_risk_movement(stu_obj)
        stu_obj["attendance_trend_label"] = get_attendance_trend_label(attendance_history)
        stu_obj["intervention"] = {
            "status": getattr(student_row, "intervention_status", None) or "none",
            "type": getattr(student_row, "intervention_type", None),
            "last_updated": getattr(student_row, "intervention_last_updated", None).isoformat() if getattr(student_row, "intervention_last_updated", None) else None,
        }
            
        students_list.append(stu_obj)

    _risk_order = {"HIGH": 3, "MEDIUM": 2, "LOW": 1}
    students_list.sort(
        key=lambda s: (
            -(float(s.get("risk_score", 0) or 0)),
            -_risk_order.get(s.get("risk", {}).get("level"), 0),
            s.get("marks", 0) or 0,
        )
    )

    try:
        future_risk = predict_future_risk(students_list, {
            "attendance": attendance_threshold,
            "marks": 15,
            "assignment": total_assignments
        })
    except Exception:
        future_risk = 0
        
    try:
        expected_marks = forecast_performance(students_list)
    except Exception:
        expected_marks = 0.0
        
    data_points = len(student_ids) + len(mark_records)
    trend_data_points = len(attendance_values)
    confidence = "LOW"
    if data_points >= 10:
        confidence = "HIGH"
    elif data_points >= 5:
        confidence = "MEDIUM"

    if trend_data_points > 0:
        confidence_reason = f"based on {trend_data_points} days of consistent attendance data"
    else:
        confidence_reason = f"based on {data_points} available academic data points"
        
    predictions = {
        "future_risk_students": future_risk,
        # Kept for backward compatibility, but now derived from real aggregates.
        "expected_attendance": overall_attendance_pct,
        "expected_avg_marks": marks_summary.get("avg_mid2", 0.0),
        "confidence": confidence,
        "trend_direction": attendance_trend_direction,
        "confidence_reason": confidence_reason
    }

    # Weakest subject tracking (decision-support)
    # 1) Compute average marks per subject for this faculty filter
    subj_marks = {}
    mid1_marks = {}
    mid2_marks = {}

    for m in mark_records:
        if m.marks is None:
            continue
        try:
            val = float(m.marks)
        except Exception:
            continue

        sid = m.subject_id
        subj_marks.setdefault(sid, []).append(val)

        exam = str(m.exam or "").strip().lower()
        if "mid-1" in exam or "mid1" in exam:
            mid1_marks.setdefault(sid, []).append(val)
        elif "mid-2" in exam or "mid2" in exam:
            mid2_marks.setdefault(sid, []).append(val)

    weakest_subj_id = None
    weakest_avg = None
    for sid, vals in subj_marks.items():
        if not vals:
            continue
        avg = sum(vals) / len(vals)
        if weakest_avg is None or avg < weakest_avg:
            weakest_avg = avg
            weakest_subj_id = sid

    weakest_subj_name = subject_map.get(weakest_subj_id, "None") if weakest_subj_id is not None else "None"

    # 2) Determine whether weakest subject is trending down (mid1 -> mid2)
    weakest_trend = "stable"
    if weakest_subj_id is not None:
        m1_vals = mid1_marks.get(weakest_subj_id, [])
        m2_vals = mid2_marks.get(weakest_subj_id, [])
        if m1_vals and m2_vals:
            m1_avg = sum(m1_vals) / len(m1_vals)
            m2_avg = sum(m2_vals) / len(m2_vals)
            if (m1_avg - m2_avg) > 0:
                weakest_trend = "declining"

    weakest_subject = {"name": weakest_subj_name, "trend": weakest_trend}
    weakest_subject["reason"] = "lowest avg marks" + (" + declining trend in Mid 2" if weakest_trend == "declining" else "")
    weakest_subject["reason_lines"] = []
    if weakest_subj_id is not None:
        weakest_subject["reason_lines"].append("Lowest average marks among recorded assessments")
        if weakest_trend == "declining":
            weakest_subject["reason_lines"].append("Declining performance from Mid 1 to Mid 2")

    # Risk distribution based on the final per-student risk classification.
    risk_distribution = {"low": 0, "medium": 0, "high": 0}
    for s in students_list:
        lvl = (s.get("risk") or {}).get("level", "LOW")
        if lvl == "HIGH":
            risk_distribution["high"] += 1
        elif lvl == "MEDIUM":
            risk_distribution["medium"] += 1
        else:
            risk_distribution["low"] += 1

    risk_counts = {
        "HIGH": risk_distribution["high"],
        "MEDIUM": risk_distribution["medium"],
        "LOW": risk_distribution["low"],
    }

    mid_analysis = {
        "improved": improved_count,
        "declined": declined_count,
        "stable": same_count,
        "total": total_students,
        "avg_mid1": round(avg_mid1, 2),
        "avg_mid2": round(avg_mid2, 2),
        "trend": performance_trend
    }

    top_risks = []
    if mid_risk_count > 0:
        top_risks.append({
            "type": "Low Marks",
            "count": mid_risk_count,
            "reason": "Marks below 15 in Mid exams"
        })
    if len(attendance_risk_students) > 0:
        top_risks.append({
            "type": "Low Attendance",
            "count": len(attendance_risk_students),
            "reason": "Attendance below 75%"
        })

    try:
        _thr_ins = float(attendance_threshold)
    except Exception:
        _thr_ins = 75.0

    avg_for_narrative = chart_series_avg if attendance_values else overall_attendance_pct
    last_pct = round(float(attendance_values[-1]), 1) if attendance_values else round(overall_attendance_pct, 1)
    trend_period_label = "days" if (trend_view or "days") == "days" else ("weeks" if (trend_view or "days") == "weeks" else "months")

    low_attendance_days = 0
    for v in reversed(attendance_values):
        if float(v) < _thr_ins:
            low_attendance_days += 1
        else:
            break

    # ========== STRUCTURED INSIGHTS ENGINE ==========
    insights = []

    if attendance_trend_direction == "declining" and attendance_change_percent > 0 and avg_for_narrative < _thr_ins:
        insights.append({
            "title": "Critical Attendance Drop",
            "message": f"Attendance dropped by {attendance_change_percent}% and is below {_thr_ins:.0f}% (chart avg {round(avg_for_narrative, 1)}%).",
            "reason": f"Attendance fell from {round(float(attendance_values[0]), 1) if attendance_values else round(avg_for_narrative, 1)}% to {last_pct}% over the last {len(attendance_values) if attendance_values else 0} {trend_period_label}",
            "priority": "high",
            "action": "Schedule mandatory classes or send a clear attendance alert to the cohort.",
            "severity": "high"
        })
    elif avg_for_narrative < _thr_ins:
        insights.append({
            "title": "Attendance Below Safe Level",
            "message": f"Average attendance is {round(avg_for_narrative, 1)}%, under the {_thr_ins:.0f}% target ({attendance_trend_direction} trend in chart).",
            "reason": f"{low_attendance_days} consecutive {trend_period_label} are below {_thr_ins:.0f}% and current attendance is {last_pct}%",
            "priority": "high",
            "action": "Notify low-attendance students and plan one catch-up session this week.",
            "severity": "high"
        })

    mid_tracked = improved_count + declined_count + same_count
    if mid_tracked > 0:
        _prio = "high" if declined_count > improved_count else "medium"
        insights.append({
            "title": "Mixed Performance Trend",
            "message": f"Only {improved_count} out of {mid_tracked} students improved from Mid 1 to Mid 2.",
            "reason": f"{declined_count} students scored lower in Mid 2 than Mid 1, while only {improved_count} improved in the same exam cycle",
            "priority": _prio,
            "action": "Focus on weak students individually; review Mid 1 gaps before the next exam.",
            "severity": "high" if _prio == "high" else "medium"
        })
    elif mark_counts.get("mid2", 0) == 0 and mark_counts.get("mid1", 0) == 0:
        insights.append({
            "title": "Mid Exam Data Missing",
            "message": "No Mid 1/Mid 2 marks found for this cohort yet.",
            "reason": f"0 Mid records were found for the selected {trend_period_label} window, so Mid 1 vs Mid 2 comparison cannot be computed",
            "priority": "medium",
            "action": "Upload or verify Mid marks so performance comparisons can drive actions.",
            "severity": "medium"
        })

    if mid_risk_count > 0:
        insights.append({
            "title": "Top Risk: Low Mid Scores",
            "message": f"{mid_risk_count} students have Mid marks under 15.",
            "reason": f"{mid_risk_count} of {max(total_students, 1)} students are below 15 marks when comparing Mid 1 and Mid 2 results",
            "priority": "high",
            "action": "Run a short diagnostic quiz and assign targeted practice for those students.",
            "severity": "high"
        })

    if len(attendance_risk_students) >= 3:
        insights.append({
            "title": "Top Risk: Weak Attendance",
            "message": f"{len(attendance_risk_students)} students are under the {_thr_ins:.0f}% attendance bar.",
            "reason": f"{len(attendance_risk_students)} of {max(total_students, 1)} students are below {_thr_ins:.0f}% attendance across the selected time range",
            "priority": "medium",
            "action": "Send attendance nudges and offer optional office hours.",
            "severity": "medium"
        })

    if performance_trend == "declining" and mid_tracked > 0:
        insights.append({
            "title": "Class Mid Trend Declining",
            "message": f"Avg Mid 1 {round(avg_mid1, 1)} → Avg Mid 2 {round(avg_mid2, 1)}; cohort is slipping.",
            "reason": f"Class average dropped by {round(avg_mid1 - avg_mid2, 1)} points between Mid 1 and Mid 2 in the current cycle",
            "priority": "high",
            "action": "Re-teach the weakest modules and add a short formative assessment.",
            "severity": "high"
        })

    # Ensure every insight exposes severity + reason for UI + dedupe by title
    _seen_titles = set()
    _seen_signatures = set()
    _deduped = []
    for ins in insights:
        t = ins.get("title") or ""
        if t in _seen_titles:
            continue
        _seen_titles.add(t)
        _msg = str(ins.get("message") or "").lower()
        _sig = " ".join(_msg.split()[:10])
        if _sig and _sig in _seen_signatures:
            continue
        if _sig:
            _seen_signatures.add(_sig)
        if not ins.get("severity"):
            p = (ins.get("priority") or "low").lower()
            ins["severity"] = "high" if p == "high" else ("medium" if p == "medium" else "low")
        if not ins.get("reason"):
            ins["reason"] = "Based on available class data"
        _deduped.append(ins)
    insights = _deduped

    if len(insights) == 0:
        insights = [
            {
                "title": "Maintain Momentum",
                "message": marks_trend_summary or "No severe signals; keep monitoring attendance and Mid performance weekly.",
                "reason": f"0 high-priority drops detected in the last {len(attendance_values) if attendance_values else 0} {trend_period_label} compared to threshold {_thr_ins:.0f}%",
                "priority": "low",
                "action": "Keep the current plan and revisit this dashboard before the next evaluation.",
                "severity": "low"
            },
            {
                "title": "Data hygiene",
                "message": "Ensure attendance and Mid marks stay up to date so this view stays actionable.",
                "reason": f"Current dataset includes {data_points} records; weekly updates improve comparison quality across future windows",
                "priority": "low",
                "action": "Schedule a weekly five-minute review of this screen.",
                "severity": "low"
            },
        ]
    elif len(insights) == 1:
        insights.append({
            "title": "Secondary check",
            "message": "Review cohort Mid spread and attendance outliers even when primary signals are calm.",
            "reason": f"Only 1 primary signal was triggered in this {trend_period_label} window; additional comparison checks reduce blind spots",
            "priority": "low",
            "action": "Spot-check three borderline students this week.",
            "severity": "low"
        })

    insights = insights[:4]

    priority_students = [s for s in students_list if (s.get("risk") or {}).get("level") == "HIGH"]
    priority_summary = {
        "count": len(priority_students),
        "breakdown": {
            "low_attendance": sum(1 for s in priority_students if isinstance(s.get("attendance"), (int, float)) and float(s.get("attendance")) < _thr_ins),
            "low_marks": sum(1 for s in priority_students if isinstance(s.get("marks"), (int, float)) and float(s.get("marks")) < 15),
            "declining": sum(1 for s in priority_students if isinstance(s.get("mid1"), (int, float)) and isinstance(s.get("mid2"), (int, float)) and float(s.get("mid2")) < float(s.get("mid1"))),
        },
    }

    if avg_for_narrative < _thr_ins:
        att_sum_status = "critical"
    elif attendance_trend_direction == "declining":
        att_sum_status = "warning"
    else:
        att_sum_status = "healthy"

    if attendance_trend_direction == "declining" and attendance_change_percent > 0:
        att_sum_msg = f"Attendance dropped by {attendance_change_percent}%"
        if avg_for_narrative < _thr_ins:
            att_sum_msg += " and is below safe level"
    elif avg_for_narrative < _thr_ins:
        att_sum_msg = f"Average attendance is {round(avg_for_narrative, 1)}% and is below safe level"
    else:
        att_sum_msg = f"Average attendance is {round(avg_for_narrative, 1)}% ({attendance_trend_direction} trend)"

    # Decision-support trend insight + recommended actions
    try:
        trendInsight = "No attendance trend computed for this filter set."
        if attendance_values:
            lvl = "Below safe level" if last_pct < _thr_ins else "At or above safe level"
            if attendance_trend_direction == "declining" and attendance_change_percent > 0:
                trendInsight = (
                    f"Attendance dropped by {attendance_change_percent}% in the selected window. "
                    f"Current: {last_pct}% ({lvl})."
                )
            elif attendance_trend_direction == "improving":
                trendInsight = (
                    f"Attendance is improving in the selected window. "
                    f"Current: {last_pct}% ({lvl})."
                )
            else:
                trendInsight = f"Attendance is stable. Current: {last_pct}% ({lvl})."

        recommended_actions = []
        for ins in insights:
            a = (ins.get("action") or "").strip()
            if a and a not in recommended_actions:
                recommended_actions.append(a)
        if not recommended_actions:
            recommended_actions = ["Maintain current performance"]
    except Exception:
        trendInsight = "No sufficient data available for selected filters."
        recommended_actions = ["Maintain current performance"]

    response_data = {
        "students": students_list,
        "predictions": predictions,
        "weakest_subject": weakest_subject,
        "attendance_trend": attendance_trend,
        "trend_summary": trend_summary,
        "attendance_chart_annotation": attendance_chart_annotation,
        "attendance_trend_direction": attendance_trend_direction,
        "marks_trend_summary": marks_trend_summary,
        "mid_comparison_summary": mid_comparison_summary,
        "mid_analysis": mid_analysis,
        "mid_comparison": mid_comparison,
        "top_risks": top_risks,
        "risk_distribution": risk_distribution,
        "risk_counts": risk_counts,
        "insights": insights,
        "trendInsight": trendInsight,
        "recommended_actions": recommended_actions,
        "priority_summary": priority_summary,
        "filters": {"year": year, "section": section, "subject_id": subject_id, "subject_name": selected_subject_name},
        "faculty_scope": {
            "total_subjects": len(subject_ids),
            "assigned_classes": len(class_pairs),
            "total_students": total_students
        },
        "attendance_summary": {
            "overall_percentage": overall_attendance_pct,
            "present_count": present_count,
            "absent_count": absent_count,
            "total_records": total_attendance,
            "by_subject": attendance_subject_breakdown,
            "average": round(avg_for_narrative, 2),
            "threshold": round(_thr_ins, 1),
            "status": att_sum_status,
            "trend": attendance_trend_direction,
            "message": att_sum_msg
        },
        "assignment_summary": {
            "total_assignments": total_assignments,
            "submitted_count": submission_count,
            "late_submissions": late_submissions,
            "submission_rate": submission_rate
        },
        "marks_summary": marks_summary,
        "risk_summary": {
            "attendance_risk_count": len(attendance_risk_students),
            "marks_risk_count": marks_risk_count,
            "assignment_engagement_risk_count": assignment_engagement_risk_count
        },
        "top_students": top_students
    }

    try:
        recommendations = generate_recommendations(response_data)
    except Exception:
        recommendations = []

    response_data["recommendations"] = recommendations

    normalized_alerts = []
    high_risk_students = [s for s in students_list if str((s.get("risk") or {}).get("level") or "").upper() == "HIGH"]
    medium_risk_students = [s for s in students_list if str((s.get("risk") or {}).get("level") or "").upper() == "MEDIUM"]

    if high_risk_students:
        normalized_alerts.append({
            "title": "High risk students detected",
            "message": f"{len(high_risk_students)} students are in HIGH risk category.",
            "action": "Review high-risk students",
            "priority": "high",
            "type": "academic",
        })

    if medium_risk_students:
        normalized_alerts.append({
            "title": "Medium risk students need monitoring",
            "message": f"{len(medium_risk_students)} students are in MEDIUM risk category.",
            "action": "Track medium-risk students",
            "priority": "medium",
            "type": "academic",
        })

    if not normalized_alerts and students_list:
        normalized_alerts.append({
            "title": "No active academic alerts",
            "message": "No validated risk alerts were generated for the selected filters.",
            "action": "Continue monitoring",
            "priority": "low",
            "type": "academic",
        })

    response_data["alerts"] = normalized_alerts[:5]

    return response_data


@app.post("/intervention/update")
def update_student_intervention(
    payload: dict = Body(...),
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "faculty":
        raise HTTPException(status_code=403, detail="Faculty only")

    student_id = payload.get("student_id")
    status = str(payload.get("status") or "none").lower()
    intervention_type = payload.get("type")

    if student_id is None:
        raise HTTPException(status_code=400, detail="student_id is required")

    if status not in {"none", "planned", "done"}:
        raise HTTPException(status_code=400, detail="Invalid intervention status")

    if intervention_type is not None:
        intervention_type = str(intervention_type).strip().lower() or None
    if intervention_type not in {None, "call", "extra_class", "1on1"}:
        raise HTTPException(status_code=400, detail="Invalid intervention type")

    student = db.query(Student).filter(Student.student_id == student_id).first()
    if not student:
        raise HTTPException(status_code=404, detail="Student not found")

    if status == "none":
        student.intervention_status = "none"
        student.intervention_type = None
    else:
        student.intervention_status = status
        student.intervention_type = intervention_type

    student.intervention_last_updated = datetime.utcnow()
    db.commit()
    db.refresh(student)

    return {
        "student_id": student.student_id,
        "intervention": {
            "status": student.intervention_status or "none",
            "type": student.intervention_type,
            "last_updated": student.intervention_last_updated.isoformat() if student.intervention_last_updated else None,
        },
    }

@app.post("/faculty/apply-scaling")
async def apply_scaling(
    file: UploadFile = File(...),
    year: int = Form(...),
    section: str = Form(...),
    subject_id: int = Form(...),
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user["role"] != "faculty":
        raise HTTPException(status_code=403, detail="Faculty only")

    try:
        import pandas as pd  # type: ignore
        from io import BytesIO
        import tempfile
        import os
        from fastapi.responses import FileResponse  # type: ignore
    except ImportError:
        return {"error": "Dependencies not installed"}

    try:
        contents = await file.read()
        df = pd.read_excel(BytesIO(contents), dtype=str)
        
        df.columns = [str(col).strip().lower() for col in df.columns]

        def find_col(keyword):
            for col in df.columns:
                if keyword in col:
                    return col
            return None

        assign_cols = [col for col in df.columns if "assignment" in col and "total" not in col and "scaled" not in col]
        assign_scaled_col = find_col("assignment total")
        mid1_col = find_col("mid 1")
        mid1_scaled_col = find_col("mid 1 (scaled")
        mid2_col = find_col("mid 2")
        mid2_scaled_col = find_col("mid 2 (scaled")
        mid_combined_col = find_col("mid combined")
        internal_col = find_col("internal total")

        def safe(val):
            try:
                return float(val)
            except:
                return 0

        for idx, row in df.iterrows():
            assignment_sum = sum([safe(row[c]) for c in assign_cols])
            max_assign = len(assign_cols) * 10
            assignment_scaled = (assignment_sum / max_assign) * 10 if max_assign else 0

            mid1 = safe(row.get(mid1_col))
            mid2 = safe(row.get(mid2_col))

            mid1_scaled = (mid1 / 30) * 20 if mid1 else 0
            mid2_scaled = (mid2 / 30) * 20 if mid2 else 0

            valid_mids = [m for m in [mid1_scaled, mid2_scaled] if m > 0]
            if valid_mids:
                mid_combined = (sum(valid_mids) / (len(valid_mids) * 20)) * 20  # type: ignore
            else:
                mid_combined = 0

            internal_total = assignment_scaled + mid_combined

            if assign_scaled_col:
                df.at[idx, assign_scaled_col] = round(assignment_scaled, 2)  # type: ignore
            if mid1_scaled_col:
                df.at[idx, mid1_scaled_col] = round(mid1_scaled, 2)  # type: ignore
            if mid2_scaled_col:
                df.at[idx, mid2_scaled_col] = round(mid2_scaled, 2)  # type: ignore
            if mid_combined_col:
                df.at[idx, mid_combined_col] = round(mid_combined, 2)  # type: ignore
            if internal_col:
                df.at[idx, internal_col] = round(internal_total, 2)  # type: ignore

        temp_dir = tempfile.gettempdir()
        file_prefix = f"scaled_output_{subject_id}_{year}_{section}_{datetime.now().strftime('%Y%m%d%H%M%S')}"
        excel_path = os.path.join(temp_dir, f"{file_prefix}.xlsx")

        df.to_excel(excel_path, index=False)
        
        from models import ScalingLog  # type: ignore
        log = ScalingLog(
            faculty_id=current_user["user_id"],
            action_type="excel_scaling",
            year=year,
            section=section,
            subject_id=subject_id,
            file_name=f"{file_prefix}.xlsx",
            snapshot_data=[]
        )
        db.add(log)
        db.commit()

        return FileResponse(
            excel_path,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            filename="scaled_marks.xlsx",
            headers={"Access-Control-Expose-Headers": "Content-Disposition"}
        )

    except Exception as e:
        return {"error": f"Scaling failed: {str(e)}"}

@app.post("/faculty/undo-scaling")
def undo_scaling(
    req: ApplyScalingRequest,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    from models import ScaledMark, ScalingLog  # type: ignore
    if current_user["role"] != "faculty":
        raise HTTPException(status_code=403, detail="Faculty only")
        
    last_log = db.query(ScalingLog).filter(
        ScalingLog.year == req.year,
        ScalingLog.section == req.section,
        ScalingLog.subject_id == req.subject_id,
        ScalingLog.action_type.in_(["apply", "recalculate"])
    ).order_by(ScalingLog.timestamp.desc()).first()
    
    if not last_log:
        return {"error": "No previous scaling records found to undo."}
        
    snapshot = last_log.snapshot_data
    if snapshot is None:
        snapshot = []
        
    db.query(ScaledMark).filter(
        ScaledMark.subject_id == req.subject_id,
        ScaledMark.year == req.year,
        ScaledMark.section == req.section
    ).delete()
    
    for ss in snapshot:
        db.add(ScaledMark(
            student_id=ss["student_id"],
            subject_id=req.subject_id,
            year=req.year,
            section=req.section,
            assignment_scaled=ss.get("assignment_scaled"),
            mid1_scaled=ss.get("mid1_scaled"),
            mid2_scaled=ss.get("mid2_scaled"),
            mid_combined=ss.get("mid_combined"),
            internal_total=ss.get("internal_total"),
            semester_marks=ss.get("semester_marks"),
            final_total=ss.get("final_total")
        ))
        
    un_log = ScalingLog(
        faculty_id=current_user["user_id"],
        action_type="undo",
        year=req.year,
        section=req.section,
        subject_id=req.subject_id,
        snapshot_data=[]
    )
    db.add(un_log)
    db.commit()
    
    return {"message": "Undo successful. Reverted to previous scaling version."}

@app.get("/faculty/scaling-logs")
def get_scaling_logs(
    year: int = Query(...),
    section: str = Query(...),
    subject_id: int = Query(...),
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    try:
        from models import ScalingLog, User  # type: ignore
        logs = db.query(ScalingLog, User).join(
            User, ScalingLog.faculty_id == User.user_id
        ).filter(
            ScalingLog.year == year,
            ScalingLog.section == section,
            ScalingLog.subject_id == subject_id
        ).order_by(ScalingLog.timestamp.desc()).all()
        
        result = []
        for log, user in logs:
            result.append({
                "action_type": log.action_type,
                "faculty_name": user.name,
                "timestamp": log.timestamp,
                "file_name": log.file_name
            })
        return {"success": True, "logs": result}
    except Exception as e:
        return {"success": False, "error": str(e)}

